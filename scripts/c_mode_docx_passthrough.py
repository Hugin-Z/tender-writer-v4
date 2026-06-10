# -*- coding: utf-8 -*-
"""
c_mode_docx_passthrough.py · Word 上游原样切用(v2 单元 5)

目的:当招标文件上游是 docx 时,对"变量字段数量 ≤ 阈值"的简单 C 模式 Part,
     跳过 template+variables 流程,直接从原 docx 按 source_anchor 切对应段落
     复制成 filled.docx,只做字体归一化。

适用范围:
- tender_brief.json 的 source_meta.source_format == "docx"
- Part.production_mode == "C" 且 sub_mode == "C-template"
- variables.yaml 变量数 ≤ --max-vars(默认 5)
- source_anchor.type == "text"(table 型暂不支持)

用法:
    # 单 Part 判断是否适用并直切
    ./run_script.bat c_mode_docx_passthrough.py \\
        --project <项目> --part <N>

    # 批量:扫所有 C-template Part,能直切的直切,不能的打印原因
    ./run_script.bat c_mode_docx_passthrough.py \\
        --project <项目> --all

注意:
- 直切的 filled.docx 不走 docxtpl 模板替换,没有"【待填:xxx】"红色占位
- 保留原 docx 的段落结构 + 字体归一化(统一宋体、正文字号)
- 适合纯模板文字(签字盖章位、固定承诺段)这类变量极少的 Part
"""
from __future__ import annotations

import argparse
import re
import shutil
import sys
from copy import deepcopy
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    import yaml
except ImportError:
    print("[错误] 缺少依赖(python-docx / pyyaml)", file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent


def _safe_name(name: str) -> str:
    s = re.sub(r'^[一二三四五六七八九十\d]+[、.]\s*', '', name)
    s = re.sub(r'[(].+?[)]', '', s)
    s = re.sub(r'[(].+?[)]', '', s)
    return s.strip() or 'part'


def _normalize_run_font(run, east_asia="宋体", ascii_font="Times New Roman"):
    """字体归一化:sets eastAsia=宋体,ascii=Times New Roman。颜色保持不动。"""
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def _raw_lines_to_docx_para_index(raw_lines_for_ai: list, tender_docx: Path,
                                  start_line: int, end_line: int) -> tuple[int, int]:
    """
    把 raw.txt 行号范围映射到 docx 段落索引范围。

    策略:取 raw_lines[start_line].text 作为标识文本,在 tender_docx 的所有段落中
    找第一个含此文本的段落作为起点;end_line 同理(向前找最后一个匹配)。

    返回 (para_start_idx, para_end_idx),若找不到返回 (-1, -1)。
    """
    src_doc = Document(str(tender_docx))
    paragraphs = src_doc.paragraphs
    # 取 raw_lines 在 start_line 处的非空 text 作为起点标记
    start_marker = ""
    for ln in raw_lines_for_ai:
        if ln.get("line_no") == start_line:
            start_marker = (ln.get("text") or "").strip()[:30]
            break
    end_marker = ""
    for ln in raw_lines_for_ai:
        if ln.get("line_no") == end_line - 1:  # end_line 是 exclusive
            end_marker = (ln.get("text") or "").strip()[:30]
            break
    if not start_marker:
        return -1, -1

    start_idx = -1
    end_idx = -1
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if start_idx < 0 and start_marker and start_marker in text:
            start_idx = i
        if start_idx >= 0 and end_marker and end_marker in text:
            end_idx = i + 1  # exclusive
    if start_idx < 0:
        return -1, -1
    if end_idx < 0:
        end_idx = min(start_idx + 20, len(paragraphs))  # fallback: 20 段
    return start_idx, end_idx


def passthrough_part(project: str, part_idx: int, max_vars: int,
                     dry_run: bool) -> dict:
    """对单个 Part 做直切判断 + 执行。"""
    project_dir = ROOT / 'projects' / project
    import json
    brief_path = project_dir / 'output' / 'tender_brief.json'
    if not brief_path.exists():
        return {'status': 'error', 'msg': f'brief 缺失: {brief_path}'}
    brief = json.loads(brief_path.read_text(encoding='utf-8'))

    # 1. 源文件必须是 docx
    source_format = brief.get('source_meta', {}).get('source_format', '')
    if source_format != 'docx':
        return {'status': 'skipped',
                'msg': f'source_format={source_format!r}(非 docx,直切不适用)'}

    # 2. 定位 Part
    parts = brief.get('response_file_parts', [])
    if part_idx >= len(parts):
        return {'status': 'error', 'msg': f'Part 索引越界: {part_idx}'}
    part = parts[part_idx]
    name = part.get('name', '')
    mode = part.get('production_mode', '')
    sub = part.get('sub_mode', '')

    if mode != 'C' or sub != 'C-template':
        return {'status': 'skipped',
                'msg': f'mode={mode!r}/sub={sub!r}(非 C-template,直切不适用)'}

    # 3. source_anchor 必须是 text 型
    anchor = part.get('source_anchor', {}) or {}
    if anchor.get('type') != 'text':
        return {'status': 'skipped',
                'msg': f'source_anchor.type={anchor.get("type")!r}'
                       f'(table 型直切不支持,需走老链路)'}

    # 4. 变量数必须 ≤ max_vars
    part_dir = project_dir / 'output' / 'c_mode' / _safe_name(name)
    vars_path = part_dir / 'variables.yaml'
    if vars_path.exists():
        with vars_path.open(encoding='utf-8') as f:
            vars_data = yaml.safe_load(f) or {}
        n_vars = len(vars_data.get('variables', []))
    else:
        n_vars = 0  # 还没产出 variables.yaml,无法判断,默认 0

    if n_vars > max_vars:
        return {'status': 'skipped',
                'msg': f'变量数 {n_vars} > 阈值 {max_vars}(走老链路 template+fill)'}

    # 5. 找源 docx 路径
    source_file = brief.get('source_file', '')
    tender_docx = ROOT / source_file if not Path(source_file).is_absolute() \
        else Path(source_file)
    if not tender_docx.exists():
        return {'status': 'error', 'msg': f'源 docx 不存在: {tender_docx}'}

    # 6. 映射行号到段落索引
    raw_lines = brief.get('raw_lines_for_ai', [])
    start_line = anchor.get('start_line', 0)
    end_line = anchor.get('end_line', 0)
    para_start, para_end = _raw_lines_to_docx_para_index(
        raw_lines, tender_docx, start_line, end_line
    )
    if para_start < 0:
        return {'status': 'error',
                'msg': f'无法从 raw_lines 行号 {start_line}-{end_line} 映射到 docx 段落'}

    if dry_run:
        return {'status': 'ok',
                'msg': f'直切可行:变量 {n_vars} 个 ≤ {max_vars};'
                       f'docx 段落 {para_start}-{para_end}(共 {para_end - para_start} 段)'}

    # 7. 执行直切:从源 docx 复制段落到新 filled.docx
    src_doc = Document(str(tender_docx))
    out_doc = Document()

    # 先清空 out_doc 默认空段落
    while len(out_doc.paragraphs) > 0:
        p = out_doc.paragraphs[0]
        p._element.getparent().remove(p._element)

    body = out_doc.element.body
    for i in range(para_start, para_end):
        if i >= len(src_doc.paragraphs):
            break
        src_p = src_doc.paragraphs[i]
        # 深拷贝段落元素
        new_p = deepcopy(src_p._element)
        body.append(new_p)

    # 字体归一化:扫所有 run 设 eastAsia=宋体, ascii=Times New Roman
    for p in out_doc.paragraphs:
        for run in p.runs:
            _normalize_run_font(run)

    # 保存到 Part 目录的 filled.docx
    filled_path = part_dir / 'filled.docx'
    part_dir.mkdir(parents=True, exist_ok=True)
    out_doc.save(str(filled_path))

    return {'status': 'ok',
            'msg': f'直切完成 → {filled_path}(段落 {para_end - para_start} 段,'
                   f'变量 {n_vars} 个)'}


def main():
    p = argparse.ArgumentParser(
        description='Word 上游原样切用(v2 单元 5)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    p.add_argument('--project', required=True)
    grp = p.add_mutually_exclusive_group(required=True)
    grp.add_argument('--part', type=int)
    grp.add_argument('--all', action='store_true')
    p.add_argument('--max-vars', type=int, default=5,
                   help='变量数阈值(≤ 此值适用直切),默认 5')
    p.add_argument('--dry-run', action='store_true',
                   help='仅判断哪些 Part 可直切,不实际执行')
    args = p.parse_args()

    import json
    brief_path = ROOT / 'projects' / args.project / 'output' / 'tender_brief.json'
    if not brief_path.exists():
        print(f"[错误] 找不到 brief", file=sys.stderr)
        sys.exit(1)

    # 用函数级闸门
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from brief_schema import load_brief_guarded
    brief = load_brief_guarded(brief_path)
    parts = brief.get('response_file_parts', [])

    if args.all:
        target = list(range(len(parts)))
    else:
        target = [args.part]

    results = []
    for idx in target:
        res = passthrough_part(args.project, idx, args.max_vars, args.dry_run)
        results.append((idx, parts[idx].get('name', '')
                        if idx < len(parts) else '<OOB>', res))

    print()
    print("=" * 60)
    print(f"Word 上游直切汇总 ({args.project}, dry_run={args.dry_run})")
    print("=" * 60)
    ok_cnt = sum(1 for _, _, r in results if r['status'] == 'ok')
    skip_cnt = sum(1 for _, _, r in results if r['status'] == 'skipped')
    err_cnt = sum(1 for _, _, r in results if r['status'] == 'error')
    for idx, name, r in results:
        icon = {'ok': '[OK]', 'error': '[ERR]', 'skipped': '[SKIP]'}[r['status']]
        print(f"  {icon} Part[{idx}] {name[:30]:30s}  {r['msg']}")
    print()
    print(f"总计: {ok_cnt} 直切 / {skip_cnt} 跳过 / {err_cnt} 失败")
    if err_cnt > 0:
        sys.exit(1)


if __name__ == '__main__':
    main()
