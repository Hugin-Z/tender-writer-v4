# -*- coding: utf-8 -*-
"""
c_mode_extract.py · C 模式模板提取脚本（阶段 4-C Step 1-3.5）

用法:
    # Step 1-2: 提取原文,打印并写入 raw.txt
    python scripts/c_mode_extract.py --project <项目名> --part <索引> --extract-text

    # Step 3.5: 从人工/AI 产出的 intermediate.json 构建 template.docx + variables.yaml
    python scripts/c_mode_extract.py --project <项目名> --part <索引> \\
        --build-from-json <intermediate.json 路径>

说明:
    Step 3 (AI 语义识别) 本轮由用户手工完成——把 raw.txt 发给 AI,按
    schema 输出 intermediate.json,粘贴回来。LLM API 集成留给后续轮次。
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("[错误] 缺少 python-docx 依赖,请先双击 install.bat 安装依赖。",
          file=sys.stderr)
    sys.exit(1)

try:
    import yaml
except ImportError:
    print("[错误] 缺少 PyYAML 依赖,请先双击 install.bat 安装依赖。",
          file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent


def _safe_name(name: str) -> str:
    """把 Part 名转成文件系统安全的目录名。"""
    s = re.sub(r'^[一二三四五六七八九十\d]+[、.]\s*', '', name)
    s = re.sub(r'[（(].+?[)）]', '', s)
    s = s.strip()
    return s or 'part'


def load_brief(project_dir: Path) -> dict:
    brief_path = project_dir / 'output' / 'tender_brief.json'
    if not brief_path.exists():
        print(f"[错误] 找不到 tender_brief.json: {brief_path}", file=sys.stderr)
        sys.exit(1)
    return json.loads(brief_path.read_text(encoding='utf-8'))


def get_c_part(brief: dict, part_index: int) -> dict:
    parts = brief.get('response_file_parts', [])
    if part_index < 0 or part_index >= len(parts):
        print(f"[错误] --part {part_index} 超出范围(0..{len(parts)-1})", file=sys.stderr)
        sys.exit(1)
    part = parts[part_index]
    mode = part.get('production_mode')
    if mode != 'C':
        print(f"[错误] Part[{part_index}] '{part.get('name')}' 的生产模式是 {mode},不是 C。",
              file=sys.stderr)
        sys.exit(1)
    return part


def _render_table_as_md(table: dict) -> str:
    """把 TableObject 渲染为伪 markdown 表格(便于 AI 读)。"""
    lines = []
    lines.append(f"# table_id: {table['table_id']}")
    lines.append(f"# evidence: {table['evidence']}")
    headers = table.get('headers', [])
    rows = table.get('rows', [])
    if headers:
        lines.append('| ' + ' | '.join(headers) + ' |')
        lines.append('|' + '|'.join(['---'] * len(headers)) + '|')
    for row in rows:
        lines.append('| ' + ' | '.join(row) + ' |')
    return '\n'.join(lines)


def cmd_extract_text(args, project_dir: Path):
    # V56: 走 resolve_source_anchor 统一入口,支持 text 和 table 两种 type
    from brief_schema import resolve_source_anchor, validate_source_anchor

    brief = load_brief(project_dir)
    part = get_c_part(brief, args.part)
    part_dir_name = _safe_name(part['name'])
    out_dir = project_dir / 'output' / 'c_mode' / part_dir_name
    out_dir.mkdir(parents=True, exist_ok=True)

    anchor = part.get('source_anchor')
    if not anchor:
        print(f"[错误] Part[{args.part}] '{part.get('name')}' 缺少 source_anchor",
              file=sys.stderr)
        sys.exit(1)

    try:
        anchor_type = validate_source_anchor(anchor)
        resolved = resolve_source_anchor(brief, anchor)
    except ValueError as e:
        print(f"[错误] source_anchor 校验失败: {e}", file=sys.stderr)
        sys.exit(1)

    raw_path = out_dir / 'raw.txt'

    if anchor_type == 'text':
        # resolved: [(line_no, text), ...]
        text = '\n'.join(t for _, t in resolved)
        raw_path.write_text(text, encoding='utf-8')
        start = anchor.get('start_line')
        end = anchor.get('end_line')
        print(f"[完成] Part[{args.part}] '{part['name']}' (text 型) 原文已写入: {raw_path}")
        print(f"  行号范围: {start}-{end}")
        print(f"  行数: {len(resolved)}")
    else:
        # table 型: resolved = [TableObject, ...]
        sections = [_render_table_as_md(t) for t in resolved]
        text = '\n\n'.join(sections)
        raw_path.write_text(text, encoding='utf-8')
        table_ids = anchor.get('table_ids', [])
        print(f"[完成] Part[{args.part}] '{part['name']}' (table 型) 原文已写入: {raw_path}")
        print(f"  table_ids: {table_ids}")
        print(f"  表格数: {len(resolved)}")

    print()
    print("=" * 60)
    print("下一步(Step 3):把 raw.txt 内容发给 AI(Claude/ChatGPT),附上")
    print("SKILL.md 阶段 4-C 中的 prompt。AI 产出 intermediate.json,")
    print(f"保存到: {out_dir / 'intermediate.json'}")
    print()
    print("然后运行:")
    print(f"  python scripts/c_mode_extract.py --project {args.project} "
          f"--part {args.part} \\")
    print(f"    --build-from-json {out_dir / 'intermediate.json'}")
    print("=" * 60)
    print()
    print("--- raw.txt 内容 ---")
    print(text)


def _render_table_block(doc, block: dict) -> None:
    """V56.3/C.3: 渲染 table block 为 Word 原生表格。"""
    headers = block.get('headers', [])
    rows = block.get('rows', [])
    n_cols = len(headers)
    if n_cols == 0:
        raise ValueError("table block headers 为空")
    bad = [i for i, r in enumerate(rows) if len(r) != n_cols]
    if bad:
        raise ValueError(
            f"table block rows 列数不一致: headers={n_cols}, "
            f"异常行 index={bad}"
        )

    table = doc.add_table(rows=len(rows) + 1, cols=n_cols)
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        table.cell(0, i).text = str(header)

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, cell_spec in enumerate(row):
            cell_type = cell_spec.get('type')
            if cell_type == 'fixed':
                content = cell_spec.get('content', '')
            elif cell_type == 'variable':
                content = f"{{{{{cell_spec['name']}}}}}"
            else:
                raise ValueError(f"unknown cell type: {cell_type}")
            table.cell(row_idx, col_idx).text = content


def render_template_docx(blocks: list[dict], out_path: Path) -> int:
    """按中间 JSON 的 blocks 渲染 template.docx。

    文本 block → 段落(docxtpl 渲染时保持原样)
    变量 block → 在段落中用 {{name}} 占位
    text 内相邻的 variable 会被合并到同一段落。

    v2 补丁 2(字体修复):调用 apply_default_styles 为 Normal / Heading 样式
    写入宋体(eastAsia)+ Times New Roman(ascii)。不调用会导致 run fallback 到
    docDefaults 的 minorEastAsia 主题字体,WPS/Office 在主题文件缺失时 fallback
    到系统日文字体 MS 明朝。
    """
    from docx_builder import apply_default_styles

    doc = Document()
    apply_default_styles(doc)
    para_count = 0
    current_para_texts: list[str] = []

    def flush_para():
        nonlocal para_count
        if current_para_texts:
            text = ''.join(current_para_texts)
            if text.strip():
                doc.add_paragraph(text)
                para_count += 1
            current_para_texts.clear()

    for blk in blocks:
        t = blk.get('type')
        if t == 'text':
            content = blk.get('content', '')
            # 按 \n 切段落
            parts = content.split('\n')
            for i, part in enumerate(parts):
                current_para_texts.append(part)
                if i < len(parts) - 1:
                    flush_para()
        elif t == 'variable':
            name = blk.get('name', '').strip()
            if not name:
                continue
            current_para_texts.append(f'{{{{{name}}}}}')
        elif t == 'table':
            # C.3: flush pending text, render Word native table
            flush_para()
            _render_table_block(doc, blk)
        else:
            # 未知 block 类型,硬失败
            raise ValueError(f"未知 block type: {t}")
    flush_para()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return para_count


def _spec_to_var_entry(spec: dict) -> dict | None:
    """把一个 variable 规格(block 或 table cell)转为 variables.yaml 条目。"""
    name = (spec.get('name') or '').strip()
    if not name:
        return None
    return {
        'name': name,
        'description': spec.get('description', ''),
        'source': spec.get('suggested_source', 'manual'),
        'path': spec.get('suggested_path', ''),
        'required': True,
    }


def build_variables_yaml(blocks: list[dict], part: dict, part_index: int) -> dict:
    """从 blocks 构建 variables.yaml 结构(支持 table block 内的 variable cell)。"""
    anchor = part.get('source_anchor', {})
    vars_list = []
    seen: set[str] = set()

    def _add(spec):
        entry = _spec_to_var_entry(spec)
        if entry and entry['name'] not in seen:
            seen.add(entry['name'])
            vars_list.append(entry)

    for blk in blocks:
        t = blk.get('type')
        if t == 'variable':
            _add(blk)
        elif t == 'table':
            # C.3: 深入 table cell 收集 variable
            for row in blk.get('rows', []):
                for cell in row:
                    if cell.get('type') == 'variable':
                        _add(cell)
        # text block 不产生变量

    return {
        'template_name': _safe_name(part['name']),
        'part_index': part_index,
        'part_name': part['name'],
        'source_lines': [anchor.get('start_line'), anchor.get('end_line')],
        'variables': vars_list,
    }


def _build_instructions_md(part: dict, part_index: int, data: dict, out_dir: Path) -> None:
    """V60 C-reference: 渲染 instructions.md (YAML front matter + markdown 正文)。

    intermediate.json 结构(C-reference 用):
        {
          "front_matter": {  # 必填: part_name, sub_mode, source_anchor; 推荐: 5 项
              "part_name": "...",
              "sub_mode": "C-reference",
              "source_anchor": {...},
              "production_channel": "...",       # 推荐
              "operation_steps": [...],          # 推荐
              "inputs_required": [...],          # 推荐
              "dependencies": [...],             # 推荐
              "caveats": [...]                   # 推荐
          },
          "body_markdown": "# ...\\n..."          # AI 自由组织正文
        }
    """
    fm = data.get('front_matter')
    body = data.get('body_markdown', '')
    if not fm:
        print(f"[错误] C-reference intermediate.json 缺少 front_matter", file=sys.stderr)
        sys.exit(1)
    # 必填校验
    for key in ('part_name', 'sub_mode', 'source_anchor'):
        if key not in fm:
            print(f"[错误] front_matter 缺少必填字段 '{key}'", file=sys.stderr)
            sys.exit(1)
    if fm['sub_mode'] != 'C-reference':
        print(f"[错误] front_matter.sub_mode='{fm['sub_mode']}',"
              f"应为 'C-reference'", file=sys.stderr)
        sys.exit(1)

    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / 'instructions.md'

    fm_yaml = yaml.safe_dump(fm, allow_unicode=True, sort_keys=False).rstrip()
    content = f"---\n{fm_yaml}\n---\n\n{body}\n"
    out_path.write_text(content, encoding='utf-8')

    print(f"[完成] Part[{part_index}] '{part['name']}' (C-reference)")
    print(f"  instructions.md: {out_path}")
    print(f"  front_matter 字段数: {len(fm)}")
    print(f"  body markdown 字符数: {len(body)}")
    print()
    print("=" * 60)
    print("下一步: 用户 review instructions.md, 确认无误。C-reference 不需要 fill,")
    print("c_mode_fill 仅做存在性校验。整标合并器(未启动)按 instructions 处理。")
    print("=" * 60)


def cmd_build_from_json(args, project_dir: Path):
    # V60: sub_mode 分流
    from brief_schema import resolve_sub_mode

    brief = load_brief(project_dir)
    part = get_c_part(brief, args.part)
    sub_mode = resolve_sub_mode(part)

    if sub_mode == 'C-attachment':
        raise NotImplementedError(
            f"Part[{args.part}] '{part['name']}' sub_mode=C-attachment 当前挂档,"
            f"见 business_model §8 #N20"
        )

    part_dir_name = _safe_name(part['name'])
    out_dir = project_dir / 'output' / 'c_mode' / part_dir_name

    json_path = Path(args.build_from_json)
    if not json_path.exists():
        print(f"[错误] intermediate.json 不存在: {json_path}", file=sys.stderr)
        sys.exit(1)

    data = json.loads(json_path.read_text(encoding='utf-8'))

    if sub_mode == 'C-reference':
        # C-reference: intermediate.json 携带 instructions.md 内容信息
        # 直接将其渲染为 instructions.md(YAML front matter + markdown 正文)
        _build_instructions_md(part, args.part, data, out_dir)
        return

    # 以下为 C-template 路径(默认)
    blocks = data.get('blocks', [])
    if not blocks:
        print(f"[错误] intermediate.json 中 blocks 为空", file=sys.stderr)
        sys.exit(1)

    text_count = sum(1 for b in blocks if b.get('type') == 'text')
    var_count = sum(1 for b in blocks if b.get('type') == 'variable')
    table_count = sum(1 for b in blocks if b.get('type') == 'table')
    table_var_cells = sum(
        1 for b in blocks if b.get('type') == 'table'
        for row in b.get('rows', [])
        for cell in row if cell.get('type') == 'variable'
    )

    template_path = out_dir / 'template.docx'
    vars_path = out_dir / 'variables.yaml'

    para_count = render_template_docx(blocks, template_path)

    vars_data = build_variables_yaml(blocks, part, args.part)
    with open(vars_path, 'w', encoding='utf-8') as f:
        yaml.safe_dump(vars_data, f, allow_unicode=True, sort_keys=False)

    print(f"[完成] Part[{args.part}] '{part['name']}'")
    print(f"  template.docx: {template_path} (段落 {para_count} 个)")
    print(f"  variables.yaml: {vars_path} (变量 {len(vars_data['variables'])} 个)")
    print(f"  intermediate: text blocks {text_count}, variable blocks {var_count}, "
          f"table blocks {table_count} (含 {table_var_cells} 个变量单元格)")
    print()
    print("=" * 60)
    print("下一步(Step 4):用户 review template.docx 和 variables.yaml,")
    print("确认无误后执行 c_mode_fill.py:")
    print(f"  python scripts/c_mode_fill.py --project {args.project} --part {args.part}")
    print("=" * 60)


def main():
    parser = argparse.ArgumentParser(
        description='C 模式模板提取(阶段 4-C Step 1-3.5)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument('--project', required=True, help='项目名(projects/ 下子目录名)')
    parser.add_argument('--part', type=int, required=True,
                        help='response_file_parts 索引(0 基)')
    g = parser.add_mutually_exclusive_group(required=True)
    g.add_argument('--extract-text', action='store_true',
                   help='Step 1-2: 提取原文到 raw.txt')
    g.add_argument('--build-from-json',
                   help='Step 3.5: 从 intermediate.json 构建 template + variables')
    args = parser.parse_args()

    project_dir = ROOT / 'projects' / args.project
    if not project_dir.exists():
        print(f"[错误] 项目目录不存在: {project_dir}", file=sys.stderr)
        sys.exit(1)

    # V53: 未 review 则硬失败
    from brief_schema import ensure_reviewed
    ensure_reviewed(project_dir / 'output')

    if args.extract_text:
        cmd_extract_text(args, project_dir)
    else:
        cmd_build_from_json(args, project_dir)


if __name__ == '__main__':
    main()
