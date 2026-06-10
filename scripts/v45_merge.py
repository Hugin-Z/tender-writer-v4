# -*- coding: utf-8 -*-
"""
v45_merge.py · V45 整标合并器(V62 落地)

职责:按招标文件第六章响应文件格式顺序,合并 A/B/C 各 Part 产出 →
     final_tender_package/{final_response.docx, operations_checklist.md, pending_manual_work.md}

合并顺序由 parse_tender 从 response_file_parts 动态推导(v2 已从硬编码改为动态)。

C-reference 处理:不进 final_response.docx,内容转 operations_checklist.md。
.pending_marker 处理:占位照合,附加到 pending_manual_work.md。

用法:
    python scripts/v45_merge.py --project demo_cadre_training
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docxcompose.composer import Composer
except ImportError:
    print("[错误] 缺少依赖(python-docx / docxcompose)。请先 pip install", file=sys.stderr)
    sys.exit(1)

try:
    import yaml
except ImportError:
    print("[错误] 缺少 PyYAML", file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent


def _safe_name(name: str) -> str:
    s = re.sub(r'^[一二三四五六七八九十\d]+[、.]\s*', '', name)
    s = re.sub(r'[（(].+?[)）]', '', s)
    s = s.strip()
    return s or 'part'


# v2 单元 7 P1 修复:动态生成合并顺序,不再硬编码另一项目的 Part 映射。
# 每项: (part_index 或 None, 来源类型, 描述)
# 来源类型: c_template / c_reference / b_mode / a_mode / inapplicable


def build_merge_order(brief: dict) -> list[tuple[int | None, str, str]]:
    """
    按 response_file_parts 顺序 + 每个 Part 的 production_mode/sub_mode
    动态生成合并顺序。
    """
    parts = brief.get('response_file_parts', [])
    order: list[tuple[int | None, str, str]] = []

    # Part 11 技术部分(A 模式)只出现一次(tender_response.docx),避免重复追加
    a_mode_appended = False

    for i, p in enumerate(parts):
        name = p.get('name', f'Part {i}')
        mode = (p.get('production_mode') or '').strip().upper()
        sub = (p.get('sub_mode') or '').strip()

        if mode == 'A':
            if a_mode_appended:
                # A 模式 Part 全部来源同一份 tender_response.docx,不重复
                continue
            order.append((i, 'a_mode', name))
            a_mode_appended = True
        elif mode == 'B':
            order.append((i, 'b_mode', name))
        elif mode == 'C':
            if sub == 'C-reference':
                order.append((i, 'c_reference', name))
            elif sub == 'C-template':
                order.append((i, 'c_template', name))
            elif sub == 'C-attachment':
                # #N20 挂档类型,按 SKILL.md 不进 v1.1
                print(f"[信息] Part[{i}] '{name}' sub_mode=C-attachment,"
                      "合并器跳过(#N20 挂档)", file=sys.stderr)
                continue
            else:
                print(f"[警告] Part[{i}] '{name}' production_mode=C 但 sub_mode 未识别:"
                      f"{sub!r},合并器跳过", file=sys.stderr)
                continue
        elif mode == 'D':
            # D 模式不在 v1.1 范围
            order.append((i, 'inapplicable', name))
        elif mode in ('', '不适用', 'N/A', 'NONE'):
            order.append((i, 'inapplicable', name))
        else:
            print(f"[警告] Part[{i}] '{name}' production_mode={mode!r} 未识别,"
                  f"按 inapplicable 处理", file=sys.stderr)
            order.append((i, 'inapplicable', name))

    return order


def _get_part(brief: dict, idx: int | None) -> dict | None:
    if idx is None:
        return None
    parts = brief.get('response_file_parts', [])
    if 0 <= idx < len(parts):
        return parts[idx]
    return None


def _create_inapplicable_doc(tmp_dir: Path, title: str) -> Path:
    """产生「不适用」占位 docx。

    v2 补丁 2:divider / inapplicable 是 composer.append 的源,其中第一个
    (divider)会成为合并 master,其 styles.xml 决定 final_response.docx 的
    Normal 样式。必须调 apply_default_styles 让 Normal 样式有字体。
    """
    from docx_builder import apply_default_styles
    doc = Document()
    apply_default_styles(doc)
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    p2 = doc.add_paragraph()
    run2 = p2.add_run('本项目不适用。')
    run2.italic = True
    out = tmp_dir / f"_inapplicable_{title.replace('/', '_')[:20]}.docx"
    doc.save(str(out))
    return out


def _create_part_divider(tmp_dir: Path, title: str) -> Path:
    """产生 Part 分隔页(标题段)。保证合并后仍能看出 Part 边界。

    v2 补丁 2:divider 通常是第一个 append 进 master 的文档,其 styles.xml
    决定合并后的 Normal 样式。必须调 apply_default_styles。
    """
    from docx_builder import apply_default_styles
    doc = Document()
    apply_default_styles(doc)
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    out = tmp_dir / f"_divider_{title.replace('/', '_')[:20]}.docx"
    doc.save(str(out))
    return out


def _render_c_reference_section(part: dict, instructions_path: Path) -> str:
    """读 C-reference instructions.md,提炼为 operations_checklist markdown 段落。"""
    content = instructions_path.read_text(encoding='utf-8')
    # 分离 front matter 和 body
    fm = {}
    body = content
    if content.startswith('---\n'):
        end = content.find('\n---\n', 4)
        if end > 0:
            fm_text = content[4:end]
            body = content[end + 5:]
            try:
                fm = yaml.safe_load(fm_text) or {}
            except yaml.YAMLError:
                fm = {}

    lines = [f"## {part.get('name', '<no name>')}"]
    lines.append('')
    pc = fm.get('production_channel')
    if pc:
        lines.append(f"**产出渠道**: {pc}")
        lines.append('')
    steps = fm.get('operation_steps') or []
    if steps:
        lines.append('**操作步骤**:')
        for i, s in enumerate(steps, 1):
            lines.append(f"{i}. {s}")
        lines.append('')
    inputs = fm.get('inputs_required') or []
    if inputs:
        lines.append('**依赖输入**:')
        for it in inputs:
            lines.append(f"- {it}")
        lines.append('')
    deps = fm.get('dependencies') or []
    if deps:
        lines.append('**前置依赖**:')
        for d in deps:
            lines.append(f"- {d}")
        lines.append('')
    caveats = fm.get('caveats') or []
    if caveats:
        lines.append('**注意事项**:')
        for c in caveats:
            lines.append(f"- {c}")
        lines.append('')
    lines.append('---')
    lines.append('')
    lines.append('### 原 instructions.md 正文')
    lines.append('')
    lines.append(body.strip())
    return '\n'.join(lines)


def main():
    parser = argparse.ArgumentParser(description='V45 整标合并器(V62)')
    parser.add_argument('--project', required=True, help='项目名')
    args = parser.parse_args()

    project_dir = ROOT / 'projects' / args.project

    # V3-3: timing hook
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from _timing_hook import stage_timer

    with stage_timer("v45_merge", project_dir):
        _main_body(args, project_dir)


def _main_body(args, project_dir):
    # V53 gate
    from brief_schema import ensure_reviewed
    ensure_reviewed(project_dir / 'output')

    output_dir = project_dir / 'output'
    with open(output_dir / 'tender_brief.json', 'r', encoding='utf-8') as f:
        brief = json.load(f)

    pkg_dir = project_dir / 'final_tender_package'
    pkg_dir.mkdir(parents=True, exist_ok=True)

    # 临时目录(不进 tracked_outputs)
    import tempfile
    tmp_dir = Path(tempfile.gettempdir()) / f"v45_merge_{args.project}"
    tmp_dir.mkdir(parents=True, exist_ok=True)

    # 用 Composer 合并:以第一个文档为 master,其他追加
    final_response_path = pkg_dir / 'final_response.docx'
    operations_lines = [
        '# 操作指引(operations_checklist)',
        '',
        '以下章节由外部系统(电子平台等)生成,不在 `final_response.docx` 中出现。',
        '投标方需按下列操作清单在对应平台完成交付。',
        '',
    ]
    pending_entries = []

    # 收集每项合并片段
    fragment_paths: list[tuple[str, Path]] = []

    merge_order = build_merge_order(brief)
    print(f"[信息] 动态生成合并顺序(共 {len(merge_order)} 项,"
          f"来源:tender_brief.response_file_parts)", file=sys.stderr)
    c_ref_written = 0
    for part_idx, kind, title in merge_order:
        part = _get_part(brief, part_idx)
        part_name = part.get('name') if part else title
        part_dir_name = _safe_name(part_name) if part else _safe_name(title)

        if kind == 'c_template':
            src = output_dir / 'c_mode' / part_dir_name / 'filled.docx'
            if not src.exists():
                print(f"[警告] C-template filled.docx 缺失: {src}", file=sys.stderr)
                continue
            fragment_paths.append((title, src))

        elif kind == 'c_reference':
            instructions = output_dir / 'c_mode' / part_dir_name / 'instructions.md'
            if not instructions.exists():
                print(f"[警告] C-reference instructions.md 缺失: {instructions}", file=sys.stderr)
                continue
            # 不进 final_response.docx(见 #N22)
            section = _render_c_reference_section(part, instructions)
            operations_lines.append(section)
            operations_lines.append('')
            c_ref_written += 1
            print(f"[信息] {title} → operations_checklist.md(不进 final_response.docx)")

        elif kind == 'b_mode':
            src = output_dir / 'b_mode' / part_dir_name / 'assembled.docx'
            if not src.exists():
                print(f"[警告] B 模式 assembled.docx 缺失: {src}", file=sys.stderr)
                continue
            fragment_paths.append((title, src))
            # 检测 .pending_marker
            marker = output_dir / 'b_mode' / part_dir_name / '.pending_marker'
            if marker.exists():
                pending_entries.append({
                    'part_title': title,
                    'part_dir': str(part_dir_name),
                    'marker_path': str(marker.relative_to(project_dir)),
                    'assembled_path': str(src.relative_to(project_dir)),
                })

        elif kind == 'a_mode':
            src = output_dir / 'tender_response.docx'
            if not src.exists():
                print(f"[警告] A 模式 tender_response.docx 缺失: {src}", file=sys.stderr)
                continue
            fragment_paths.append((title, src))

        elif kind == 'inapplicable':
            frag = _create_inapplicable_doc(tmp_dir, title)
            fragment_paths.append((title, frag))

    if not fragment_paths:
        print("[错误] 没有可合并的片段", file=sys.stderr)
        sys.exit(1)

    # 创建第一个分隔页作为 master,后续追加分隔页+片段
    master_title, master_frag = fragment_paths[0]
    master_divider = _create_part_divider(tmp_dir, master_title)
    master = Document(str(master_divider))
    composer = Composer(master)
    composer.append(Document(str(master_frag)))

    for title, frag in fragment_paths[1:]:
        divider = _create_part_divider(tmp_dir, title)
        composer.append(Document(str(divider)))
        composer.append(Document(str(frag)))

    composer.save(str(final_response_path))

    # 写 operations_checklist.md
    ops_path = pkg_dir / 'operations_checklist.md'
    ops_path.write_text('\n'.join(operations_lines) + '\n', encoding='utf-8')

    # 写 pending_manual_work.md
    pending_lines = [
        '# 待人工填充清单(pending_manual_work)',
        '',
        'V45 合并器检测到以下 Part 仍带 `.pending_marker`,即占位内容尚未被',
        '供应商替换为真实内容。投标方交付前请逐项核对,完成后手工删除 marker。',
        '',
    ]
    if pending_entries:
        for e in pending_entries:
            pending_lines.append(f"## {e['part_title']}")
            pending_lines.append('')
            pending_lines.append(f"- 产出目录: `{e['part_dir']}`")
            pending_lines.append(f"- assembled 路径: `{e['assembled_path']}`")
            pending_lines.append(f"- marker 路径: `{e['marker_path']}`")
            pending_lines.append('- 完成填充后: 手工删除 `.pending_marker`')
            pending_lines.append('')
    else:
        pending_lines.append('当前无待填充 Part(所有 B 模式产出 marker 已清除)。')
        pending_lines.append('')

    pending_path = pkg_dir / 'pending_manual_work.md'
    pending_path.write_text('\n'.join(pending_lines), encoding='utf-8')

    # 简报
    print()
    print("=" * 60)
    print(f"[完成] V45 整标合并完成 → {pkg_dir}")
    print(f"  final_response.docx:        {final_response_path}"
          f" ({final_response_path.stat().st_size:,} bytes)")
    print(f"  operations_checklist.md:    {ops_path}"
          f" ({ops_path.stat().st_size:,} bytes)")
    print(f"  pending_manual_work.md:     {pending_path}"
          f" ({pending_path.stat().st_size:,} bytes)")
    print(f"  合并片段数:   {len(fragment_paths)}")
    # v3.0.2: 区分声明数 vs 实际写入数,缺失能立即识别
    c_ref_declared = sum(1 for _, kind, _ in merge_order if kind == 'c_reference')
    print(f"  C-reference: {c_ref_written}/{c_ref_declared}(成功写入/声明数量)")
    print(f"  待人工填充:   {len(pending_entries)}(见 pending_manual_work.md)")
    print("=" * 60)


if __name__ == '__main__':
    main()
