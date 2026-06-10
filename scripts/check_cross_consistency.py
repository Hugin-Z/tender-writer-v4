# -*- coding: utf-8 -*-
"""
check_cross_consistency.py · 跨章节一致性检查

用途:
    在 compliance_check.py 之前跑一次,捕获 AI 写作中常见的"自相矛盾"类错误。
    典型案例:团队规模 17 人 × 人天成本 → 超预算;承诺 D+65 超 60 天工期。

与 compliance_check 的关系:
    compliance_check 关注"评分项是否覆盖 / ★▲ 是否响应 / 模板残留 / 格式",
    本脚本关注"AI 输出的数字/时间/金额之间是否自相矛盾"。两者互补,都要跑。

用法:
    ./run_script.bat check_cross_consistency.py \\
        projects/{项目}/output/tender_response.docx \\
        --brief projects/{项目}/output/tender_brief.json \\
        --per-person-day-yuan 1500  # 可选,默认 1500 元/人·天(行业中位数)
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from docx import Document


# ─────────────────────────────────────────────
# 数字抽取工具
# ─────────────────────────────────────────────

_MONEY_RE = re.compile(r"(\d+(?:\.\d+)?)\s*万元|(\d{4,})\s*元")
_DAYS_RE = re.compile(r"(\d+)\s*(日历日|天|个日历日)")
_PERSON_RE = re.compile(r"(\d+)\s*(人|位)")
_D_PLUS_RE = re.compile(r"D\s*\+\s*(\d+)")

# v2.0.1:团队语境词表。只有匹配"N 人/N 位"前 20 字窗口里含以下任一词,
# 才视为团队规模数字,纳入成本估算。否则视为培训规模 / 服务对象 / 无关计数,过滤掉。
_TEAM_CONTEXT_WORDS = (
    "项目组", "团队", "项目经理", "技术骨干", "核心人员",
    "专家", "工程师", "成员", "人员配置", "项目部",
)


def extract_docx_text(docx_path: Path) -> str:
    """把 docx 的段落和表格单元格拼成纯文本,供后续正则抽取。"""
    doc = Document(str(docx_path))
    chunks = []
    for p in doc.paragraphs:
        chunks.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def parse_money_in_yuan(text: str) -> list[tuple[float, str]]:
    """返回 (金额元, 原文片段) 列表。万元 × 10000。"""
    out = []
    for m in _MONEY_RE.finditer(text):
        wan = m.group(1)
        yuan = m.group(2)
        if wan:
            try:
                out.append((float(wan) * 10000, m.group(0)))
            except ValueError:
                pass
        elif yuan:
            try:
                out.append((float(yuan), m.group(0)))
            except ValueError:
                pass
    return out


def parse_days(text: str) -> list[tuple[int, str]]:
    return [(int(m.group(1)), m.group(0)) for m in _DAYS_RE.finditer(text)]


def parse_d_plus(text: str) -> list[int]:
    return [int(m.group(1)) for m in _D_PLUS_RE.finditer(text)]


def parse_people(text: str) -> list[int]:
    """
    v2.0.1:只保留"N 人/N 位"前 20 字窗口里含团队语境词的匹配。
    过滤培训规模(如"培训不少于 100 人")/ 服务对象(如"服务 3000 万群众")/
    其他无关计数,避免团队成本估算误判。
    """
    out = []
    for m in _PERSON_RE.finditer(text):
        start = max(0, m.start() - 20)
        ctx = text[start:m.start()]
        if any(w in ctx for w in _TEAM_CONTEXT_WORDS):
            out.append(int(m.group(1)))
    return out


# ─────────────────────────────────────────────
# 检查项
# ─────────────────────────────────────────────

def check_duration_vs_dplus(budget_days: int | None,
                            d_plus_values: list[int]) -> list[str]:
    """
    检查承诺时间节点(D+N)不得超过服务期限(budget_days)。

    分级阈值:
    - D+N > duration_days × 2 → [失败](明显矛盾,如 D+130 对 60 天工期)
    - duration_days < D+N ≤ duration_days × 2 → [警告](可能是后续服务/评审配合,需人工复核)
    - D+N ≤ duration_days → [通过]
    """
    issues = []
    if budget_days is None or budget_days <= 0:
        issues.append(f"[信息] 未能从 tender_brief 获取服务期限,跳过 D+N vs 工期检查")
        return issues
    uniq = sorted(set(d_plus_values))
    hard_violations = [d for d in uniq if d > budget_days * 2]
    soft_violations = [d for d in uniq
                       if budget_days < d <= budget_days * 2]
    for d in hard_violations:
        issues.append(
            f"[失败] D+{d} 明显超出服务期限 {budget_days} 天(超 {d - budget_days} 天),"
            f"明显矛盾"
        )
    for d in soft_violations:
        issues.append(
            f"[警告] D+{d} 超出服务期限 {budget_days} 天(超 {d - budget_days} 天),"
            f"若属后续服务/评审配合可忽略,属承诺节点需修改"
        )
    if not hard_violations and not soft_violations:
        issues.append(
            f"[通过] docx 中 D+N 节点({uniq})均在服务期 {budget_days} 天内"
        )
    return issues


def check_team_cost_vs_budget(people_counts: list[int],
                              duration_days: int | None,
                              budget_yuan: float | None,
                              per_person_day_yuan: float = 1500.0,
                              tolerance: float = 1.5) -> list[str]:
    """
    检查团队规模 × 人天成本 是否在合同预算 × tolerance 内。

    计算:总人天 = 团队总人数 × 工期天数(取 docx 中最大团队规模 + 预算工期)
          估算成本 = 总人天 × per_person_day_yuan
    告警阈值:估算成本 > 合同预算 × tolerance(默认 1.5 倍)
    """
    issues = []
    if not people_counts:
        issues.append("[信息] docx 中未抽到团队人数,跳过团队成本检查")
        return issues
    if duration_days is None or duration_days <= 0:
        issues.append("[信息] 未获取服务期限,跳过团队成本检查")
        return issues
    if budget_yuan is None or budget_yuan <= 0:
        issues.append("[信息] 未获取合同预算,跳过团队成本检查")
        return issues

    max_people = max(people_counts)
    # 假设团队 50% 投入,考虑到有专家只是部分时间参与
    effective_people = max_people * 0.5
    estimated_cost = effective_people * duration_days * per_person_day_yuan
    threshold = budget_yuan * tolerance

    status = "失败" if estimated_cost > threshold else "通过"
    issues.append(
        f"[{status}] 团队成本估算: "
        f"{max_people} 人 × 50% 投入率 × {duration_days} 天 × {per_person_day_yuan} 元/人·天 "
        f"= {estimated_cost:,.0f} 元;合同预算 {budget_yuan:,.0f} 元(含 {tolerance}× 容忍度)"
    )
    if estimated_cost > threshold:
        issues.append(
            f"[建议] 团队规模/投入率/工期/预算至少一个维度需调整;"
            f"人数最大值 {max_people} 可能偏大,检查是否含专家或外协"
        )
    return issues


def check_key_numbers_consistency(text: str, budget_yuan: float | None) -> list[str]:
    """
    检查 docx 中出现的关键数字与 tender_brief 是否一致。
    若 docx 出现的金额 > 预算 10 倍,判失败(明显矛盾);1-10 倍内判警告(可能是分项/保险额等)。
    """
    issues = []
    if budget_yuan is None or budget_yuan <= 0:
        return issues
    docx_moneys = parse_money_in_yuan(text)
    hard = [m for m, _ in docx_moneys if m > budget_yuan * 10]
    soft = [m for m, _ in docx_moneys
            if budget_yuan < m <= budget_yuan * 10]
    if hard:
        uniq = sorted(set(hard))[:5]
        issues.append(
            f"[失败] docx 中金额 {uniq} 大幅超出预算 {budget_yuan:,.0f} 元(10 倍以上),"
            f"疑似编造或单位错误"
        )
    if soft:
        uniq = sorted(set(soft))[:5]
        issues.append(
            f"[警告] docx 中金额 {uniq} 超出预算 {budget_yuan:,.0f} 元(1-10 倍),"
            f"若属分项/保险额/行业参考金额可忽略"
        )
    return issues


# ─────────────────────────────────────────────
# V3-9: 项 4/5/6 新增检查
# ─────────────────────────────────────────────

# 项 5 章节识别关键词
RESUME_SECTION_KEYWORDS = [
    "简历", "师资", "团队成员", "人员配置", "项目人员", "骨干", "人员",
]
ORG_SECTION_KEYWORDS = [
    "组织架构", "组织结构", "项目组织", "人员架构", "团队架构", "团队组成",
]

# 项 5 姓名+职称正则(2-4 字中文姓名 + 必选连接符 + 0-3 字间隔 + 职称白名单)
# V3-9 commit 1 内 NAME_RE 跨人验证: "张三负责整体管理,李四担任技术负责人" 不抓张三+技术负责人
# 必选连接符: 逗号/顿号 OR 指示词(担任/为/出任/是/负责)
# 0-3 字间隔(不是 0-10 字),避免误抓非姓名 4 字串("整体管理"/"具体负责")+ 职称
NAME_RE = re.compile(
    r"([一-龥]{2,4})"
    r"(?:"
    r"[,，、]\s*(?:担任|为|出任|是|负责)?\s*"
    r"|"
    r"(?:担任|为|出任|是|负责)\s*"
    r")"
    r"(?:[一-龥]{0,3})?"
    r"(项目经理|项目负责人|技术负责人|主讲|讲师|顾问|专家|总监|副总|经理"
    r"|架构师|工程师|设计师|实施|运维)"
)

# 项 6 章节交叉引用正则(必须含引用动词 + 编号)
REF_RE = re.compile(
    r"(?:详见|见|参见|参考|参照)\s*"
    r"(?:第\s*)?"
    r"(\d+(?:\.\d+){0,2})"
    r"\s*(?:章|节|条|款)?"
)
HEADING_NUM_RE = re.compile(r"^(\d+(?:\.\d+){0,2})")


def _normalize_for_token(s: str) -> list[str]:
    """V3-9 项 4: 把字段切成 4 字滑窗 token 集合(用于 jaccard 相似度)。"""
    s = re.sub(r"\s+", "", s)
    s = re.sub(r"[，。、（）()【】《》\"\"''""]", "", s)
    if len(s) < 4:
        return [s] if s else []
    return [s[i:i + 4] for i in range(len(s) - 3)]


def _jaccard(a: list[str], b: list[str]) -> float:
    """V3-9 项 4: jaccard 相似度,空集合返回 0.0。"""
    sa, sb = set(a), set(b)
    if not sa or not sb:
        return 0.0
    return len(sa & sb) / len(sa | sb)


def check_field_consistency(
    extracted: dict,
    docx_text: str,
    docx_headings: list[str] | None = None,
) -> list[str]:
    """V3-9 项 4: 检查 brief vs docx 字段一致性,允许代称模式。

    project_name: 优先扫 docx 前 500 字符 + Heading 1 标题(项目名最常见在封面/章节标题);
                  若上述均无 jaccard ≥ 0.4 候选,降级全文"项目"附近 30 字窗口
    buyer_name / buyer_agency_name: docx 完整出现 → 通过;0 次 + 含代称 → 信息;0 次 + 无代称 → 警告
    """
    issues = []
    pn = (extracted.get("project_name") or "").strip()
    bn = (extracted.get("buyer_name") or "").strip()
    ba = (extracted.get("buyer_agency_name") or "").strip()

    if pn:
        # 第一道:substring 完整匹配。brief 字面去空格标点后在 docx 中完整出现 → 通过
        # 避免短窗口 jaccard 因分母稀释产生误报(clean fixture 中 brief = docx 但
        # 候选窗口截断到 brief 前段时 jaccard 反而 < 0.7)
        _strip_re = re.compile(r"[\s，。、（）()【】《》]")
        normalized_pn = _strip_re.sub("", pn)
        normalized_docx = _strip_re.sub("", docx_text)
        if normalized_pn and normalized_pn in normalized_docx:
            issues.append(f"[通过] 项目名一致性: brief 字面完整出现在 docx")
        else:
            # 第二道:候选窗口 + jaccard 4 字滑窗
            pn_tokens = _normalize_for_token(pn)
            candidates = [docx_text[:500]]
            if docx_headings:
                candidates.extend(docx_headings)

            best_jaccard = max(
                (_jaccard(pn_tokens, _normalize_for_token(c)) for c in candidates),
                default=0.0,
            )

            if best_jaccard < 0.4:
                fallback_cands = []
                for m in re.finditer(r"项目", docx_text):
                    start = max(0, m.start() - 30)
                    end = min(len(docx_text), m.end() + 5)
                    fallback_cands.append(docx_text[start:end])
                if fallback_cands:
                    fallback_jaccard = max(
                        _jaccard(pn_tokens, _normalize_for_token(c)) for c in fallback_cands
                    )
                    if fallback_jaccard > best_jaccard:
                        best_jaccard = fallback_jaccard
                        candidates.extend(fallback_cands)

            if best_jaccard >= 0.7:
                issues.append(f"[通过] 项目名一致性: brief vs docx jaccard={best_jaccard:.2f}")
            elif best_jaccard >= 0.4:
                best = max(candidates, key=lambda c: _jaccard(pn_tokens, _normalize_for_token(c)))
                issues.append(
                    f"[警告] 项目名 brief vs docx 部分不一致: jaccard={best_jaccard:.2f}\n"
                    f"           brief: {pn!r}\n"
                    f"           docx 最近候选: {best.strip()[:100]!r}"
                )
            else:
                issues.append(
                    f"[失败] 项目名 brief vs docx 严重不一致: jaccard={best_jaccard:.2f}\n"
                    f"           brief: {pn!r}"
                )

    for label, val in [("采购方 buyer_name", bn), ("代理 buyer_agency_name", ba)]:
        if not val:
            continue
        normalized_docx = re.sub(r"\s+", "", docx_text)
        normalized_val = re.sub(r"\s+", "", val)
        if normalized_val in normalized_docx:
            issues.append(f"[通过] {label} docx 完整出现")
        else:
            has_proxy_term = any(
                term in docx_text for term in ["采购人", "招标方", "委托方", "甲方"]
            )
            if has_proxy_term:
                issues.append(
                    f"[信息] {label}={val!r} docx 用代称(采购人/招标方/委托方/甲方),不强校验"
                )
            else:
                issues.append(f"[警告] {label}={val!r} docx 既无完整名也无代称")

    return issues


def _split_sections_by_heading(doc) -> dict[str, list[str]]:
    """V3-9 项 5: 按 Heading 1/2/3 切分章节,返回 {章节标题: [段落文本]}。"""
    sections = {}
    cur_title = None
    cur_paras = []
    for p in doc.paragraphs:
        style = getattr(p.style, "name", "")
        if "Heading" in style and p.text.strip():
            if cur_title is not None:
                sections[cur_title] = cur_paras
            cur_title = p.text.strip()
            cur_paras = []
        elif p.text.strip():
            cur_paras.append(p.text)
    if cur_title is not None:
        sections[cur_title] = cur_paras
    return sections


def check_resume_org_consistency(docx_path: Path) -> list[str]:
    """V3-9 项 5: 检查简历章节 vs 组织架构图 中的人名职称对应关系。

    简历章节: Heading 标题含简历/师资/团队成员/人员配置/项目人员/骨干/人员
    架构章节: Heading 标题含组织架构/组织结构/项目组织/人员架构/团队架构/团队组成
    """
    issues = []
    doc = Document(str(docx_path))
    sections = _split_sections_by_heading(doc)

    resume_sections = {t: paras for t, paras in sections.items()
                       if any(kw in t for kw in RESUME_SECTION_KEYWORDS)}
    org_sections = {t: paras for t, paras in sections.items()
                    if any(kw in t for kw in ORG_SECTION_KEYWORDS)}

    if not resume_sections:
        issues.append("[信息] docx 未识别到简历章节(标题含简历/师资/团队成员等),跳过项 5")
        return issues
    if not org_sections:
        issues.append("[信息] docx 未识别到组织架构章节(标题含组织架构/项目组织等),跳过项 5")
        return issues

    # 占位符 + 职称语义词(中文姓名一般不含这些字面)过滤
    NAME_SKIP_PHRASES = [
        "待填", "XX", "示例", "某某",
        "负责", "牵头", "工程", "架构", "设计", "担任", "统筹",
        "经理", "主讲", "顾问", "专家", "总监", "副总", "实施", "运维",
    ]

    def extract_pairs(paras: list[str]) -> dict[str, set[str]]:
        d = {}
        for para in paras:
            for m in NAME_RE.finditer(para):
                name, title = m.group(1), m.group(2)
                if any(skip in name for skip in NAME_SKIP_PHRASES):
                    continue
                d.setdefault(name, set()).add(title)
        return d

    resume_paras = sum(resume_sections.values(), [])
    org_paras = sum(org_sections.values(), [])
    resume_pairs = extract_pairs(resume_paras)
    org_pairs = extract_pairs(org_paras)

    only_in_resume = set(resume_pairs) - set(org_pairs)
    only_in_org = set(org_pairs) - set(resume_pairs)

    for name in sorted(only_in_resume):
        issues.append(
            f"[警告] 人员 {name!r} 在简历章节出现但组织架构图未列(职称: {sorted(resume_pairs[name])})"
        )
    for name in sorted(only_in_org):
        issues.append(
            f"[警告] 人员 {name!r} 在组织架构图出现但简历章节未列(职称: {sorted(org_pairs[name])})"
        )

    common = set(resume_pairs) & set(org_pairs)
    for name in sorted(common):
        if resume_pairs[name] & org_pairs[name]:
            issues.append(f"[通过] 人员 {name!r} 在简历和架构中职称一致或交集非空")
        else:
            issues.append(
                f"[失败] 人员 {name!r} 在简历(职称 {sorted(resume_pairs[name])}) 和架构图"
                f"(职称 {sorted(org_pairs[name])}) 中职称完全不一致"
            )

    if not resume_pairs and not org_pairs:
        issues.append("[信息] 简历/架构章节均未抽到具名人员,跳过(可能用占位符 / 描述)")

    return issues


def check_section_ref_validity(docx_path: Path) -> list[str]:
    """V3-9 项 6: 检查 docx 内章节交叉引用编号是否对应实际章节。

    引用模式: 详见 / 见 / 参见 / 参考 / 参照 + 编号(如 1.2 / 3.1.1)
    extract_docx_text 已含表格 cell 拼接,自动覆盖表格内引用。
    """
    issues = []
    doc = Document(str(docx_path))

    valid_nums = set()
    for p in doc.paragraphs:
        style = getattr(p.style, "name", "")
        if "Heading" not in style:
            continue
        title = p.text.strip()
        m = HEADING_NUM_RE.match(title)
        if m:
            valid_nums.add(m.group(1))

    if not valid_nums:
        issues.append("[信息] docx 未识别到带编号的 Heading(如 1.1 / 2.3.1),跳过项 6")
        return issues

    text = extract_docx_text(docx_path)
    bad_refs = []
    good_count = 0
    for m in REF_RE.finditer(text):
        num = m.group(1)
        end_pos = m.end()
        if end_pos < len(text) and text[end_pos] in "元万亿千百":
            continue
        if num in valid_nums:
            good_count += 1
        else:
            parent = num.rsplit(".", 1)[0] if "." in num else None
            if parent and parent in valid_nums:
                good_count += 1
            else:
                bad_refs.append((num, m.group(0)))

    if not bad_refs and good_count == 0:
        issues.append("[信息] docx 未识别到典型章节交叉引用(详见/见/参见 + 编号),跳过项 6")
        return issues

    if bad_refs:
        for num, raw in bad_refs[:10]:
            issues.append(
                f"[失败] 章节引用 {raw!r} 指向编号 {num!r},但 docx 章节体系不存在该编号"
            )
    if good_count:
        issues.append(f"[通过] {good_count} 处章节交叉引用编号有效")

    return issues


# ─────────────────────────────────────────────
# 主入口
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="跨章节一致性检查(v2)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("docx_path", help="tender_response.docx 或 final_response.docx 路径")
    parser.add_argument("--brief", required=True, help="tender_brief.json 路径")
    parser.add_argument("--per-person-day-yuan", type=float, default=1500.0,
                        help="人天成本估算基准(元/人·天,默认 1500)")
    parser.add_argument("--tolerance", type=float, default=1.5,
                        help="成本告警容忍度倍数(默认 1.5,即估算 > 预算 × 1.5 告警)")
    args = parser.parse_args()

    # V3-3: timing hook
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from _timing_hook import stage_timer
    # brief 在 project/output/tender_brief.json,推 project_dir
    _project_dir = Path(args.brief).resolve().parent.parent

    with stage_timer("check_cross_consistency", _project_dir):
        _main_body(args)


def _main_body(args):
    docx_path = Path(args.docx_path)
    brief_path = Path(args.brief)

    if not docx_path.exists():
        print(f"[错误] 找不到 docx: {docx_path}", file=sys.stderr)
        sys.exit(1)
    if not brief_path.exists():
        print(f"[错误] 找不到 tender_brief.json: {brief_path}", file=sys.stderr)
        sys.exit(1)

    # v2 单元 7 P16:用函数级闸门读 brief(而不是直接 json.loads 绕过 gate)
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from brief_schema import load_brief_guarded
    brief_data = load_brief_guarded(brief_path)
    extracted = brief_data.get("extracted", {})

    # 解析服务期限天数(从 extracted.duration 抽)
    duration_str = extracted.get("duration", "") or ""
    duration_match = _DAYS_RE.search(duration_str)
    duration_days = int(duration_match.group(1)) if duration_match else None

    # 解析合同预算金额(从 extracted.budget 抽)
    # budget_str 可能形如:"500000.00元(人民币伍拾万元整)" 或 "50万元" 或 "50 万元"
    # 规则:第一个带单位的数字决定金额;单位=元直接用,单位=万元/万 乘以 10000
    budget_str = (extracted.get("budget", "") or "").replace(",", "")
    budget_yuan = None
    m_wanyuan = re.search(r"(\d+(?:\.\d+)?)\s*万元", budget_str)
    m_yuan = re.search(r"(\d+(?:\.\d+)?)\s*元", budget_str)
    if m_wanyuan and (not m_yuan or m_wanyuan.start() < m_yuan.start()):
        budget_yuan = float(m_wanyuan.group(1)) * 10000
    elif m_yuan:
        budget_yuan = float(m_yuan.group(1))

    # 抽 docx 文本
    text = extract_docx_text(docx_path)

    d_plus_values = parse_d_plus(text)
    people_counts = parse_people(text)

    print("=" * 60)
    print("跨章节一致性检查报告")
    print("=" * 60)
    print(f"docx:  {docx_path}")
    print(f"brief: {brief_path}")
    print(f"服务期限: {duration_days} 天" if duration_days else "服务期限: 未识别")
    print(f"合同预算: {budget_yuan:,.0f} 元" if budget_yuan else "合同预算: 未识别")
    print(f"docx 抽到 D+N 节点: {sorted(set(d_plus_values))}")
    print(f"docx 抽到团队人数样本: {sorted(set(people_counts))[:10]}")
    print()

    all_issues = []

    print("[检查 1] 承诺时间节点 vs 服务期限")
    issues = check_duration_vs_dplus(duration_days, d_plus_values)
    for msg in issues:
        print(f"  {msg}")
    all_issues.extend(issues)
    print()

    print("[检查 2] 团队成本估算 vs 合同预算")
    issues = check_team_cost_vs_budget(
        people_counts, duration_days, budget_yuan,
        per_person_day_yuan=args.per_person_day_yuan,
        tolerance=args.tolerance,
    )
    for msg in issues:
        print(f"  {msg}")
    all_issues.extend(issues)
    print()

    print("[检查 3] docx 中的金额数字与预算一致性")
    issues = check_key_numbers_consistency(text, budget_yuan)
    if not issues:
        print("  [通过] 未发现明显异常金额")
    for msg in issues:
        print(f"  {msg}")
    all_issues.extend(issues)
    print()

    # V3-9: 抽 Heading 1 标题列表给项 4 做候选窗口
    docx_headings = []
    for p in Document(str(docx_path)).paragraphs:
        style = getattr(p.style, "name", "")
        if style == "Heading 1" and p.text.strip():
            docx_headings.append(p.text.strip())

    print("[检查 4] 字段一致性: 项目名 / 采购方 / 代理")
    issues = check_field_consistency(extracted, text, docx_headings=docx_headings)
    for msg in issues:
        print(f"  {msg}")
    all_issues.extend(issues)
    print()

    print("[检查 5] 简历 / 组织架构 人名一致性")
    issues = check_resume_org_consistency(docx_path)
    for msg in issues:
        print(f"  {msg}")
    all_issues.extend(issues)
    print()

    print("[检查 6] 章节交叉引用编号正确性")
    issues = check_section_ref_validity(docx_path)
    for msg in issues:
        print(f"  {msg}")
    all_issues.extend(issues)
    print()

    # 汇总
    fail_count = sum(1 for m in all_issues if m.startswith("[失败]"))
    warn_count = sum(1 for m in all_issues if m.startswith("[警告]"))
    print("=" * 60)
    print(f"跨章节一致性检查汇总: {fail_count} 失败 / {warn_count} 警告")
    print("=" * 60)

    if fail_count > 0:
        sys.exit(1)


if __name__ == "__main__":
    main()
