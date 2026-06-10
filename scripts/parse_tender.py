# -*- coding: utf-8 -*-
"""
parse_tender.py · 招标文件解析脚本(阶段 1)

用途:
    把招标文件(PDF 或 docx)解析成结构化的 tender_brief.json 和
    tender_brief.md。识别招标文件中关键章节(评分办法、资格要求、
    实质性响应、废标条款、格式要求等)并定位。

使用方法:
    通过项目根目录的 run_script.bat 调用,不要直接调用 python:
        run_script.bat parse_tender.py "<招标文件绝对路径>"
        run_script.bat parse_tender.py "<招标文件绝对路径>" --out output

参数:
    tender_path : 招标文件路径,支持 .pdf 和 .docx
    --out       : 输出目录,默认为当前工作目录下的 output/

输出:
    <out>/tender_brief.json   结构化 JSON 解析结果
    <out>/tender_brief.md     基于 templates/stage_samples/tender_brief.md 填充的 markdown
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path

# V3-6: 扫描版 PDF 检测阈值(全文平均字/页)
# 实测文字版招标 PDF 平均 466~700+ 字/页,扫描版图片 PDF 平均 0 字/页,9× 安全余量。
_SCANNED_PDF_AVG_THRESHOLD = 50
_SCANNED_PDF_TOTAL_THRESHOLD = 100  # 兜底:1-2 页极短文档边界


class ScannedPdfDetected(Exception):
    """V3-6: 扫描版 PDF 检测命中,无 --ocr-text fallback 时抛出。"""

    def __init__(self, n_pages: int, total_chars: int, avg_per_page: float, pdf_path: Path):
        self.n_pages = n_pages
        self.total_chars = total_chars
        self.avg_per_page = avg_per_page
        self.pdf_path = pdf_path
        super().__init__(
            f"PDF 疑似扫描版:{n_pages} 页 / {total_chars} 字符 / "
            f"平均 {avg_per_page:.1f} 字/页(阈值 {_SCANNED_PDF_AVG_THRESHOLD})"
        )


def read_pdf(path: Path, ocr_fallback_txt: Path | None = None) -> str:
    """读取 PDF 文件,返回纯文本(每页之间用 \\n\\n 分隔)。

    V3-6: 扫描版检测
    - 全文平均字/页 < 50 (或总字符 < 100) 视为扫描版
    - 有 ocr_fallback_txt 参数 → 读 txt 作为 raw_text 返回
    - 无参数 → 抛 ScannedPdfDetected,main 捕获后友好报错退出
    """
    try:
        import pdfplumber
    except ImportError:
        print("[错误] 缺少 pdfplumber 依赖。请先双击 install.bat 安装依赖。", file=sys.stderr)
        sys.exit(1)

    pages_text = []
    with pdfplumber.open(str(path)) as pdf:
        n_pages = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            pages_text.append(text)

    full_text = "\n\n".join(pages_text)
    total_chars = sum(len(t) for t in pages_text)
    avg_per_page = total_chars / n_pages if n_pages else 0
    is_scanned = (
        n_pages == 0
        or avg_per_page < _SCANNED_PDF_AVG_THRESHOLD
        or total_chars < _SCANNED_PDF_TOTAL_THRESHOLD
    )

    if is_scanned:
        if ocr_fallback_txt is None:
            raise ScannedPdfDetected(
                n_pages=n_pages,
                total_chars=total_chars,
                avg_per_page=avg_per_page,
                pdf_path=path,
            )
        if not ocr_fallback_txt.exists():
            print(
                f"[错误] --ocr-text 指定的文件不存在:{ocr_fallback_txt}",
                file=sys.stderr,
            )
            sys.exit(1)
        return ocr_fallback_txt.read_text(encoding="utf-8")

    return full_text


def print_scanned_pdf_help(e: ScannedPdfDetected) -> None:
    """V3-6: 扫描版 PDF 检测命中时的用户友好引导(列 3 条路径)。"""
    msg = f"""[错误] 检测到疑似扫描版 PDF:
  文件:{e.pdf_path}
  全文 {e.n_pages} 页 / {e.total_chars} 字符 / 平均 {e.avg_per_page:.1f} 字/页
  阈值:{_SCANNED_PDF_AVG_THRESHOLD} 字/页(实测文字版招标 PDF 平均 466~700+ 字/页)

扫描版 PDF 没有内嵌文字流,pdfplumber 无法提取文本。请按以下三条路径之一处理:

[路径 1] 本地 OCR(推荐技术用户)
  pip install ocrmypdf
  装 tesseract(含 chi_sim 中文包)
  ocrmypdf -l chi_sim "<原 PDF>" "<输出 PDF>"
  然后用输出 PDF 重跑 parse_tender

[路径 2] 外部 OCR / AI 多模态服务
  上传扫描版 PDF 到豆包/通义/Claude/GPT 等多模态 AI
  让其转出纯文本(注意核对 ★/▲ 条款逐字)
  存为 <pdf_name>.txt,重跑:
    parse_tender.py "<原 PDF>" --ocr-text "<pdf_name>.txt"

[路径 3] 找原版文字 PDF
  联系采购方 / 招标代理获取从 Word 直接导出的文字版 PDF
  最稳妥,不引入 OCR 错字风险

详见 docs/FAQ.md Q4。

[退出] 扫描版 PDF 处理需用户介入,parse_tender 暂停。"""
    print(msg, file=sys.stderr)


def extract_all_tables(pdf_path: Path) -> list:
    """
    V56: 用 pdfplumber 按页提取所有表格,返回 tables 列表。

    每个 table: {table_id, page_num, headers, rows, evidence}
    - table_id: 全文编号(t_001, t_002, ...)跨页连续
    - page_num: 1-indexed 页码
    - headers: 第一行作为表头,不做跨行表头合并
    - rows: 剩余行(list of list of str),None/空值规范化为 ""
    - evidence: "page_{N}_table_{M}" 位置描述

    仅支持 PDF。非 PDF 返回空列表。
    """
    if pdf_path.suffix.lower() != '.pdf':
        return []

    try:
        import pdfplumber
    except ImportError:
        print("[警告] extract_all_tables 跳过: 缺少 pdfplumber", file=sys.stderr)
        return []

    tables = []
    global_counter = 0

    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_idx, page in enumerate(pdf.pages, start=1):
            page_tables = page.extract_tables() or []
            for t_idx, raw_table in enumerate(page_tables, start=1):
                if not raw_table or len(raw_table) < 1:
                    continue

                global_counter += 1
                table_id = f't_{global_counter:03d}'

                normalized = [
                    [cell if cell is not None else '' for cell in row]
                    for row in raw_table
                ]

                headers = normalized[0] if normalized else []
                rows = normalized[1:] if len(normalized) > 1 else []

                tables.append({
                    'table_id': table_id,
                    'page_num': page_idx,
                    'headers': headers,
                    'rows': rows,
                    'evidence': f'page_{page_idx}_table_{t_idx}',
                })

    return tables


def read_docx(path: Path) -> str:
    """读取 docx 文件,返回纯文本(段落用 \n 分隔,表格用制表符分隔)"""
    try:
        from docx import Document
    except ImportError:
        print("[错误] 缺少 python-docx 依赖。请先双击 install.bat 安装依赖。", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(path))
    parts = []
    # 段落
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # 表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append("\t".join(cells))
    return "\n".join(parts)


def normalize_text(text: str) -> str:
    """文本预处理:繁简转换兜底、去除多余空白"""
    try:
        from opencc import OpenCC
        cc = OpenCC("t2s")
        text = cc.convert(text)
    except Exception:
        # 没装 opencc 或转换失败时,跳过繁简转换
        pass
    # 把多个连续空白压成单空格,但保留换行
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


# 章节编号正则：只匹配"第X章"级标题，不匹配章内子标题（一、/1.1 等）
_CHAPTER_NUM_PATTERN = re.compile(
    r"^第[一二三四五六七八九十百零\d]+章\s"
)

# 评分项标记正则：匹配"(20分)"等带括号的分值标记（全角/半角/含空格）
SCORE_MARKER = re.compile(r'[（(]\s*\d+\s*分\s*[)）]')


def extract_raw_lines_with_features(text: str) -> list[dict]:
    """把纯文本按行编号，附带特征供 AI 标注章节用。

    每行输出:
    - line_no: 行号(从 0 开始)
    - text: 行内容(strip 后)
    - length: 行长度
    - is_standalone: 是否为独立短行(长度 <= 60 且不以标点结尾)
    - has_chapter_num: 是否匹配章节编号正则
    - has_dots_leader: 是否包含目录引导符(连续 3 个以上的点号)
    """
    lines = text.split("\n")
    result = []
    for idx, raw_line in enumerate(lines):
        stripped = raw_line.strip()
        if not stripped:
            continue
        has_dots = bool(re.search(r"[.．·…]{3,}", stripped))
        has_chapter_num = bool(_CHAPTER_NUM_PATTERN.match(stripped))
        is_standalone = (len(stripped) <= 60
                         and not stripped.endswith(("，", "。", "；", "、", ",", ".", ";"))
                         and not has_dots)
        result.append({
            "line_no": idx,
            "text": stripped,
            "length": len(stripped),
            "is_standalone": is_standalone,
            "has_chapter_num": has_chapter_num,
            "has_dots_leader": has_dots,
        })
    return result


def extract_section_by_anchors(text: str, anchors: list[dict]) -> dict:
    """根据 AI 标注的 section_anchors 抽取各章节文本。

    anchors schema (由 AI 阶段 1 Step 2 填充):
        section      : str   - 章节名(SKILL.md 要求的标准名)
        start_line   : int   - 起始行号(inclusive, 0-based)
        end_line     : int   - 结束行号(exclusive, Python 切片惯例)
        confidence   : str   - "high" / "medium" / "low"
        is_embedded  : bool  - 是否嵌入到上层章节中(可选,默认 false)
        embedded_in  : str|null - 上层章节名(is_embedded=true 时填,可选)
        evidence     : str   - 判断依据

    本函数只使用 section/start_line/end_line 三个字段做文本切取,
    其余字段(含 is_embedded/embedded_in)透传保留,不消费。

    返回: {"评审办法": {"start": 669, "end": 889, "content": "..."}, ...}
    """
    lines = text.split("\n")
    sections = {}
    for anchor in anchors:
        name = anchor.get("section", "")
        start = anchor.get("start_line", 0)
        end = anchor.get("end_line", len(lines))
        content = "\n".join(lines[start:end]).strip()
        sections[name] = {"start": start, "end": end, "content": content}
    return sections


# v1 (V74+V75+V76): 删除以下函数:
#   - extract_budget / extract_cap_price / extract_duration
#   - extract_procurement_method / extract_delivery_location
#   - extract_project_name / extract_project_number
#   - extract_buyer_name / extract_buyer_agency_name
#   - extract_qualifications (+ QUAL_PATTERNS)
#   - _wrap_draft / REGEX_DRAFT_PREFIX
# 原因: 违反 skill 哲学三条红线。对非结构化招标文件做语义字段映射属主 agent 职责。
# 新流程: extracted 9 字符串字段 + qualifications 由 AI 在阶段 1 Step 2A 直接读
#         raw_lines_for_ai / raw_text 填充。详见 business_model §9 + §8.3 #N16。


def extract_substantial_marks(text: str) -> list:
    """提取所有 ★ / ▲ 标记的条款(取所在行作为上下文)"""
    marks = []
    for idx, line in enumerate(text.split("\n")):
        if "★" in line or "▲" in line:
            marks.append({"line_no": idx, "text": line.strip()})
    return marks


def extract_part_list_candidates(raw_text: str, section_anchors: list[dict]) -> list[dict]:
    """抽取 Part 清单候选段落（脚本抽候选，AI 判断）。

    按业务模型 §2.2 三个常见位置抽候选：
    1. 章节标题正则匹配"投标/响应文件(组成|格式|编写要求)"独立章节
       → 整章文本作为一条候选
    2. 投标须知前附表正则匹配"前附表" → 表格段作为一条候选
    3. 投标须知正文按关键词反推：
       - "响应文件应包括"
       - "投标文件由...组成"
       - "递交...份"
       命中段落 ±10 行作为一条候选

    返回: [{"source": "chapter"|"front_table"|"keyword_match",
             "text": "候选段落原文", "start_line": int, "end_line": int}, ...]
    """
    lines = raw_text.split("\n")
    candidates = []

    # 预扫描: 所有章节标题行的位置（用于确定章节边界）
    chapter_heading_lines = []
    for i, line in enumerate(lines):
        s = line.strip()
        if s and _CHAPTER_NUM_PATTERN.match(s) and len(s) <= 60:
            chapter_heading_lines.append(i)

    def _find_next_chapter(start: int) -> int:
        """从 start 行之后找下一个章节标题行作为边界"""
        for ch in chapter_heading_lines:
            if ch > start:
                return ch
        return len(lines)

    # --- 候选 1: 章节标题匹配 "投标/响应文件(组成|格式|编写要求)" ---
    part_list_chapter_pat = re.compile(
        r"(投标|响应|报价).{0,5}文件.{0,5}(组成|格式|编写要求|构成|编制)"
    )
    for i, line in enumerate(lines):
        s = line.strip()
        if not s or len(s) > 60:
            continue
        if part_list_chapter_pat.search(s):
            end = _find_next_chapter(i)
            candidates.append({
                "source": "chapter",
                "text": "\n".join(lines[i:end]),
                "start_line": i,
                "end_line": end,
            })

    # --- 候选 2: 前附表（标题行特征过滤，AI 判断去重）---
    # 真实标题行如"供应商须知前附表"(≤25字)或"评审办法前附表（综合评估法）"(≤25字)
    # 正文引用如"见供应商须知前附表第X条"通常更长或含"见""详见"等引用词
    front_table_title = re.compile(r"^.{0,10}前附表[\s（(]?")
    front_table_ref = re.compile(r"见|详见|按照|根据|规定")
    for i, line in enumerate(lines):
        s = line.strip()
        if not front_table_title.match(s):
            continue
        if len(s) > 25:
            continue
        if front_table_ref.search(s):
            continue
        end = _find_next_chapter(i)
        candidates.append({
            "source": "front_table",
            "text": "\n".join(lines[i:end]),
            "start_line": i,
            "end_line": end,
        })

    # --- 候选 3: 关键词匹配 ---
    keyword_patterns = [
        re.compile(r"响应文件应包括"),
        re.compile(r"投标文件由.{0,20}组成"),
        re.compile(r"递交.{0,20}份"),
    ]
    seen_lines = set()
    for i, line in enumerate(lines):
        for kp in keyword_patterns:
            if kp.search(line) and i not in seen_lines:
                seen_lines.add(i)
                start = max(0, i - 10)
                end = min(len(lines), i + 11)
                candidates.append({
                    "source": "keyword_match",
                    "text": "\n".join(lines[start:end]),
                    "start_line": start,
                    "end_line": end,
                })

    return candidates


def extract_score_items_raw_positions(raw_text: str, section_anchors: list[dict]) -> list[dict]:
    """评分项粗位置识别（只切位置，不解析分值，不判断真伪）。

    在 section_anchors 中 lookup "评审办法" 章节范围，
    按 SCORE_MARKER 正则找到所有匹配行，从一个匹配点所在行
    切到下一个匹配点所在行的前一行。最后一片延伸到章节末尾。

    严禁做的事：不解析分值数字、不判断是真评分项还是档次描述行、
    不加长度阈值过滤、不和 build_scoring_matrix.py 共用任何逻辑。

    返回: [{"raw_text": str, "start_line": int, "end_line": int}, ...]
    """
    lines = raw_text.split("\n")

    # lookup "评审办法"，不做别名映射
    scoring_anchor = None
    for anchor in section_anchors:
        if anchor.get("section") == "评审办法":
            scoring_anchor = anchor
            break

    if scoring_anchor is None:
        print("[警告] 未在 section_anchors 中找到'评审办法'章节，"
              "score_items_raw_positions 置空。"
              "AI 后续需自行定位评分项位置。", file=sys.stderr)
        return []

    start_line = scoring_anchor["start_line"]
    end_line = scoring_anchor["end_line"]

    # 在评审办法范围内找所有 SCORE_MARKER 匹配行（同一行只算一个起点）
    match_lines = []
    for i in range(start_line, min(end_line, len(lines))):
        if SCORE_MARKER.search(lines[i]):
            match_lines.append(i)

    if not match_lines:
        print(f"[警告] 评审办法章节范围 L{start_line}-L{end_line} 内"
              f"未匹配到 (\\d+分) 模式，score_items_raw_positions 置空。"
              "可能原因：本项目用裸数字表示分值。"
              "AI 后续需自行从评审办法原文识别评分项。", file=sys.stderr)
        return []

    # 切片：匹配点 N → 匹配点 N+1 前一行；最后一片 → 章节末尾
    positions = []
    for idx, line_no in enumerate(match_lines):
        if idx + 1 < len(match_lines):
            slice_end = match_lines[idx + 1]
        else:
            slice_end = end_line
        text_block = "\n".join(lines[line_no:slice_end])
        positions.append({
            "raw_text": text_block,
            "start_line": line_no,
            "end_line": slice_end,
        })

    return positions


def parse_tender(tender_path: Path, ocr_fallback_txt: Path | None = None) -> dict:
    """主解析函数,返回结构化的 dict。

    章节定位不再由脚本完成——脚本只产出 raw_lines_for_ai 供 AI 标注,
    AI 标注 section_anchors 后由 extract_section_by_anchors() 抽取内容。

    V3-6: ocr_fallback_txt 参数仅对 PDF 路径生效,docx 路径不受影响。
    """
    suffix = tender_path.suffix.lower()
    if suffix == ".pdf":
        raw = read_pdf(tender_path, ocr_fallback_txt=ocr_fallback_txt)
    elif suffix == ".docx":
        raw = read_docx(tender_path)
    else:
        print(f"[错误] 不支持的文件格式:{suffix}。仅支持 .pdf 和 .docx", file=sys.stderr)
        sys.exit(1)

    text = normalize_text(raw)

    # 提取带特征的行列表,供 AI 标注章节
    raw_lines = extract_raw_lines_with_features(text)

    # 全文级正则提取(不依赖章节定位)
    result = {
        "source_file": str(tender_path),
        "char_count": len(text),
        "raw_lines_for_ai": raw_lines,
        # V56: PDF 表格二维结构(跨页连续编号 t_001, t_002, ...)
        "tables": extract_all_tables(tender_path),
        "section_anchors": [],  # 由 AI 阶段 1 Step 2 填充,schema 含 is_embedded/embedded_in
        "sections": {},         # 由 AI 填充 section_anchors 后,脚本可按需抽取
        "extracted": {
            # v1.2+ v1: 由 AI 读 raw_lines_for_ai / raw_text 直接填
            # 不再用 regex 初提,消除 __REGEX_DRAFT__ 机制(见 business_model §9 + §8.3 #N16)
            "procurement_method": "",
            "budget": "",
            "cap_price": "",
            "duration": "",
            "delivery_location": "",
            "project_name": "",
            "project_number": "",
            "buyer_name": "",
            "buyer_agency_name": "",
            "qualifications": [],
            "substantial_response_marks": extract_substantial_marks(text)[:50],
            # v2:项目类型,由 AI 阶段 1 Step 2A 判定填入,generate_outline 按此选模板
            # 合法值:engineering(工程) / platform(平台) / research(课题研究) / planning(规划编制) / other
            # 判据:采购需求关键词 + 最终交付物类型 + 服务期限长短(详见 SKILL.md 阶段 1 Step 2A)
            "project_type": "",
        },
        # v2 单元 5:记录上游文件类型(pdf/docx),影响 C 模式是否走 Word 原样切用分支
        "source_meta": {
            "source_format": tender_path.suffix.lstrip(".").lower(),
        },
        "part_list_candidates": [],           # 下方填充
        "score_items_raw_positions": [],       # 下方填充
        "raw_text": text,
    }

    # Part 清单候选段落抽取（脚本抽候选，AI 判断）
    result["part_list_candidates"] = extract_part_list_candidates(
        text, result["section_anchors"]
    )

    # 评分项粗位置切片（依赖 section_anchors 中的"评审办法"条目）
    # 首次运行时 section_anchors 为空，输出 []；AI 填充 section_anchors 后可重新调用
    result["score_items_raw_positions"] = extract_score_items_raw_positions(
        text, result["section_anchors"]
    )

    return result


def render_brief_md(result: dict) -> str:
    """
    生成阶段 1 的 tender_brief.md 初稿。
    自动能提取的字段先落入简报，其余字段保留 "【待补充】" 占位，
    供 AI 和用户在后续 review 中继续补全。
    """
    sections = result.get("sections", {})
    extracted = result.get("extracted", {})

    def get_section_preview(key):
        alias_map = {
            "qualification": ["qualification", "投标人资格要求", "供应商资格要求"],
            "scoring": ["scoring", "评审办法", "评分办法"],
            "substantial_response": ["substantial_response", "实质性响应条款"],
            "disqualification": ["disqualification", "废标条款"],
            "format_requirement": ["format_requirement", "响应文件格式", "格式要求"],
            "submission": ["submission", "响应文件格式", "投标文件构成与装订"],
            "open_bid": ["open_bid", "响应文件格式", "关键时间节点"],
        }
        sec = None
        for alias in alias_map.get(key, [key]):
            sec = sections.get(alias)
            if sec:
                break
        if not sec:
            return "【待 AI 标注 section_anchors 后自动填充,或人工补充】"
        full_text = sec.get("content", "").strip()
        if full_text:
            return full_text[:500]
        return sec.get("content_preview", "").strip() or "【章节内容为空,请人工核查】"

    qual_lines = []
    quals = extracted.get("qualifications", [])
    if quals:
        for q in quals:
            qual_lines.append(f"- **{q['name']}**:{q['snippet']}")
    else:
        qual_lines.append("- 【未在资格要求章节中识别到常见资质项,请人工补充】")

    sub_lines = []
    marks = extracted.get("substantial_response_marks", [])
    if marks:
        for m in marks[:30]:
            sub_lines.append(f"- (行 {m['line_no']}) {m['text']}")
        if len(marks) > 30:
            sub_lines.append(f"- ……(共 {len(marks)} 条 ★/▲ 标记,此处仅显示前 30 条)")
    else:
        sub_lines.append("- 【未检测到 ★/▲ 标记的实质性响应条款,请人工核查招标文件原文】")

    md = f"""# 招标文件解读简报(tender_brief)

> **来源文件**:{result.get('source_file', '')}
> **字符总数**:{result.get('char_count', 0)}
> **生成时间**:由 parse_tender.py 自动生成

> ⚠️ 本文档是后续所有阶段的**唯一事实来源(single source of truth)**。
> 凡是标注为"【待补充】"或"【未在招标文件中定位到】"的字段,
> **必须由用户人工核查并补充**,严禁让模型凭记忆脑补。

---

## 提醒维护者补清单

（阶段 1 Step 3 填充）

---

## 招标文件矛盾/不一致记录

（阶段 1 Step 3 填充）

---

## 一、项目基本信息

- **项目名称**:【待补充,请从招标文件首页/封面提取】
- **采购人**:【待补充】
- **采购代理机构**:【待补充】
- **项目编号**:【待补充】
- **采购方式**:{extracted.get('procurement_method', '') or '【未自动识别,请人工补充】'}
- **预算金额**:{extracted.get('budget', '') or '【未自动识别,请人工补充】'}
- **最高限价**:{extracted.get('cap_price', '') or '【未自动识别,请人工补充】'}
- **总工期**:{extracted.get('duration', '') or '【未自动识别,请人工补充】'}
- **关键里程碑**:【待补充】
- **交付地点**:{extracted.get('delivery_location', '') or '【未自动识别,请人工补充】'}

---

## 二、投标人资格要求

{chr(10).join(qual_lines)}

> 章节原文摘要(前 500 字符):
>
> {get_section_preview('qualification')}

---

## 三、评分办法

> ⚠️ 这是技术标编制最关键的章节,请人工通读章节原文并填写下表。

### 评分维度权重

- 技术分:【__】 分
- 商务分:【__】 分
- 价格分:【__】 分
- 合计:100 分

### 章节原文摘要(前 500 字符)

{get_section_preview('scoring')}

---

## 四、实质性响应条款(★/▲)

{chr(10).join(sub_lines)}

> 章节原文摘要(前 500 字符):
>
> {get_section_preview('substantial_response')}

---

## 五、废标条款

{get_section_preview('disqualification')}

---

## 六、格式要求

{get_section_preview('format_requirement')}

> 关键格式参数(请人工核查并补充):
> - 字体字号:【待补充】
> - 行距:【待补充】
> - 页边距:【待补充】
> - 封面样式:【待补充】
> - 目录要求:【待补充】

---

## 七、投标文件构成与装订

{get_section_preview('submission')}

> - 正本份数:【待补充】
> - 副本份数:【待补充】
> - 电子版份数:【待补充】
> - 装订方式:【待补充】

---

## 八、关键时间节点

{get_section_preview('open_bid')}

> - 投标截止时间:【待补充】
> - 开标时间:【待补充】
> - 开标地点:【待补充】

---

## 九、项目所在地本地化信息(撰写时必须基于真实数据)

> ⚠️ 严禁使用“该地区资源丰富、产业兴旺”等通用套话。
> 不知道的字段宁可留空让用户补,也不要编造。

- **项目所在地**:【待补充,精确到县/区/乡镇】
- **行政区划**:【待补充】
- **常住人口**:【待补充】
- **GDP / 财政收入**:【待补充】
- **主导产业**:【待补充】
- **信息化基础**:【待补充】
- **特殊地理 / 气候因素**:【待补充】

---

## 十、人工核查 checklist(请用户逐项打钩)

- [ ] 项目基本信息已补全(名称、预算、工期、里程碑)
- [ ] 资格要求已逐项核对,我方满足
- [ ] 评分办法各维度权重已填写
- [ ] 所有 ★/▲ 条款已列出并标注响应方式
- [ ] 废标条款已通读,关键风险点已标注
- [ ] 格式要求已记录(字号、行距、页边距)
- [ ] 投标文件构成与份数已确认
- [ ] 投标截止时间已记录,留出充足缓冲
- [ ] 项目所在地的本地化信息已补充或明确标注"待我后续补充"

---

> ✅ 本简报经用户确认无误后,才能进入阶段 2(评分矩阵构建)。
"""
    return md


def main():
    parser = argparse.ArgumentParser(
        description="招标文件解析(阶段 1)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("tender_path", help="招标文件路径(.pdf 或 .docx)")
    parser.add_argument("--out", default="output", help="输出目录,默认 output/")
    parser.add_argument(
        "--force",
        action="store_true",
        help="目标 output/tender_brief.json 已存在时强制覆盖(默认拒绝,防止误覆盖已填字段)",
    )
    parser.add_argument(
        "--ocr-text",
        default=None,
        help=(
            "V3-6: 扫描版 PDF 的 OCR 结果 txt 路径。命中扫描版检测时,"
            "作为 raw_text 替代,跳过 PDF 文字提取。详见 docs/FAQ.md Q4。"
        ),
    )
    args = parser.parse_args()

    # V3-3: timing hook
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from _timing_hook import stage_timer
    _out_path = Path(args.out)
    _project_dir = _out_path.parent if _out_path.name == "output" else _out_path

    with stage_timer("parse_tender", _project_dir):
        _main_body(args)


def _main_body(args):
    tender_path = Path(args.tender_path)
    out_dir = Path(args.out)
    json_path = out_dir / "tender_brief.json"
    raw_path = out_dir / "tender_raw.txt"

    if not tender_path.exists():
        print(f"[错误] 找不到招标文件:{tender_path}", file=sys.stderr)
        sys.exit(1)

    out_dir.mkdir(parents=True, exist_ok=True)

    # V80: 重跑保护 — 默认拒绝覆盖已存在的 tender_brief.json
    if json_path.exists() and not args.force:
        print(
            f"[错误] 目标 tender_brief.json 已存在: {json_path}\n"
            f"\n"
            f"  parse_tender 是初次生成工具,不是增量更新工具。\n"
            f"  重跑会将 extracted 字段重置为空骨架,覆盖由 AI/用户填充的已有值。\n"
            f"\n"
            f"  如果这是有意的覆盖(如招标文件更新、parse_tender 逻辑调整等),\n"
            f"  加 --force 参数显式允许:\n"
            f"    ./run_script.bat parse_tender.py \"<招标文件路径>\" --force\n"
            f"\n"
            f"  如果不是有意覆盖,不要运行 parse_tender。\n"
            f"  修改 tender_brief.json 字段直接编辑 json 文件后手工重建 .reviewed 标记。",
            file=sys.stderr,
        )
        sys.exit(1)

    if json_path.exists() and args.force:
        print(
            f"[警告] --force 模式: 将覆盖已存在的 {json_path}\n"
            f"  extracted 字段会重置为空骨架,需要 AI 重新填充并由用户重新 review。\n"
            f"  ⚠️ 当前 ensure_reviewed 闸门只检查 .reviewed 文件存在性,不校验 hash,\n"
            f"     下游脚本仍会放行使用空骨架数据。\n"
            f"  重跑前请手动删除 output/tender_brief.reviewed 标记,完成新一轮 review 后重建。",
            file=sys.stderr,
        )

    print(f"[信息] 正在解析:{tender_path}")
    ocr_fallback = Path(args.ocr_text) if args.ocr_text else None
    try:
        result = parse_tender(tender_path, ocr_fallback_txt=ocr_fallback)
    except ScannedPdfDetected as e:
        print_scanned_pdf_help(e)
        sys.exit(2)

    # 写 JSON(去掉 raw_text 以减小体积,raw_text 单独写一个文件供调试)
    json_result = {k: v for k, v in result.items() if k != "raw_text"}
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(json_result, f, ensure_ascii=False, indent=2)
    print(f"[完成] JSON 已写入:{json_path}")

    with open(raw_path, "w", encoding="utf-8") as f:
        f.write(result["raw_text"])
    print(f"[完成] 原文文本已写入:{raw_path}")

    # 写 markdown
    md_text = render_brief_md(result)
    md_path = out_dir / "tender_brief.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_text)
    print(f"[完成] markdown 已写入:{md_path}")

    print()
    print("=" * 60)
    print("阶段 1 脚本部分完成。下一步:")
    print(f"  1. AI 读取 {json_path} 中的 raw_lines_for_ai")
    print("  2. AI 标注 section_anchors(含 is_embedded/embedded_in),写回 JSON")
    print(f"  3. 跑 update_score_positions.py {json_path} 生成 score_items_raw_positions")
    print(f"  4. 用户 review {md_path},补全【待补充】字段")
    print("  5. AI 读 raw_lines_for_ai 直接填 extracted 字段的 9 个字符串字段")
    print("     (procurement_method / budget / cap_price / duration / delivery_location /")
    print("      project_name / project_number / buyer_name / buyer_agency_name)")
    print("     + qualifications 列表。见 SKILL.md 阶段 1 Step 2A / business_model §9")
    print("  6. AI 生成 response_file_parts;对 production_mode='C' 的 Part,")
    print("     按 business_model §8 #N20 判据判定 sub_mode 并写入")
    print("  7. 用户确认后,进入阶段 2:")
    print("     .\\run_script.bat build_scoring_matrix.py output/tender_brief.json")
    print("=" * 60)


if __name__ == "__main__":
    main()
