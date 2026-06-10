# -*- coding: utf-8 -*-
"""
docx_builder.py · 标书 docx 构建工具模块

用途:
    本模块提供构建中文政府类标书 docx 所需的所有底层工具函数,
    包括字体字号设置、页边距设置、自动目录、SEQ 域图表编号、繁简
    转换兜底等。被其他脚本(parse_tender / build_scoring_matrix /
    阶段 4 的章节追加)以 import 方式复用。

使用方法:
    本模块既可作为模块被 import,也可直接命令行运行以创建一份
    空的标书骨架(封面 + 目录 + 页眉页脚 + 默认样式):
        run_script.bat docx_builder.py --out output/tender_response.docx \\
            --project "智慧城市综合管理平台" --bidder "示例科技有限公司"

参数(命令行模式):
    --out      : 输出 docx 路径,默认 output/tender_response.docx
    --project  : 项目名称(写到封面)
    --bidder   : 投标人名称(写到封面)
    --date     : 投标日期(默认今天)

设计要点:
    1. 中文字体必须用 XML 设置 w:eastAsia,只设 font.name 不生效
    2. 默认页边距:上 2.54 / 下 2.54 / 左 3.17 / 右 3.17 cm
    3. 默认正文:宋体小四 14pt(对应需求规格说明书 GB 标准)
       默认标题:黑体三号(H1)/ 小三(H2)/ 四号(H3)/ 小四(H4)
    4. 目录使用 Word TOC 域,Word 打开后按 F9 更新
    5. 图表编号使用 Word SEQ 域,自动连续
    6. opencc 兜底繁简转换:写入文档前对所有中文做 t2s 转换
"""

import argparse
import re
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[错误] 缺少 python-docx 依赖。请先双击 install.bat 安装依赖。", file=sys.stderr)
    sys.exit(1)


# ============================================================
# 一、繁简转换兜底
# ============================================================

_OPENCC_CONVERTER = None


def _get_opencc():
    """懒加载 opencc 转换器,避免每次调用都初始化"""
    global _OPENCC_CONVERTER
    if _OPENCC_CONVERTER is None:
        try:
            from opencc import OpenCC
            _OPENCC_CONVERTER = OpenCC("t2s")
        except Exception:
            _OPENCC_CONVERTER = False  # 标记为不可用
    return _OPENCC_CONVERTER if _OPENCC_CONVERTER else None


def to_simplified(text: str) -> str:
    """繁体转简体兜底。opencc 不可用时返回原文。"""
    cc = _get_opencc()
    if cc:
        try:
            return cc.convert(text)
        except Exception:
            return text
    return text


# ============================================================
# 二、字体字号设置(核心)
# ============================================================

# v2:全文默认黑色(避免 Word 主题色把标题/正文渲染成灰/蓝)
DEFAULT_TEXT_COLOR = RGBColor(0, 0, 0)


def set_run_font(run, ascii_font="Times New Roman", east_asia="宋体",
                 size_pt=14, bold=False, color=None, italic=False):
    """
    设置某个 run 的中英文字体字号。
    中文字体必须通过 XML 设置 w:eastAsia,只设 font.name 不生效。

    v2:color 默认强制为 RGB(0,0,0);italic 新增支持内联解析。
    """
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    # v2:默认黑色,避免 Word 主题渲染成灰/蓝
    run.font.color.rgb = color if color is not None else DEFAULT_TEXT_COLOR

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:cs"), ascii_font)


# ============================================================
# 三、默认样式应用
# ============================================================

# 默认样式表(对应 references/doc_format_spec.md)
DEFAULT_STYLES = {
    "Normal":     {"east_asia": "宋体",   "size": 14,   "bold": False},  # 正文 四号 14pt
    "Heading 1":  {"east_asia": "黑体",   "size": 16,   "bold": True},   # H1 三号 16pt
    "Heading 2":  {"east_asia": "黑体",   "size": 15,   "bold": True},   # H2 小三 15pt
    "Heading 3":  {"east_asia": "黑体",   "size": 14,   "bold": True},   # H3 四号 14pt
    "Heading 4":  {"east_asia": "黑体",   "size": 12,   "bold": True},   # H4 小四 12pt
}


def apply_default_styles(doc: Document):
    """
    把默认样式表应用到文档的 Normal / Heading 1~4。
    v2:强制所有样式颜色为黑色 RGB(0,0,0),覆盖 Word 主题色。
    """
    for style_name, conf in DEFAULT_STYLES.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(conf["size"])
        font.bold = conf["bold"]
        # v2:强制黑色
        font.color.rgb = DEFAULT_TEXT_COLOR
        # 设置 east_asia
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:eastAsia"), conf["east_asia"])
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")


# ============================================================
# 四、页面设置
# ============================================================

def set_page_margins(doc: Document,
                     top_cm=2.54, bottom_cm=2.54,
                     left_cm=3.17, right_cm=3.17):
    """设置所有 section 的页边距"""
    for section in doc.sections:
        section.top_margin = Cm(top_cm)
        section.bottom_margin = Cm(bottom_cm)
        section.left_margin = Cm(left_cm)
        section.right_margin = Cm(right_cm)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)


# ============================================================
# 五、段落辅助
# ============================================================

def add_paragraph(doc: Document, text: str,
                  style="Normal", first_line_indent_chars=2,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  line_spacing=1.5):
    """
    添加一段正文,自动套用首行缩进 2 字符 + 1.5 倍行距。
    text 会先经过繁简转换。
    """
    text = to_simplified(text)
    p = doc.add_paragraph(style=style)
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if first_line_indent_chars > 0 and style == "Normal":
        # 正文 14pt × 2 字符 ≈ 28pt 缩进(中文每字字宽 ≈ 字号)
        pf.first_line_indent = Pt(14 * first_line_indent_chars)

    run = p.add_run(text)
    if style == "Normal":
        set_run_font(run, east_asia="宋体", size_pt=14, bold=False)
    return p


def add_chapter(doc: Document, title: str, level: int = 1):
    """添加一章/一节标题(level=1~4)"""
    title = to_simplified(title)
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.first_line_indent = Pt(0)

    conf = DEFAULT_STYLES.get(style, DEFAULT_STYLES["Heading 1"])
    run = p.add_run(title)
    set_run_font(run, east_asia=conf["east_asia"],
                 size_pt=conf["size"], bold=conf["bold"])
    return p


def _body_size_to_table_size(body_size_pt: float) -> float:
    """v2:表格字号相对规则。
    正文四号(14)→ 表格小四(12);正文小四(12)→ 表格五号(10.5);正文小五(9)→ 表格小五(9)。
    """
    if body_size_pt >= 14:
        return 12.0
    if body_size_pt >= 12:
        return 10.5
    return max(body_size_pt, 9.0)


def add_table(doc: Document, headers: list, rows: list,
              header_east_asia="宋体",
              body_east_asia="宋体",
              size_pt=None):
    """
    添加一张表格。表头加粗居中,表体内容根据长度自动选择居中或左对齐。

    v2:
    - 默认中文字体从 "仿宋_GB2312" 改为 "宋体"(与正文统一)
    - size_pt=None 时按正文字号自动下调(正文 14 → 表格 12;正文 12 → 表格 10.5)
    """
    if size_pt is None:
        body_size = DEFAULT_STYLES.get("Normal", {}).get("size", 14)
        size_pt = _body_size_to_table_size(body_size)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(to_simplified(str(h)))
        set_run_font(run, east_asia=header_east_asia, size_pt=size_pt, bold=True)

    # 表体
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            text = to_simplified(str(val))
            # 短文本居中,长文本左对齐
            if len(text) <= 12:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            set_run_font(run, east_asia=body_east_asia, size_pt=size_pt, bold=False)
    return table


# ============================================================
# 六、Word 域代码:目录(TOC)和图表编号(SEQ)
# ============================================================

def insert_toc(paragraph):
    """
    在指定段落处插入 Word 自动目录(TOC 域)。
    Word 打开后,选中目录区域按 F9 即可更新。
    """
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在 Word 中按 F9 更新"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._element.append(fld_begin)
    run._element.append(instr_text)
    run._element.append(fld_separate)
    run._element.append(placeholder)
    run._element.append(fld_end)


def insert_seq_field(paragraph, seq_name="图"):
    """
    在指定段落处插入 SEQ 域,实现自动编号。
    例如:在图注段落用 add_run("图 ") + insert_seq_field(p, "图")
    可得到 "图 1" "图 2" 这样的自动编号。
    """
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" SEQ {seq_name} \\* ARABIC "

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._element.append(fld_begin)
    run._element.append(instr_text)
    run._element.append(fld_end)


def _add_placeholder_box(doc: Document, text: str = "【图片占位】"):
    """
    v2:添加一个"图片占位区块"——带四边框的居中空白段落,
    提示用户在此处插入真正的图片。
    实现:用段落属性 <w:pBdr> 给段落加四边框;段落内容是"【图片占位】"字样。
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # 加四边框
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "808080")
        pbdr.append(b)
    p_pr.append(pbdr)

    run = p.add_run(text)
    set_run_font(run, east_asia="宋体", size_pt=10.5, bold=False,
                 color=RGBColor(0x80, 0x80, 0x80))
    return p


def add_figure_caption(doc: Document, caption_text: str,
                       with_placeholder: bool = True):
    """
    添加图注:图 X xxx(SEQ 自动编号)

    v2:with_placeholder=True 时在图注前加一个"图片占位区块"
    (带边框的居中空白段落),使图占位在 docx 中有真正的视觉区块,而不仅仅是字面字符。
    """
    if with_placeholder:
        _add_placeholder_box(doc)

    caption_text = to_simplified(caption_text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)

    run1 = p.add_run("图 ")
    set_run_font(run1, east_asia="宋体", size_pt=10.5, bold=False)
    insert_seq_field(p, "图")
    run2 = p.add_run(f" {caption_text}")
    set_run_font(run2, east_asia="宋体", size_pt=10.5, bold=False)
    return p


# v2:docx 生成后的空格清理后处理
# 作为 AI prompt 约束(R3)的兜底:即使 AI 输出了带空格的 markdown,
# docx 渲染后仍然把中文-数字/中文-英文等之间的空格清除。
_CLEAN_PATTERNS = [
    (re.compile(r"([一-鿿])[ \t]+(\d)"), r"\1\2"),
    (re.compile(r"(\d)[ \t]+([一-鿿])"), r"\1\2"),
    (re.compile(r"([一-鿿])[ \t]+([A-Za-z])"), r"\1\2"),
    (re.compile(r"([A-Za-z])[ \t]+([一-鿿])"), r"\1\2"),
]


def clean_text_whitespace(text: str) -> str:
    """按 R3 规则清理中英/中数之间的多余空格。多次应用直到稳定。"""
    prev = None
    out = text
    # 重复应用直到不再变化(最多 5 次,实际 1-2 次足够)
    for _ in range(5):
        if out == prev:
            break
        prev = out
        for pat, repl in _CLEAN_PATTERNS:
            out = pat.sub(repl, out)
    return out


def clean_docx_whitespace(doc: Document) -> int:
    """
    扫整个 docx 的所有段落/表格单元格 run,应用 clean_text_whitespace。
    返回清理的 run 数。

    注意:只改 run.text,不动字体/格式;跳过 run 中不含中文的情况(纯英文段落、纯代码)。
    """
    changed = 0

    def _process_paragraphs(paragraphs):
        nonlocal changed
        for para in paragraphs:
            for run in para.runs:
                t = run.text
                if not t:
                    continue
                # 只处理含中文的 run
                if not re.search(r"[一-鿿]", t):
                    continue
                new_t = clean_text_whitespace(t)
                if new_t != t:
                    run.text = new_t
                    changed += 1

    _process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _process_paragraphs(cell.paragraphs)
    return changed


def add_table_caption(doc: Document, caption_text: str):
    """添加表注:表 X xxx(SEQ 自动编号),通常放在表格上方"""
    caption_text = to_simplified(caption_text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    run1 = p.add_run("表 ")
    set_run_font(run1, east_asia="宋体", size_pt=10.5, bold=False)
    insert_seq_field(p, "表")
    run2 = p.add_run(f" {caption_text}")
    set_run_font(run2, east_asia="宋体", size_pt=10.5, bold=False)
    return p


# ============================================================
# 七、封面
# ============================================================

def add_cover_page(doc: Document, project_name: str, bidder_name: str,
                   bid_date: str = ""):
    """
    添加封面页:项目名(主标题) + 投标人 + 日期。
    封面与正文之间插入分页符。
    """
    project_name = to_simplified(project_name)
    bidder_name = to_simplified(bidder_name)

    # 顶部留空 6 行
    for _ in range(6):
        empty = doc.add_paragraph()
        empty.paragraph_format.space_before = Pt(0)
        empty.paragraph_format.space_after = Pt(0)

    # 主标题
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run("投 标 文 件")
    set_run_font(r1, east_asia="黑体", size_pt=22, bold=True)

    # 空两行
    for _ in range(2):
        doc.add_paragraph()

    # 副标题:项目名
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(project_name)
    set_run_font(r2, east_asia="宋体", size_pt=18, bold=True)

    # 空一行
    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(0)
    r3 = p3.add_run("(技术标)")
    set_run_font(r3, east_asia="宋体", size_pt=18, bold=True)

    # 底部投标人 + 日期(留出空白)
    for _ in range(10):
        doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(0)
    p4.paragraph_format.space_after = Pt(0)
    r4 = p4.add_run(f"投标人:{bidder_name}")
    set_run_font(r4, east_asia="宋体", size_pt=12, bold=False)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.space_before = Pt(0)
    p5.paragraph_format.space_after = Pt(0)
    r5 = p5.add_run(f"日  期:{bid_date or date.today().strftime('%Y 年 %m 月 %d 日')}")
    set_run_font(r5, east_asia="宋体", size_pt=12, bold=False)

    # 分页
    doc.add_page_break()


# ============================================================
# 八、目录页
# ============================================================

def add_toc_page(doc: Document):
    """添加目录页:'目  录' 标题 + TOC 域 + 分页"""
    # "目  录" 二字
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)
    r = p_title.add_run("目  录")
    set_run_font(r, east_asia="黑体", size_pt=16, bold=True)

    # TOC 域
    p_toc = doc.add_paragraph()
    insert_toc(p_toc)

    # 分页
    doc.add_page_break()


# ============================================================
# 九、创建空骨架
# ============================================================

def create_tender_doc(out_path: Path,
                      project_name: str,
                      bidder_name: str,
                      bid_date: str = ""):
    """
    创建一份空的标书 docx 骨架,包含:
    - 默认样式(字体字号、页边距)
    - 封面
    - 目录(TOC 域,Word 中按 F9 更新)
    - 等待追加正文章节
    """
    doc = Document()
    apply_default_styles(doc)
    set_page_margins(doc)
    add_cover_page(doc, project_name, bidder_name, bid_date)
    add_toc_page(doc)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# ============================================================
# 十、正文容器(section-only 模式)
# ============================================================

def _activate_list_styles(doc: Document):
    """
    预激活 Word 内置列表样式。python-docx 创建的空白文档不一定包含
    所有 List 样式,通过"插入占位段落 + 立即删除"的方式将样式注入
    styles.xml。
    """
    styles_to_activate = [
        'List Bullet', 'List Bullet 2',
        'List Number', 'List Number 2',
    ]
    placeholders = []
    for style_name in styles_to_activate:
        try:
            p = doc.add_paragraph(style=style_name)
            placeholders.append(p)
        except KeyError:
            raise RuntimeError(
                f"Word 内置样式 '{style_name}' 在 python-docx 默认 docx 中不可用,"
                f"create_section_doc 无法预激活。请检查 python-docx 版本或手动注入样式。"
            )
    for p in placeholders:
        p._element.getparent().remove(p._element)


def create_section_doc(out_path: Path) -> Path:
    """
    创建一份只包含样式和页边距的空白 docx 容器,用于阶段 4 追加
    正文章节。不包含封面、目录、页眉、页脚。

    与 create_tender_doc 的区别:
    - create_tender_doc: 完整单文档骨架(封面+目录+正文容器)
    - create_section_doc: 仅正文容器,用于作为整标的一个 Part
    """
    doc = Document()
    apply_default_styles(doc)
    set_page_margins(doc)
    _activate_list_styles(doc)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# ============================================================
# 十一、命令行入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="标书 docx 构建工具(可命令行创建空骨架)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("--out", default="output/tender_response.docx",
                        help="输出 docx 路径")
    parser.add_argument("--project", default="【项目名称】",
                        help="项目名称(写到封面)")
    parser.add_argument("--bidder", default="【投标人名称】",
                        help="投标人名称(写到封面)")
    parser.add_argument("--date", default="",
                        help="投标日期,默认今天")
    parser.add_argument("--section-only", action="store_true",
                        help="仅创建正文容器(不含封面/目录),用于整标合并")
    args = parser.parse_args()

    if args.section_only:
        out_path = create_section_doc(Path(args.out))
        print(f"[完成] 正文容器已生成:{out_path}")
        print()
        print("=" * 60)
        print("下一步:进入阶段 4(分章节撰写)")
        print("  - 每写一章,通过 append_chapter.py 追加到此容器中")
        print("  - 本容器不含封面/目录,由后续整标合并器统一处理")
        print("=" * 60)
    else:
        out_path = create_tender_doc(
            Path(args.out),
            project_name=args.project,
            bidder_name=args.bidder,
            bid_date=args.date,
        )
        print(f"[完成] 标书空骨架已生成:{out_path}")
        print()
        print("=" * 60)
        print("下一步:进入阶段 4(分章节撰写)")
        print("  - 每写一章,通过 append_chapter.py 追加到此骨架中")
        print("  - 写完后用 Word 打开,选中目录按 F9 更新")
        print("=" * 60)


if __name__ == "__main__":
    main()
