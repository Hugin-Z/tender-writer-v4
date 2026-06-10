# -*- coding: utf-8 -*-
"""
extract_text.py · 通用文本提取工具(供素材摄入与 triage 流程调用)

用途:
    输入一个 PDF 或 docx 文件路径,把纯文本输出到 stdout,并把
    文件 sha256(用于摄入去重)输出到 stderr。被以下场景调用:
      - 已知分类的素材摄入流程(章节九)
      - 未知分类的 triage 流程(章节十)

使用方法:
    通过项目根目录的 run_script.bat 调用,不要直接调用 python:
        run_script.bat extract_text.py "<文件绝对路径>"

参数:
    file_path : 输入文件路径,支持 .pdf 和 .docx
                .doc 旧格式不支持,会提示用户先转换为 docx

输出:
    stdout : 提取的纯文本(UTF-8)
    stderr : 一行 "SHA256: <64 位十六进制>"
    exit 0 : 成功
    exit 1 : 文件不存在 / 格式不支持 / 解析失败
"""

import argparse
import hashlib
import sys
from pathlib import Path


def compute_sha256(path: Path) -> str:
    """计算文件的 sha256 摘要"""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def extract_pdf(path: Path) -> str:
    """通过 pdfplumber 提取 PDF 全部文字"""
    try:
        import pdfplumber
    except ImportError:
        print("[错误] 缺少 pdfplumber 依赖。请先双击 install.bat 安装依赖。", file=sys.stderr)
        sys.exit(1)

    parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            parts.append(text)
    return "\n\n".join(parts)


def extract_docx(path: Path) -> str:
    """通过 python-docx 提取段落与表格"""
    try:
        from docx import Document
    except ImportError:
        print("[错误] 缺少 python-docx 依赖。请先双击 install.bat 安装依赖。", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(path))
    parts = []
    # 段落
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # 表格(每行用制表符分隔单元格)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append("\t".join(cells))
    return "\n".join(parts)


def main():
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    if hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")

    parser = argparse.ArgumentParser(
        description="通用文本提取工具(PDF/docx → 纯文本)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("file_path", help="输入文件路径(.pdf 或 .docx)")
    args = parser.parse_args()

    path = Path(args.file_path)
    if not path.exists():
        print(f"[错误] 找不到文件:{path}", file=sys.stderr)
        sys.exit(1)

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = extract_pdf(path)
    elif suffix == ".docx":
        text = extract_docx(path)
    elif suffix == ".doc":
        print("[错误] 旧版 .doc 格式不支持。请先用 Word 另存为 .docx 后重试。", file=sys.stderr)
        sys.exit(1)
    else:
        print(f"[错误] 不支持的文件格式:{suffix}。仅支持 .pdf 和 .docx", file=sys.stderr)
        sys.exit(1)

    # sha256 写到 stderr,避免污染 stdout 的纯文本
    sha = compute_sha256(path)
    print(f"SHA256: {sha}", file=sys.stderr)

    # 纯文本写到 stdout
    sys.stdout.write(text)
    sys.stdout.flush()


if __name__ == "__main__":
    main()
