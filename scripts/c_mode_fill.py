# -*- coding: utf-8 -*-
"""
c_mode_fill.py · C 模式模板填充脚本(阶段 4-C Step 5-6)

用法:
    python scripts/c_mode_fill.py --project <项目名> --part <索引>

输入:
    projects/<项目>/output/c_mode/<part_name>/template.docx
    projects/<项目>/output/c_mode/<part_name>/variables.yaml
    projects/<项目>/output/tender_brief.json
    companies.yaml

行为(v2 改为非交互):
    按 variables.yaml 的 source+path 取值;无法解析的变量(pending_user /
    manual)**不进交互**,直接在 filled.docx 中写入"【待填:变量描述】"
    显式占位,让用户在 Word 里搜索"【待填:"逐项替换。
    产出 filled.docx。
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docxtpl import DocxTemplate
    from docx import Document
    from docx.shared import RGBColor
    from copy import deepcopy
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[错误] 缺少 docxtpl/python-docx 依赖,请先双击 install.bat 安装依赖。",
          file=sys.stderr)
    sys.exit(1)

try:
    import yaml
except ImportError:
    print("[错误] 缺少 PyYAML 依赖,请先双击 install.bat 安装依赖。",
          file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent
PENDING_USER = '__PENDING_USER__'
SKIP = '__SKIP__'


def _safe_name(name: str) -> str:
    s = re.sub(r'^[一二三四五六七八九十\d]+[、.]\s*', '', name)
    s = re.sub(r'[（(].+?[)）]', '', s)
    s = s.strip()
    return s or 'part'


def get_by_path(data, path: str):
    """按点分路径取值,找不到返回 None。"""
    if not path:
        return None
    keys = path.split('.')
    cur = data
    for k in keys:
        if isinstance(cur, dict):
            if k not in cur:
                return None
            cur = cur[k]
        elif isinstance(cur, list):
            try:
                cur = cur[int(k)]
            except (ValueError, IndexError):
                return None
        else:
            return None
    return cur


def find_company(companies_data: dict, company_id: str) -> dict | None:
    for c in companies_data.get('companies', []):
        if c.get('id') == company_id:
            return c
    return None


def resolve_value(var: dict, brief: dict, companies_data: dict) -> tuple:
    """
    按 var 的 source+path 取值。
    返回 (value, status) — status 是 "resolved" / "pending_user" / "missing" / "manual"

    v2 补丁 1:source=='companies' 时,path 的 company_id 部分不再盲信字面,
    一律用 brief.extracted.bidding_entity 覆盖。
    调用方在 main() 里保证 bidding_entity 已就位(否则提前 raise)。
    """
    source = var.get('source', 'manual')
    path = var.get('path', '')
    name = var.get('name', '')

    if source == 'tender_brief':
        val = get_by_path(brief, path)
        if val is None:
            return None, 'missing'
        if val == PENDING_USER:
            return val, 'pending_user'
        return val, 'resolved'

    if source == 'companies':
        # path 形如 <company_id>.<field>;company_id 一律覆盖为 bidding_entity
        keys = path.split('.', 1)
        if len(keys) != 2:
            return None, 'missing'
        _, field = keys[0], keys[1]
        bidding_entity = (brief.get('extracted') or {}).get('bidding_entity')
        if not bidding_entity:
            # main() 已在入口 raise,此处双保险
            return None, 'missing'
        company = find_company(companies_data, bidding_entity)
        if company is None:
            return None, 'missing'
        val = get_by_path(company, field)
        if val is None:
            return None, 'missing'
        if val == PENDING_USER:
            return val, 'pending_user'
        return val, 'resolved'

    if source == 'manual':
        return None, 'manual'

    return None, 'missing'


PLACEHOLDER_PREFIX = '【待填:'
PLACEHOLDER_SUFFIX = '】'


def make_placeholder(var: dict) -> str:
    """
    v2:非交互占位生成(纯字符串,由后处理着色)。
    文字形如"【待填:变量描述】"。fill 完成后,post_process_highlight_placeholders
    会扫 filled.docx 把"【待填:...】"片段拆出并改为红色加粗 run。
    """
    desc = var.get('description', '') or var.get('name', 'unknown')
    return f'{PLACEHOLDER_PREFIX}{desc}{PLACEHOLDER_SUFFIX}'


def _set_run_red_bold(run):
    """把一个 run 改为红色加粗(保留原字体/字号)。"""
    rPr = run._element.get_or_add_rPr()
    # 清除已有 color 节点
    for c in rPr.findall(qn('w:color')):
        rPr.remove(c)
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), 'FF0000')
    rPr.append(color_el)
    # 加 bold
    for b in rPr.findall(qn('w:b')):
        rPr.remove(b)
    b_el = OxmlElement('w:b')
    rPr.append(b_el)


def _split_run_for_placeholder(run) -> int:
    """
    如果 run.text 含"【待填:xxx】" 片段,把 run 拆成多段,
    中间"【待填:xxx】"部分单独成 run 并改为红色加粗。
    返回拆出的占位 run 数。

    注意:一个 run 可能含多个"【待填:...】" 片段,逐一拆分。
    """
    txt = run.text
    if PLACEHOLDER_PREFIX not in txt:
        return 0

    # 找所有占位片段的 (start, end) 位置
    spans = []
    pos = 0
    while True:
        s = txt.find(PLACEHOLDER_PREFIX, pos)
        if s < 0:
            break
        e = txt.find(PLACEHOLDER_SUFFIX, s + len(PLACEHOLDER_PREFIX))
        if e < 0:
            break
        spans.append((s, e + len(PLACEHOLDER_SUFFIX)))
        pos = e + len(PLACEHOLDER_SUFFIX)

    if not spans:
        return 0

    # 按"前缀 / 占位 / 中缀 / 占位 / 后缀"拆出文本段
    segments = []
    last = 0
    for s, e in spans:
        if s > last:
            segments.append(('plain', txt[last:s]))
        segments.append(('placeholder', txt[s:e]))
        last = e
    if last < len(txt):
        segments.append(('plain', txt[last:]))

    # 第一段保留在原 run,后续段在原 run 后方依次插入新 w:r
    first_kind, first_text = segments[0]
    run.text = first_text
    if first_kind == 'placeholder':
        _set_run_red_bold(run)

    # 依次插入后续段
    current_element = run._element
    for kind, seg_text in segments[1:]:
        # 深拷贝 run 结构(含 rPr),重置文本和颜色
        new_r = deepcopy(run._element)
        # 重置 text
        for t in new_r.findall(qn('w:t')):
            new_r.remove(t)
        # 清除继承来的 color/bold,让新 run 回到 plain
        new_rPr = new_r.find(qn('w:rPr'))
        if new_rPr is not None:
            for c in new_rPr.findall(qn('w:color')):
                new_rPr.remove(c)
            for b in new_rPr.findall(qn('w:b')):
                new_rPr.remove(b)
        new_t = OxmlElement('w:t')
        new_t.set(qn('xml:space'), 'preserve')
        new_t.text = seg_text
        new_r.append(new_t)
        # 如果是 placeholder,加红色加粗
        if kind == 'placeholder':
            if new_rPr is None:
                new_rPr = OxmlElement('w:rPr')
                new_r.insert(0, new_rPr)
            color_el = OxmlElement('w:color')
            color_el.set(qn('w:val'), 'FF0000')
            new_rPr.append(color_el)
            b_el = OxmlElement('w:b')
            new_rPr.append(b_el)
        # 插入到当前元素之后
        current_element.addnext(new_r)
        current_element = new_r

    return sum(1 for k, _ in segments if k == 'placeholder')


def post_process_highlight_placeholders(filled_path: Path) -> int:
    """
    打开 filled.docx,扫所有 run 把"【待填:...】"片段拆出并改为红色加粗。
    返回高亮的占位数。
    """
    doc = Document(str(filled_path))
    highlighted = 0

    def _process_paragraphs(paragraphs):
        nonlocal highlighted
        for para in paragraphs:
            # 遍历时会插入新 run,需要快照当前 runs
            for run in list(para.runs):
                highlighted += _split_run_for_placeholder(run)

    _process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _process_paragraphs(cell.paragraphs)

    doc.save(str(filled_path))
    return highlighted


def main():
    parser = argparse.ArgumentParser(
        description='C 模式模板填充(阶段 4-C Step 5-6)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument('--project', required=True, help='项目名')
    parser.add_argument('--part', type=int, required=True,
                        help='response_file_parts 索引(0 基)')
    args = parser.parse_args()

    project_dir = ROOT / 'projects' / args.project

    # V53: 未 review 则硬失败
    from brief_schema import ensure_reviewed
    ensure_reviewed(project_dir / 'output')

    brief_path = project_dir / 'output' / 'tender_brief.json'
    if not brief_path.exists():
        print(f"[错误] 找不到 tender_brief.json: {brief_path}", file=sys.stderr)
        sys.exit(1)
    brief = json.loads(brief_path.read_text(encoding='utf-8'))

    # v2 补丁 1:bidding_entity 必须已就位
    bidding_entity = (brief.get('extracted') or {}).get('bidding_entity')
    if not bidding_entity:
        print("[错误] tender_brief.json 的 extracted.bidding_entity 未设置。",
              file=sys.stderr)
        print("  请先跑:./run_script.bat select_bidding_entity.py "
              f"--project {args.project}", file=sys.stderr)
        sys.exit(1)

    companies_path = ROOT / 'companies.yaml'
    if not companies_path.exists():
        print(f"[错误] 找不到 companies.yaml: {companies_path}", file=sys.stderr)
        sys.exit(1)
    with open(companies_path, 'r', encoding='utf-8') as f:
        companies_data = yaml.safe_load(f)

    parts = brief.get('response_file_parts', [])
    if args.part < 0 or args.part >= len(parts):
        print(f"[错误] --part 超出范围", file=sys.stderr)
        sys.exit(1)
    part_name = parts[args.part]['name']
    part_dir_name = _safe_name(part_name)
    c_dir = project_dir / 'output' / 'c_mode' / part_dir_name

    # V60: sub_mode 分流
    from brief_schema import resolve_sub_mode
    sub_mode = resolve_sub_mode(parts[args.part])

    if sub_mode == 'C-attachment':
        raise NotImplementedError(
            f"Part[{args.part}] '{part_name}' sub_mode=C-attachment 当前挂档,"
            f"见 business_model §8 #N20"
        )

    if sub_mode == 'C-reference':
        # C-reference: 无 fill 动作, 仅校验 instructions.md 存在
        instructions_path = c_dir / 'instructions.md'
        if not instructions_path.exists():
            print(f"[错误] 找不到 instructions.md: {instructions_path}", file=sys.stderr)
            print(f"  请先跑 c_mode_extract.py 产出 instructions", file=sys.stderr)
            sys.exit(1)
        size = instructions_path.stat().st_size
        print(f"[信息] Part[{args.part}] '{part_name}' (C-reference) 校验通过")
        print(f"  instructions.md: {instructions_path} ({size:,} bytes)")
        print(f"  C-reference 不产 filled.docx, 命令行接口保持一致(do nothing)。")
        return

    # 以下为 C-template 路径(默认)
    template_path = c_dir / 'template.docx'
    vars_path = c_dir / 'variables.yaml'
    filled_path = c_dir / 'filled.docx'

    if not template_path.exists():
        print(f"[错误] 找不到 template.docx: {template_path}", file=sys.stderr)
        print(f"  请先跑 c_mode_extract.py 产出模板", file=sys.stderr)
        sys.exit(1)
    if not vars_path.exists():
        print(f"[错误] 找不到 variables.yaml: {vars_path}", file=sys.stderr)
        sys.exit(1)

    with open(vars_path, 'r', encoding='utf-8') as f:
        vars_config = yaml.safe_load(f)

    variables = vars_config.get('variables', [])

    print(f"[信息] 填充 Part[{args.part}] '{part_name}',共 {len(variables)} 个变量")

    context: dict = {}
    stats = {'resolved': 0, 'placeholder_pending_user': 0,
             'placeholder_manual': 0, 'missing': 0}

    for var in variables:
        name = var['name']
        val, status = resolve_value(var, brief, companies_data)

        if status == 'resolved':
            context[name] = val
            stats['resolved'] += 1
            continue

        if status == 'missing':
            # 硬失败:path 不通(variables.yaml 里给了无效路径,需要修 yaml 不是填值)
            print(f"[错误] 变量 '{name}' path 不通: "
                  f"source={var.get('source')} path={var.get('path')} 不存在",
                  file=sys.stderr)
            sys.exit(1)

        # v2:pending_user 或 manual → 非交互,写入"【待填:xxx】"占位
        context[name] = make_placeholder(var)
        if status == 'pending_user':
            stats['placeholder_pending_user'] += 1
        else:
            stats['placeholder_manual'] += 1

    # docxtpl 渲染
    tpl = DocxTemplate(str(template_path))
    tpl.render(context)
    tpl.save(str(filled_path))

    # v2:后处理 — 把"【待填:...】"占位拆出并改为红色加粗 run
    highlighted = post_process_highlight_placeholders(filled_path)

    total_placeholder = stats['placeholder_pending_user'] + stats['placeholder_manual']
    print()
    print("=" * 60)
    print(f"[完成] filled.docx: {filled_path}")
    print(f"  已自动取值: {stats['resolved']} 个")
    print(f"  显式占位: {total_placeholder} 个 "
          f"(红色加粗 run 共 {highlighted} 处;在 Word 中搜'【待填:' 逐项替换)")
    print(f"    - pending_user 占位: {stats['placeholder_pending_user']} 个"
          f"(companies.yaml 对应字段需用户更新)")
    print(f"    - manual 占位: {stats['placeholder_manual']} 个"
          f"(项目特定变量,用户手工在 Word 中搜索【待填:】替换)")
    print("=" * 60)


if __name__ == '__main__':
    main()
