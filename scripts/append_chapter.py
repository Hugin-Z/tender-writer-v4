# -*- coding: utf-8 -*-
"""
append_chapter.py · 将 markdown 章节追加到现有 docx（阶段 4 辅助脚本）

输入:
    output/tender_response.docx
    章节 markdown 文件

能力:
    - 识别 # / ## / ### / #### 标题
    - 识别普通段落（中文相邻不加空格）
    - 识别 markdown 表格 → 调用 docx_builder.add_table
    - 识别图占位符 【图 X.Y:说明】 → 调用 docx_builder.add_figure_caption
    - 识别 - / * 无序列表和 1. 有序列表 → Word 内置 List 样式
    - 跳过 HTML 注释行
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("[错误] 缺少 python-docx 依赖，请先双击 install.bat 安装依赖。", file=sys.stderr)
    sys.exit(1)

from docx_builder import (
    add_chapter, add_paragraph, add_table, add_figure_caption,
    apply_default_styles, set_page_margins, set_run_font, to_simplified,
    clean_docx_whitespace,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt


# ─────────────────────────────────────────────
# 图占位符识别
# ─────────────────────────────────────────────

FIGURE_PATTERN = re.compile(
    r'^\*?\*?【图\s*[\d.]+\s*[:：]\s*(.+?)】\*?\*?$'
)


def is_figure_caption(line: str) -> tuple[bool, str | None]:
    """整行图注识别。返回 (is_caption, caption_text)"""
    stripped = line.strip()
    m = FIGURE_PATTERN.match(stripped)
    if m:
        return True, m.group(1).strip()
    return False, None


# ─────────────────────────────────────────────
# 列表识别
# ─────────────────────────────────────────────

LIST_STYLE_MAP = {
    ('unordered', 0): 'List Bullet',
    ('unordered', 1): 'List Bullet 2',
    ('ordered', 0): 'List Number',
    ('ordered', 1): 'List Number 2',
}

_ORDERED_RE = re.compile(r'^(\d+)\.\s+(.+)$')


def parse_list_item(line: str) -> tuple[str, int, str] | None:
    """
    尝试将一行解析为列表项。
    返回 (list_type, level, text) 或 None。
    """
    # 计算缩进级别
    stripped = line.lstrip()
    indent = len(line) - len(stripped)
    level = min(indent // 2, 1)  # 0 或 1，更深截断到 1

    if stripped.startswith("- ") or stripped.startswith("* "):
        return ('unordered', level, stripped[2:].strip())

    m = _ORDERED_RE.match(stripped)
    if m:
        return ('ordered', level, m.group(2).strip())

    return None


def add_list_item(doc, text: str, list_type: str = 'unordered', level: int = 0):
    """
    添加列表项。样式不存在时硬失败。
    v2:支持内联 markdown(**bold** / *italic* / `code`)。
    """
    style_name = LIST_STYLE_MAP.get((list_type, level))
    if style_name is None:
        raise ValueError(f"unsupported list ({list_type}, level={level})")

    try:
        p = doc.add_paragraph(style=style_name)
    except KeyError:
        raise RuntimeError(
            f"Word 内置样式 '{style_name}' 不存在于当前 docx。"
            f"需要 create_section_doc 已经预激活该样式。"
        )

    text = to_simplified(text)
    segments = _parse_inline(text)
    if not segments:
        segments = [("plain", text)]
    for kind, content in segments:
        if not content:
            continue
        run = p.add_run(content)
        if kind == "bold":
            set_run_font(run, east_asia="宋体", size_pt=14, bold=True)
        elif kind == "italic":
            set_run_font(run, east_asia="宋体", size_pt=14, italic=True)
        elif kind == "code":
            set_run_font(run, ascii_font="Consolas", east_asia="宋体",
                         size_pt=13, bold=False)
        else:
            set_run_font(run, east_asia="宋体", size_pt=14, bold=False)
    return p


# ─────────────────────────────────────────────
# Markdown 表格解析
# ─────────────────────────────────────────────

_TABLE_LINE_RE = re.compile(r'^\|.*\|$')
_SEPARATOR_RE = re.compile(r'^\|[\s\-:|]+\|$')


def is_table_line(line: str) -> bool:
    """判断一行是否为 markdown 表格行"""
    stripped = line.strip()
    return bool(_TABLE_LINE_RE.match(stripped)) and stripped.count('|') >= 3


def parse_markdown_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """
    解析连续的 markdown 表格行。
    返回 (headers, rows)。
    """
    if not lines:
        return [], []

    def split_cells(line: str) -> list[str]:
        stripped = line.strip()
        if stripped.startswith('|'):
            stripped = stripped[1:]
        if stripped.endswith('|'):
            stripped = stripped[:-1]
        return [cell.strip() for cell in stripped.split('|')]

    headers = split_cells(lines[0])
    rows = []
    for line in lines[1:]:
        if _SEPARATOR_RE.match(line.strip()):
            continue
        cells = split_cells(line)
        # Pad or trim to match header count
        while len(cells) < len(headers):
            cells.append("")
        rows.append(cells[:len(headers)])

    return headers, rows


# ─────────────────────────────────────────────
# 段落缓冲与中文合并
# ─────────────────────────────────────────────

def _is_chinese_boundary(left_char: str, right_char: str) -> bool:
    """两个字符相邻时是否应该不加空格（都是中文/中文标点）"""
    def is_cjk(c):
        return '\u4e00' <= c <= '\u9fff' or c in '，。；：、！？""''（）《》【】'
    return is_cjk(left_char) or is_cjk(right_char)


# v2:内联 markdown 解析
# 支持 **bold** / *italic* / `code` 三种内联标记,拆成多个 run
# 使用非贪婪匹配,避免跨多个标记段落的误匹配
_INLINE_PATTERN = re.compile(
    r"(\*\*[^*]+?\*\*"   # **bold** — 不含 * 的最短匹配
    r"|\*[^*\s][^*]*?\*"  # *italic* — 不含 * 的最短,且首字符非空白
    r"|`[^`]+?`)"         # `code`
)


def _parse_inline(text: str) -> list[tuple[str, str]]:
    """
    把一行文本按内联 markdown 拆成 (kind, content) 对,kind ∈ {'plain','bold','italic','code'}。
    """
    out: list[tuple[str, str]] = []
    pos = 0
    for m in _INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            out.append(("plain", text[pos:m.start()]))
        token = m.group(1)
        if token.startswith("**") and token.endswith("**"):
            out.append(("bold", token[2:-2]))
        elif token.startswith("`") and token.endswith("`"):
            out.append(("code", token[1:-1]))
        elif token.startswith("*") and token.endswith("*"):
            out.append(("italic", token[1:-1]))
        else:
            out.append(("plain", token))
        pos = m.end()
    if pos < len(text):
        out.append(("plain", text[pos:]))
    return out


def _add_paragraph_with_inline(document, text: str, style: str = "Normal",
                               size_pt: float = 14,
                               first_line_indent_chars: int = 2):
    """
    v2:支持内联 markdown 的段落添加器。
    **bold** → bold run;*italic* → italic run;`code` → 等宽字体 run。
    """
    text = to_simplified(text)
    p = document.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if first_line_indent_chars > 0 and style == "Normal":
        pf.first_line_indent = Pt(size_pt * first_line_indent_chars)

    segments = _parse_inline(text)
    if not segments:
        segments = [("plain", text)]
    for kind, content in segments:
        if not content:
            continue
        run = p.add_run(content)
        if kind == "bold":
            set_run_font(run, east_asia="宋体", size_pt=size_pt, bold=True)
        elif kind == "italic":
            set_run_font(run, east_asia="宋体", size_pt=size_pt, bold=False,
                         italic=True)
        elif kind == "code":
            set_run_font(run, ascii_font="Consolas", east_asia="宋体",
                         size_pt=size_pt - 1, bold=False)
        else:
            set_run_font(run, east_asia="宋体", size_pt=size_pt, bold=False)
    return p


def flush_paragraph(document, buffer: list[str], stats: dict | None = None) -> None:
    if not buffer:
        return
    parts = [p.strip() for p in buffer if p.strip()]
    if not parts:
        buffer.clear()
        return

    text = parts[0]
    for nxt in parts[1:]:
        if text and nxt and _is_chinese_boundary(text[-1], nxt[0]):
            text += nxt
        else:
            text += " " + nxt

    # v2:含 markdown 内联标记时走内联解析器
    if "**" in text or "`" in text or re.search(r"\*[^*\s]", text):
        _add_paragraph_with_inline(document, text)
    else:
        add_paragraph(document, text)
    if stats is not None:
        stats["paragraphs"] += 1
    buffer.clear()


# ─────────────────────────────────────────────
# 主循环（状态机）
# ─────────────────────────────────────────────

def append_markdown(document: Document, markdown_text: str) -> dict:
    """
    追加 markdown 到 docx。返回统计 dict。
    """
    paragraph_buffer: list[str] = []
    table_buffer: list[str] = []
    in_comment = False

    stats = {
        "headings": 0,
        "paragraphs": 0,
        "tables": 0,
        "figures": 0,
        "list_items": 0,
    }

    def flush_table():
        if not table_buffer:
            return
        headers, rows = parse_markdown_table(table_buffer)
        if headers:
            add_table(document, headers, rows)
            stats["tables"] += 1
        table_buffer.clear()

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        # 多行 HTML 注释处理
        if in_comment:
            if '-->' in stripped:
                in_comment = False
            continue
        if stripped.startswith('<!--'):
            if '-->' not in stripped:
                in_comment = True
            continue

        # 空行
        if not stripped:
            if table_buffer:
                flush_table()
            flush_paragraph(document, paragraph_buffer, stats)
            continue

        # 表格行
        if is_table_line(stripped):
            flush_paragraph(document, paragraph_buffer, stats)
            table_buffer.append(stripped)
            continue

        # 非表格行但表格缓冲非空 → flush 表格
        if table_buffer:
            flush_table()

        # 图占位符
        is_fig, caption_text = is_figure_caption(stripped)
        if is_fig:
            flush_paragraph(document, paragraph_buffer, stats)
            add_figure_caption(document, caption_text)
            stats["figures"] += 1
            continue

        # 标题
        if stripped.startswith('#'):
            flush_paragraph(document, paragraph_buffer, stats)
            level = len(stripped) - len(stripped.lstrip('#'))
            level = min(max(level, 1), 4)
            title = stripped[level:].strip()
            add_chapter(document, title, level=level)
            stats["headings"] += 1
            continue

        # 列表
        list_result = parse_list_item(line)
        if list_result:
            flush_paragraph(document, paragraph_buffer, stats)
            list_type, level, text = list_result
            add_list_item(document, text, list_type, level)
            stats["list_items"] += 1
            continue

        # 其他: 累积到段落缓冲
        paragraph_buffer.append(stripped)

    # 末尾 flush
    if table_buffer:
        flush_table()
    flush_paragraph(document, paragraph_buffer, stats)

    return stats


# ─────────────────────────────────────────────
# 命令行入口
# ─────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="将 markdown 章节追加到现有 docx",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("docx_path", help="现有 docx 路径")
    parser.add_argument("chapter_md", help="章节 markdown 文件路径")
    args = parser.parse_args()

    docx_path = Path(args.docx_path)
    chapter_md = Path(args.chapter_md)

    # V53: 未 review 则硬失败(docx_path 父目录即 output 目录)
    from brief_schema import ensure_reviewed
    ensure_reviewed(docx_path.parent)

    if not docx_path.exists():
        print(f"[错误] 找不到 docx 文件: {docx_path}", file=sys.stderr)
        sys.exit(1)
    if not chapter_md.exists():
        print(f"[错误] 找不到 markdown 文件: {chapter_md}", file=sys.stderr)
        sys.exit(1)

    document = Document(str(docx_path))
    stats = append_markdown(document, chapter_md.read_text(encoding="utf-8"))

    # v2:追加完后统一清理中英/中数之间的空格
    cleaned = clean_docx_whitespace(document)

    document.save(str(docx_path))

    print(f"[完成] 已将章节追加到: {docx_path}")
    print(f"  标题: {stats['headings']}, 段落: {stats['paragraphs']}, "
          f"表格: {stats['tables']}, 图注: {stats['figures']}, "
          f"列表项: {stats['list_items']}, "
          f"空格清理: {cleaned} 处 run")


if __name__ == "__main__":
    main()
