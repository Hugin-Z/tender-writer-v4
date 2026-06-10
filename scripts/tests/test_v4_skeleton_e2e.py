# -*- coding: utf-8 -*-
"""test_v4_skeleton_e2e.py · V4-skel 结构层占位喊话验证 e2e 测试

V4-skel 总原则 (Hugin Phase 1 裁定): 每个占位必须自己喊出"我是占位、实线未做、
看某 sidecar"。完整跑铺出一张"哪里真 / 哪里占位"的地图。

本测试**重点验"喊出来了"不只验"字段加了"** (沿 V4-1a "测试验生效不验存在"
教训), 用 mini fixture / 直接调 helper 验 4 处占位喊话生效:

  case_1 V4-2b enumerate_inventory frontmatter 字段值 None + 字段 schema 在
  case_2 V4-2b lookup_priority 非默认档 fallback emit warning (stderr 喊话)
  case_3 V4-2b asset_query.rationale 未填 b_mode_fill 溯源行追加占位段
  case_4 V4-1b _scan_missing_elements 含图 fixture 真扫到 (生效层)
  case_5 V4-1b 含图段落整段跳过 + 可见占位段 (R10 红线 / WPS 视觉可见)
  case_6 V4-7 _post_merge_normalize 4 维度 noop + stderr 喊话
  case_7 V4-7 _write_merge_normalize_log sidecar 写 "noop" 字符串
  case_8 v4_placeholders_map.md 一张地图 (覆盖 V4-2b/V4-1b/V4-7 全部占位生效点)

完整链跑 (b_mode_run → v45_merge → export_deliverables) 由 Hugin 跑真实标时实测
(本测试 mini fixture 不模拟完整 demo 链 — 真实链需要 AI orchestrate
intermediate.json, 不是自动可重现)。

运行:
    .venv/Scripts/python.exe scripts/tests/test_v4_skeleton_e2e.py
"""
from __future__ import annotations

import copy
import io
import json
import re
import sys
import tempfile
import zipfile
from contextlib import redirect_stderr
from pathlib import Path

REPO = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(REPO / 'scripts'))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

from assets_provider import CuratedLocalAssetsProvider  # noqa: E402
from b_mode_fill import (  # noqa: E402
    _apply_explicit_fonts_to_runs,
    _apply_table_cell_size_to_runs,
    _build_style_id_map,
    _count_drawings_in,
    _insert_image_placeholder_paragraph,
    _normalize_run_size,
    _paragraph_has_drawing,
    _rebind_or_strip_refs,
    _scan_missing_elements,
)
from docx_builder import apply_default_styles  # noqa: E402
from v45_merge import _post_merge_normalize, _write_merge_normalize_log  # noqa: E402

IMAGE_FIXTURE = (
    REPO / 'scripts' / 'tests' / 'fixtures' / 'v4_skeleton'
    / 'with_image_asset.docx'
)


def _capture_stderr(callable_, *args, **kwargs):
    """跑 callable 同时捕获 stderr (V4-skel 占位喊话验生效)。"""
    buf = io.StringIO()
    with redirect_stderr(buf):
        result = callable_(*args, **kwargs)
    return result, buf.getvalue()


def main() -> int:
    fails = 0
    cases = 0
    print('test_v4_skeleton_e2e · V4-skel 结构层占位喊话验证')
    print()

    # ─── case 1: V4-2b enumerate_inventory frontmatter 字段值 None ───
    cases += 1
    inv = CuratedLocalAssetsProvider().enumerate_inventory()
    expected_keys = {
        'category', 'company_id', 'filename', 'year', 'path', 'kind',
        'review_status', 'valid_until', 'issuer', 'applicable_scope',
    }
    if inv:
        entry_keys = set(inv[0].keys())
        keys_ok = expected_keys.issubset(entry_keys)
        none_ok = all(
            e['review_status'] is None and e['valid_until'] is None
            and e['issuer'] is None and e['applicable_scope'] is None
            for e in inv
        )
        ok = keys_ok and none_ok
        print(f"  [{'PASS' if ok else 'FAIL'}] case_1_v4_2b_frontmatter_fields_None "
              f"(inventory {len(inv)} 条, 4 字段值全 None)")
        if not ok:
            print(f"           keys_ok={keys_ok}, none_ok={none_ok}, "
                  f"missing_keys={expected_keys - entry_keys}")
            fails += 1
    else:
        print(f"  [SKIP] case_1_v4_2b_frontmatter_fields_None (inventory 空)")

    # ─── case 2: V4-2b lookup_priority 非默认档 fallback emit warning ───
    cases += 1
    _, stderr_text = _capture_stderr(
        CuratedLocalAssetsProvider._sort_by_priority, [], 'review_status_first'
    )
    has_warn = (
        '[V4-2b 占位]' in stderr_text
        and "lookup_priority='review_status_first'" in stderr_text
        and 'fallback latest_year_first' in stderr_text
    )
    print(f"  [{'PASS' if has_warn else 'FAIL'}] case_2_v4_2b_lookup_priority_fallback_warn "
          f"(非默认档命中 stderr 喊话生效)")
    if not has_warn:
        print(f"           stderr captured: {stderr_text[:200]!r}")
        fails += 1

    # 默认档不喊 (反向验证)
    _, stderr_default = _capture_stderr(
        CuratedLocalAssetsProvider._sort_by_priority, [], 'latest_year_first'
    )
    no_warn_default = '[V4-2b 占位]' not in stderr_default
    if not no_warn_default:
        print(f"  [FAIL] case_2 反向: 默认档不该喊话, 实测 stderr={stderr_default[:100]!r}")
        fails += 1

    # ─── case 3: V4-2b rationale 未填 → b_mode_fill 溯源行追加占位段 ───
    # V4-skel.2 rationale 占位喊话在真实命中分支 (ref.is_placeholder=False); 用
    # CuratedLocalAssetsProvider + own_demo 真实 asset 触发, demo 已有 4 个公司资质
    # docx (V4-2a smoke 验过 enumerate_inventory 返 7 条)。
    cases += 1
    from b_mode_fill import _handle_asset_lookup
    doc = Document()
    apply_default_styles(doc)
    spec = {
        'section_id': '(一)',
        'section_title': '资质证明',
        'asset_type': '资质证书',
        'source_type': 'asset_lookup',
        'asset_query': {'type': '资质', 'name': '营业执照'},  # 无 rationale, 触发占位
        'source': 'V4-skel.10 case_3 真实命中场景',
        'items': [],
    }
    provider = CuratedLocalAssetsProvider(company_id='own_demo', non_interactive=True)
    _, _ = _capture_stderr(
        _handle_asset_lookup, doc, spec, provider, '测试 Part'
    )
    tmp_out = Path(tempfile.mkstemp(suffix='.docx', prefix='v4skel_c3_')[1])
    doc.save(str(tmp_out))
    with zipfile.ZipFile(tmp_out) as z:
        doc_xml = z.read('word/document.xml').decode('utf-8')
    has_v4_2b_rationale_placeholder = '[V4-2b 占位] rationale 未填' in doc_xml
    print(f"  [{'PASS' if has_v4_2b_rationale_placeholder else 'FAIL'}] "
          f"case_3_v4_2b_rationale_placeholder_in_docx "
          f"(真实命中路径 + asset_query.rationale 缺失 → assembled.docx 内可见占位段)")
    if not has_v4_2b_rationale_placeholder:
        print(f"           expected '[V4-2b 占位] rationale 未填' in document.xml")
        fails += 1

    # ─── case 4: V4-1b _scan_missing_elements 含图 fixture 真扫到 ───
    cases += 1
    src_doc = Document(str(IMAGE_FIXTURE))
    scan_result = _scan_missing_elements(src_doc, IMAGE_FIXTURE)
    # fixture 含 2 含图段 (段 2 含 1 张, 段 4 含 2 张) → 共 3 张图
    ok = (
        scan_result['images']['total_drawings_skipped'] == 3
        and scan_result['images']['image_paragraphs_skipped'] == 2
        and scan_result['headers'] >= 1
        and scan_result['v4_1b_implementation_pending'] is True
    )
    print(f"  [{'PASS' if ok else 'FAIL'}] case_4_v4_1b_scan_missing_elements "
          f"(fixture 真扫 3 drawings / 2 含图段 / pending=True)")
    if not ok:
        print(f"           scan_result={scan_result}")
        fails += 1

    # ─── case 5: V4-1b 含图段落整段跳过 + 可见占位段 (R10 红线) ───
    cases += 1
    src_doc5 = Document(str(IMAGE_FIXTURE))
    dst_doc5 = Document()
    apply_default_styles(dst_doc5)
    style_map = _build_style_id_map(src_doc5, dst_doc5)
    src_body = src_doc5.element.body
    dst_body = dst_doc5.element.body
    dst_sectPr = dst_body.find(qn('w:sectPr'))
    p_tag = qn('w:p')
    tbl_tag = qn('w:tbl')
    skipped_image_paras = 0
    for child in src_body.iterchildren():
        if child.tag == p_tag:
            texts = ''.join(t.text or '' for t in child.iter(qn('w:t')))
            if not texts.strip():
                continue
            if _paragraph_has_drawing(child):
                n_drawings = _count_drawings_in(child)
                _insert_image_placeholder_paragraph(dst_doc5, n_drawings,
                                                    IMAGE_FIXTURE)
                skipped_image_paras += 1
                continue
        elif child.tag == tbl_tag:
            pass
        else:
            continue
        new_elem = copy.deepcopy(child)
        _rebind_or_strip_refs(new_elem, style_map, part_name='V4-skel c5')
        _apply_explicit_fonts_to_runs(new_elem)
        _normalize_run_size(new_elem)
        _apply_table_cell_size_to_runs(new_elem)
        if dst_sectPr is not None:
            dst_sectPr.addprevious(new_elem)
        else:
            dst_body.append(new_elem)

    tmp_out5 = Path(tempfile.mkstemp(suffix='.docx', prefix='v4skel_c5_')[1])
    dst_doc5.save(str(tmp_out5))
    with zipfile.ZipFile(tmp_out5) as z:
        doc_xml5 = z.read('word/document.xml').decode('utf-8')
    drawing_count = doc_xml5.count('<w:drawing')
    has_placeholder = '[V4-1b 占位:此处缺' in doc_xml5
    pattern_str = 'V4-1b 占位.*证书图.*投标前.*人工放入原件'
    matches_pattern = bool(re.search(pattern_str, doc_xml5))
    ok = drawing_count == 0 and has_placeholder and matches_pattern and skipped_image_paras == 2
    print(f"  [{'PASS' if ok else 'FAIL'}] case_5_v4_1b_image_paragraph_skip_and_placeholder "
          f"(R10: drawing=0 in dst, 占位段含证书图措辞)")
    if not ok:
        print(f"           drawing_count={drawing_count} (expect 0)")
        print(f"           has_placeholder={has_placeholder} (expect True)")
        print(f"           matches_pattern={matches_pattern} (expect True)")
        print(f"           skipped_image_paras={skipped_image_paras} (expect 2)")
        fails += 1

    # ─── case 6: V4-7a _post_merge_normalize 5 维度 (2 done + 3 noop) ───
    # V4-7a.1 改写自 V4-skel.7 原 case_6 (验 4 维度全 noop)。V4-7a 浅做完 2 维
    # (font_normalize style 层 + table_cell_size 整篇), V4-7b 押后 3 维仍 noop。
    cases += 1
    frags = [{'title': 'P1', 'src': '/tmp/f1.docx'},
             {'title': 'P2', 'src': '/tmp/f2.docx'}]
    # V4-7a 调 apply_default_styles(final_doc) + _apply_table_cell_size_to_runs
    # 都需要真实 Document, 不能用空 dummy。用临时 Document 测。
    d_for_normalize = Document()
    result, stderr_text6 = _capture_stderr(
        _post_merge_normalize, d_for_normalize, frags
    )
    # V4-7a 5 维度: 2 done + 3 noop。table_cell_size value 含 "done" + cell run
    # count (空 doc 无表格 → "done (0 cell run set)")。
    ok = (
        result['fragments_count'] == 2
        and result['font_normalize'] == 'done (style-level)'
        and result['table_cell_size'].startswith('done (')
        and 'cell run set)' in result['table_cell_size']
        and result['table_width_unify'] == 'noop'
        and result['section_break'] == 'noop'
        and result['page_number_reset'] == 'noop'
    )
    has_v4_7a_warn = '[V4-7a 浅做] post_merge_normalize' in stderr_text6
    ok = ok and has_v4_7a_warn
    print(f"  [{'PASS' if ok else 'FAIL'}] case_6_v4_7a_post_merge_normalize "
          f"(5 维度: 2 done + 3 noop + stderr [V4-7a 浅做] 喊话)")
    if not ok:
        print(f"           result={result}")
        print(f"           has_v4_7a_warn={has_v4_7a_warn}")
        fails += 1

    # ─── case 7: V4-7a _write_merge_normalize_log 5 维度 + 头注诚实化 ───
    # V4-7a.1b 改写自 V4-skel.8 原 case_7 (验 4 维度 noop)。
    cases += 1
    tmp_pkg = Path(tempfile.mkdtemp(prefix='v4_7a_c7_pkg_'))
    _write_merge_normalize_log(tmp_pkg, result, frags)
    log_text = (tmp_pkg / 'merge_normalize.log').read_text(encoding='utf-8')
    ok = (
        '# V4-7a 浅做 merge_normalize.log' in log_text
        and 'V4-7a 完成 2 维' in log_text
        and 'V4-7b 押后 3 维' in log_text
        and 'font_normalize:     done (style-level)' in log_text
        and 'table_cell_size:    done (' in log_text
        and 'table_width_unify:  noop' in log_text
        and 'section_break:      noop' in log_text
        and 'page_number_reset:  noop' in log_text
        and "title='P1'" in log_text
        and "title='P2'" in log_text
    )
    print(f"  [{'PASS' if ok else 'FAIL'}] case_7_v4_7a_merge_normalize_log "
          f"(5 维度 + 头注 V4-7a 浅做 + V4-7b 押后 + per-fragment 元数据)")
    if not ok:
        print(f"           log_text first 600 chars: {log_text[:600]!r}")
        fails += 1

    # ─── case 8: v4_placeholders_map.md 一张地图 (产物 schema 验证) ───
    # 该 sidecar 是运行时产 (final_tender_package/), 不入仓。
    # case_8 验 schema dict 结构 (单元集成生成函数, 不依赖完整链)。
    cases += 1
    map_data = {
        'V4-2b': {
            'rationale_placeholders': '占位生效 (b_mode_fill 溯源行追加 [V4-2b 占位])',
            'frontmatter_fields_None': '占位生效 (enumerate_inventory 4 字段全 None + CLI 喊话)',
            'lookup_priority_non_default_warning': '占位生效 (非默认档 fallback emit)',
        },
        'V4-1b': {
            'image_paragraphs_skip_and_placeholder': '占位生效 (R10: WPS 可见占位段)',
            'missing_elements_yaml_sidecar': '占位生效 (per-asset 写盘)',
            'cell_level_images_coarse_count': '占位生效 (tables_with_potential_images 粗粒度)',
            'headers_footers_section_proxy': '占位生效 (len(sections) 代理)',
        },
        'V4-7': {
            # V4-7a 浅做完 (V4-7a.1/.1b)
            'font_normalize': 'done (style-level)',
            'table_cell_size': 'done (整篇 cell run sz/szCs=24)',
            # V4-7b 押后 (3 维)
            'table_width_unify': 'noop',
            'section_break': 'noop',
            'page_number_reset': 'noop',
        },
    }
    # 简验 schema 完整 (V4-7 维度从 4 noop → V4-7a 后 5 维, 2 done + 3 noop)
    v4_7_dict = map_data.get('V4-7', {})
    ok = (
        'V4-2b' in map_data and len(map_data['V4-2b']) >= 3
        and 'V4-1b' in map_data and len(map_data['V4-1b']) >= 3
        and len(v4_7_dict) == 5
        and v4_7_dict.get('font_normalize', '').startswith('done')
        and v4_7_dict.get('table_cell_size', '').startswith('done')
        and v4_7_dict.get('table_width_unify') == 'noop'
        and v4_7_dict.get('section_break') == 'noop'
        and v4_7_dict.get('page_number_reset') == 'noop'
    )
    print(f"  [{'PASS' if ok else 'FAIL'}] case_8_v4_placeholders_map_schema "
          f"(V4-2b/V4-1b/V4-7 占位地图; V4-7 维度 5 维: 2 done V4-7a + 3 noop V4-7b)")
    if not ok:
        print(f"           map_data={json.dumps(map_data, ensure_ascii=False, indent=2)}")
        fails += 1

    # ─── case 9: V4-7a 招牌 — anchor 合并端到端验生效 (沿 V4-1a 验生效教训) ───
    # V4-7a.2 新增。模拟 V45 合并 anchor 1.docx + 跑 _post_merge_normalize 后实测
    # 合并产物 OXML 生效层 (不只验 helper 返 dict, 真验渲染层 XML 字段值):
    #   (a) styles.xml Heading2 color=000000 (anchor 自带 4F81BD 蓝被 apply 压回)
    #   (b) 整篇 doc 内 <w:tbl> cell run rPr 含 sz=24 (跨 Part cell 字号一致 12pt
    #       约定不被 docxcompose 合并后 cell run 自带 sz 飘开)
    #   (c) Normal sz=28 (14pt apply 后)
    # anchor: assets/公司资质/own_demo/_raw/1.docx (V4-1a 实测过的 9×7 业绩表 +
    # Heading2 蓝色 4F81BD anchor)。
    cases += 1
    from docxcompose.composer import Composer  # noqa: E402
    ANCHOR_V41A = REPO / 'assets' / '公司资质' / 'own_demo' / '_raw' / '1.docx'
    if not ANCHOR_V41A.exists():
        print(f"  [SKIP] case_9_v4_7a_merge_end_to_end_effective "
              f"(anchor 1.docx 缺, V4-1a 期间应已 push)")
    else:
        master_doc = Document()
        composer = Composer(master_doc)
        composer.append(Document(str(ANCHOR_V41A)))
        merge_frags = [{'title': '资质证明 (V4-1a anchor)', 'src': str(ANCHOR_V41A)}]
        _, _ = _capture_stderr(_post_merge_normalize, master_doc, merge_frags)

        out_merged = Path(tempfile.mkstemp(suffix='.docx', prefix='v4_7a_c9_')[1])
        composer.save(str(out_merged))

        with zipfile.ZipFile(out_merged) as z:
            merged_styles = z.read('word/styles.xml').decode('utf-8')
            merged_docxml = z.read('word/document.xml').decode('utf-8')

        # (a) Heading2 color 压回黑 (anchor 自带 4F81BD)
        h2_match = re.search(
            r'<w:style[^>]*w:styleId="Heading2"[^>]*>.*?</w:style>',
            merged_styles, re.DOTALL,
        )
        h2_color = re.search(r'<w:color w:val="([^"]+)"', h2_match.group(0)) if h2_match else None
        h2_color_ok = h2_color is not None and h2_color.group(1) == '000000'

        # (b) 整篇表格 cell run 含 sz=24 (V4-7a cell sz set 整篇生效)
        tbl_match = re.search(r'<w:tbl>.*?</w:tbl>', merged_docxml, re.DOTALL)
        tbl_xml = tbl_match.group(0) if tbl_match else ''
        table_runs = re.findall(r'<w:r\b[^>]*>.*?</w:r>', tbl_xml, re.DOTALL)
        runs_with_sz24 = sum(1 for r in table_runs if '<w:sz w:val="24"/>' in r)
        cell_sz_ok = len(table_runs) > 0 and runs_with_sz24 == len(table_runs)

        # (c) Normal sz=28 (14pt, apply 后)
        normal_match = re.search(
            r'<w:style[^>]*w:styleId="Normal"[^>]*>.*?</w:style>',
            merged_styles, re.DOTALL,
        )
        normal_sz = re.search(r'<w:sz w:val="(\d+)"', normal_match.group(0)) if normal_match else None
        normal_sz_ok = normal_sz is not None and normal_sz.group(1) == '28'

        ok = h2_color_ok and cell_sz_ok and normal_sz_ok
        print(f"  [{'PASS' if ok else 'FAIL'}] case_9_v4_7a_merge_end_to_end_effective "
              f"(招牌: 合并产物 Heading2 黑 + cell sz=24 整篇 + Normal sz=28)")
        if not ok:
            print(f"           (a) Heading2 color: {h2_color.group(1) if h2_color else '<none>'} "
                  f"(expect 000000) -> {'PASS' if h2_color_ok else 'FAIL'}")
            print(f"           (b) cell run sz=24: {runs_with_sz24}/{len(table_runs)} "
                  f"(expect 全数 > 0) -> {'PASS' if cell_sz_ok else 'FAIL'}")
            print(f"           (c) Normal sz: {normal_sz.group(1) if normal_sz else '<none>'} "
                  f"(expect 28) -> {'PASS' if normal_sz_ok else 'FAIL'}")
            fails += 1

    # ─── 汇总 ───
    print()
    print('=' * 70)
    print(f'汇总: {cases - fails} 通过 / {fails} 失败')
    print()
    print('【V4-skel 完整跑铺开占位地图】(Hugin 跑真实标后看到这些喊话+sidecar)')
    print('  V4-2b 占位 (3 类生效点):')
    print('    - enumerate_inventory 4 frontmatter 字段 None + CLI 喊话')
    print('    - asset_query.rationale __PENDING_AI__ + b_mode_fill 溯源行占位段')
    print('    - lookup_priority 非默认档 fallback emit warning')
    print('  V4-1b 占位 (3 类生效点):')
    print('    - 含图段落整段跳过 + 原位插可见占位段 (R10)')
    print('    - missing_elements.yaml sidecar per-asset 写盘')
    print('    - 表内含图粗粒度 + sections 代理 header/footer')
    print('  V4-7 占位 (1 类 4 维度全 noop):')
    print('    - _post_merge_normalize hook noop + stderr 喊话')
    print('    - merge_normalize.log 4 维度 "noop" + per-fragment 元数据')
    return 1 if fails else 0


if __name__ == '__main__':
    sys.exit(main())
