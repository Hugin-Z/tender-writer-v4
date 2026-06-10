# -*- coding: utf-8 -*-
"""make_image_fixture.py · 给 V4-skel C10 e2e 测试生成含 1 张图的 docx fixture。

V4-1b 占位段 / missing_elements.yaml / scan helper 验生效需含图 asset。
demo own_demo 现有 6 个 docx 都是纯文字占位 (make_demo_own_default_assets 生成),
本 generator 单独生成 1 个含图 docx 入 fixture 目录, 跟 demo 隔离。

跑法:
    .venv/Scripts/python.exe scripts/tests/fixtures/v4_skeleton/_generators/make_image_fixture.py
"""
from __future__ import annotations

import base64
import io
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print('[错误] 缺少 python-docx', file=sys.stderr)
    sys.exit(1)

FIXTURE_DIR = Path(__file__).resolve().parent.parent
OUT_PATH = FIXTURE_DIR / 'with_image_asset.docx'

# 最小合法 PNG (1x1 透明), 给 add_picture 用
PNG_MIN_BASE64 = (
    'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNkAAIA'
    'AAoAAv/lxKUAAAAASUVORK5CYII='
)


def main() -> int:
    doc = Document()
    # 段 1: 纯文字
    doc.add_paragraph('一、资质证书')
    # 段 2: 含 1 张图 (V4-1b 占位段触发点)
    img_p = doc.add_paragraph()
    img_p.add_run('图前文字: 等保三级安评报告扫描件').add_picture(
        io.BytesIO(base64.b64decode(PNG_MIN_BASE64)), width=Inches(2)
    )
    # 段 3: 纯文字
    doc.add_paragraph('颁证机构:公安部信息安全等级保护评估中心')
    # 段 4: 含 2 张图 (验 n_drawings 计数 > 1)
    img2 = doc.add_paragraph()
    img2.add_run('图: ').add_picture(
        io.BytesIO(base64.b64decode(PNG_MIN_BASE64)), width=Inches(1)
    )
    img2.add_run(' + ').add_picture(
        io.BytesIO(base64.b64decode(PNG_MIN_BASE64)), width=Inches(1)
    )
    # 段 5: 纯文字 (验非含图段仍 deepcopy)
    doc.add_paragraph('有效期:2025-04 至 2027-04')

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f'[完成] V4-skel C10 含图 fixture: {OUT_PATH}')
    print(f'  含图段数: 2 (段 2 有 1 张图, 段 4 有 2 张图)')
    print(f'  纯文字段数: 3 (段 1/3/5)')
    return 0


if __name__ == '__main__':
    sys.exit(main())
