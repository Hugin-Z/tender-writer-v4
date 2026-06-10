# -*- coding: utf-8 -*-
"""
make_test_docx.py · 生成 check_cross_consistency 测试用 docx fixture

产出 4 个 docx 落到上一级目录:
- clean_response.docx          : 项目名一致 + 简历+架构齐备含具名人员 + 引用全有效
                                 (含同段两人名场景验证 NAME_RE 不跨人)
- dirty_project_name.docx      : docx 项目名与 brief 字面不一致(仿 demo 案例)
- dirty_resume_org.docx        : 同名"张三"在简历是项目经理,在架构图是技术负责人
- dirty_section_ref.docx       : 章节最大到 4.x,但正文+表格 cell 各含 1 处错引用
                                 (8.5 节 / 9.1 节)验证 extract_docx_text 表格覆盖

运行:
    ./run_script.bat tests/fixtures/check_cross_consistency/_generators/make_test_docx.py

产物 git track。
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

OUT_DIR = Path(__file__).resolve().parent.parent

# 模拟 brief.extracted 中的字段(brief_minimal.json 内容与此对齐)
BRIEF_PROJECT_NAME = "测试演示项目 2026 年度技术服务采购"
BRIEF_BUYER_NAME = "测试采购方"
BRIEF_BUYER_AGENCY = "测试代理公司"


def add_heading(doc, level: int, text: str):
    doc.add_heading(text, level=level)


def make_clean(out_path: Path):
    """clean_response.docx: 各项一致,含同段两人名场景验证 NAME_RE 不跨人。"""
    doc = Document()
    # 封面 / 标题(项目名出现在前 500 字符)
    doc.add_heading(BRIEF_PROJECT_NAME, level=0)
    doc.add_paragraph(f"采购人:{BRIEF_BUYER_NAME}")
    doc.add_paragraph(f"采购代理:{BRIEF_BUYER_AGENCY}")

    # 第一章 项目人员(简历章节)
    add_heading(doc, 1, "1 项目人员")
    add_heading(doc, 2, "1.1 团队成员简历")
    # 关键: 同段两人名场景,验证 NAME_RE 不跨人
    doc.add_paragraph(
        "本项目由张三负责整体管理,李四担任技术负责人,王五负责具体实施工作。"
        "张三, 项目经理, 拥有 PMP 认证。"
    )

    # 第二章 团队组成(架构章节,关键词触发 ORG_SECTION_KEYWORDS)
    add_heading(doc, 1, "2 团队组成")
    add_heading(doc, 2, "2.1 组织架构")
    doc.add_paragraph(
        "项目部由项目经理统筹,技术线由技术负责人牵头,实施线由实施工程师执行。"
        "张三担任项目经理,李四担任技术负责人,王五担任实施。"
    )

    # 第三章 实施方案(含有效引用)
    add_heading(doc, 1, "3 实施方案")
    add_heading(doc, 2, "3.1 阶段划分")
    doc.add_paragraph("详见 1.1 节简历章节,具体实施按 3.1 阶段开展。")

    doc.save(str(out_path))


def make_dirty_project_name(out_path: Path):
    """dirty_project_name.docx: 项目名与 brief 字面不一致。"""
    doc = Document()
    # 故意用与 BRIEF_PROJECT_NAME 完全不同的项目名(jaccard < 0.4 触发失败)
    doc.add_heading("某省政务大数据平台 2024 年度运维支持服务", level=0)
    doc.add_paragraph(f"采购人:{BRIEF_BUYER_NAME}")
    add_heading(doc, 1, "1 概述")
    doc.add_paragraph("本项目内容与 brief 完全不一致,用于触发项 4 失败档。")
    doc.save(str(out_path))


def make_dirty_resume_org(out_path: Path):
    """dirty_resume_org.docx: 同名"张三"在简历是项目经理,在架构图是技术负责人。"""
    doc = Document()
    doc.add_heading(BRIEF_PROJECT_NAME, level=0)

    # 简历章节: 张三 担任 项目经理
    add_heading(doc, 1, "1 项目人员")
    add_heading(doc, 2, "1.1 简历")
    doc.add_paragraph("张三担任项目经理, 持有 PMP 认证, 负责整体管理。")

    # 架构章节: 张三 担任 技术负责人(冲突!)
    add_heading(doc, 1, "2 团队组成")
    add_heading(doc, 2, "2.1 组织架构")
    doc.add_paragraph("张三担任技术负责人, 牵头技术方案设计与评审。")

    doc.save(str(out_path))


def make_dirty_section_ref(out_path: Path):
    """dirty_section_ref.docx: 章节最大到 4.x,正文+表格 cell 各含错引用。"""
    doc = Document()
    doc.add_heading(BRIEF_PROJECT_NAME, level=0)

    # 章节体系: 1, 2, 3, 4, 4.1
    add_heading(doc, 1, "1 概述")
    doc.add_paragraph("本项目概述。")
    add_heading(doc, 1, "2 实施方案")
    doc.add_paragraph(
        "技术方案的实施细节,详见 8.5 节,该节描述详细流程。"  # 错引用反例 1(段落)
    )
    add_heading(doc, 1, "3 质量保障")
    doc.add_paragraph("质量保障措施按章节执行。")
    add_heading(doc, 1, "4 服务承诺")
    add_heading(doc, 2, "4.1 响应时间")
    doc.add_paragraph("响应时间 4 小时内。")

    # 表格 cell 错引用反例 2(验证 extract_docx_text 表格覆盖)
    add_heading(doc, 1, "5 附录")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "依据"
    table.cell(1, 0).text = "应急响应"
    table.cell(1, 1).text = "详见 9.1 节标准执行"  # 表格内错引用(必含 REF_RE 引用动词)

    doc.save(str(out_path))


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    targets = {
        "clean_response.docx": make_clean,
        "dirty_project_name.docx": make_dirty_project_name,
        "dirty_resume_org.docx": make_dirty_resume_org,
        "dirty_section_ref.docx": make_dirty_section_ref,
    }
    for name, maker in targets.items():
        path = OUT_DIR / name
        maker(path)
        print(f"[OK] {path} ({path.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
