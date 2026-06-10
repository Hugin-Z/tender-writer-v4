# -*- coding: utf-8 -*-
"""
make_test_docx.py · 生成 compliance_check 测试用 docx fixture

产出 3 个 docx,落到上一级目录:
- clean_response.docx     : 覆盖 scoring_matrix_minimal.csv 全部评分项,
                            Normal 样式 eastAsia=宋体 / ascii=Times New Roman
- missing_keywords.docx   : 与 clean 同结构,但删除"技术方案完整性"段
- font_unsafe.docx        : 与 clean 同结构,但 Normal 样式 eastAsia=MS Mincho
                            (故意触发 check_font_safety 失败)

所有 docx 段落格式:页边距 上下 2.54cm,左右 3.17cm(对齐 check_format 期望)。

运行:
    ./run_script.bat tests/fixtures/compliance_check/_generators/make_test_docx.py

产物 git track,无需每次跑测重新生成。仅 fixture 设计需要更新时手工重跑。
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = Path(__file__).resolve().parent.parent


def set_section_margins(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def set_normal_font(doc, east_asia: str, ascii_font: str = "Times New Roman"):
    """设置 Normal 样式的 rFonts(eastAsia + ascii + hAnsi)。"""
    normal = doc.styles["Normal"]
    rpr = normal.element.get_or_add_rPr()
    # 清掉已有的 w:rFonts(避免追加多份)
    for child in rpr.findall(qn("w:rFonts")):
        rpr.remove(child)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rpr.insert(0, rfonts)


def add_cover_paragraphs(doc):
    """加入封面/目录关键词以满足 check_format 期望。"""
    doc.add_heading("投标文件", level=1)
    doc.add_paragraph("投标人:示例投标方 ALPHA")
    doc.add_paragraph("目录")


def add_management_paragraph(doc):
    doc.add_heading("第一章 项目管理体系", level=1)
    doc.add_paragraph("本投标方在项目管理体系方面采用 PMBOK 与敏捷方法论,"
                      "建立分层管理机制并完全响应招标文件要求。")


def add_tech_paragraph(doc):
    doc.add_heading("第二章 技术方案完整性", level=1)
    doc.add_paragraph("技术方案完整性是核心要素,包括完整架构和模块化设计。"
                      "本部分内容完全满足★实质性响应条款的所有要求。")


def make_clean(out_path: Path):
    doc = Document()
    set_section_margins(doc)
    set_normal_font(doc, east_asia="宋体")
    add_cover_paragraphs(doc)
    add_management_paragraph(doc)
    add_tech_paragraph(doc)
    doc.save(str(out_path))


def make_missing_keywords(out_path: Path):
    doc = Document()
    set_section_margins(doc)
    set_normal_font(doc, east_asia="宋体")
    add_cover_paragraphs(doc)
    add_management_paragraph(doc)
    # 不加技术方案段,故意造成"技术方案完整性"漏答
    doc.save(str(out_path))


def make_font_unsafe(out_path: Path):
    doc = Document()
    set_section_margins(doc)
    set_normal_font(doc, east_asia="MS Mincho")  # 故意非白名单
    add_cover_paragraphs(doc)
    add_management_paragraph(doc)
    add_tech_paragraph(doc)
    doc.save(str(out_path))




def _inject_font_in_fonttable(docx_bytes: bytes, font_name: str) -> bytes:
    """fontTable.xml に font_name エントリを追加して返す。"""
    import io
    import zipfile

    in_buf = io.BytesIO(docx_bytes)
    out_buf = io.BytesIO()
    with zipfile.ZipFile(in_buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/fontTable.xml":
                xml_text = data.decode("utf-8")
                entry = f'<w:font w:name="{font_name}"/>'
                # Insert before closing tag
                xml_text = xml_text.replace("</w:fonts>", f"{entry}</w:fonts>")
                data = xml_text.encode("utf-8")
            zout.writestr(item, data)
    out_buf.seek(0)
    return out_buf.read()


def make_font_unsafe_table(out_path: Path):
    """Normal=宋体（白名单），fontTable に Comic Sans MS（非白名单非 boilerplate）を注入。"""
    import io
    doc = Document()
    set_section_margins(doc)
    set_normal_font(doc, east_asia="宋体")
    add_cover_paragraphs(doc)
    add_management_paragraph(doc)
    add_tech_paragraph(doc)
    buf = io.BytesIO()
    doc.save(buf)
    modified = _inject_font_in_fonttable(buf.getvalue(), "Comic Sans MS")
    out_path.write_bytes(modified)


def make_font_simsun_alias(out_path: Path):
    """Normal=SimSun（宋体の别名），fontTable はデフォルト boilerplate のみ。"""
    doc = Document()
    set_section_margins(doc)
    set_normal_font(doc, east_asia="SimSun")
    add_cover_paragraphs(doc)
    add_management_paragraph(doc)
    add_tech_paragraph(doc)
    doc.save(str(out_path))

def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    targets = {
        "clean_response.docx": make_clean,
        "missing_keywords.docx": make_missing_keywords,
        "font_unsafe.docx": make_font_unsafe,
        "font_unsafe_table.docx": make_font_unsafe_table,
        "font_simsun_alias.docx": make_font_simsun_alias,
    }
    for name, maker in targets.items():
        path = OUT_DIR / name
        maker(path)
        print(f"[OK] {path} ({path.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
