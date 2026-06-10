# -*- coding: utf-8 -*-
"""
make_test_assets.py · 生成 assets_provider 测试用 docx fixture

目录结构(对齐真实 assets/ 的 <类别>/<company_id>/ 顺序):

- 公司资质/company_a/_raw/20240601_iso27001.docx
- 公司资质/company_a/_raw/20250101_iso9001.docx
- 公司资质/company_a/_raw/20251015_business_license.docx
- 公司资质/company_a/iso9001.md(curated 兜底,仅 _raw 全空时使用,本 fixture 实际不会用到)
- 团队简历/company_a/_raw/20250301_zhang_san.docx

运行:
    ./run_script.bat tests/fixtures/assets_provider/_generators/make_test_assets.py
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


# 输出根目录:fixtures/assets_provider/(类别在前 / company_id 在后)
OUT_ROOT = Path(__file__).resolve().parent.parent
COMPANY_ID = "company_a"


def make_docx(path: Path, title: str, body: str):
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph(title)
    doc.add_paragraph(body)
    doc.save(str(path))


def main() -> int:
    # 公司资质 / company_a / _raw
    qual_raw = OUT_ROOT / "公司资质" / COMPANY_ID / "_raw"
    make_docx(
        qual_raw / "20240601_iso27001.docx",
        "ISO 27001 信息安全证书(test)",
        "test asset content: ISO 27001 颁证 2024-06-01",
    )
    make_docx(
        qual_raw / "20250101_iso9001.docx",
        "ISO 9001 质量体系证书(test)",
        "test asset content: ISO 9001 颁证 2025-01-01",
    )
    make_docx(
        qual_raw / "20251015_business_license.docx",
        "营业执照(test)",
        "test asset content: 营业执照 颁证 2025-10-15",
    )

    # 公司资质 / company_a / curated md
    md_path = OUT_ROOT / "公司资质" / COMPANY_ID / "iso9001.md"
    md_path.parent.mkdir(parents=True, exist_ok=True)
    md_path.write_text(
        "# ISO 9001 质量体系认证(curated)\n\n"
        "颁证日期:2025-01-01\n"
        "test 用 curated md 条目\n",
        encoding="utf-8",
    )

    # 团队简历 / company_a / _raw
    cv_raw = OUT_ROOT / "团队简历" / COMPANY_ID / "_raw"
    make_docx(
        cv_raw / "20250301_zhang_san.docx",
        "张三 简历(test)",
        "test asset content: 项目经理 张三 简历",
    )

    print("[OK] 生成完毕 → ", OUT_ROOT)
    return 0


if __name__ == "__main__":
    sys.exit(main())
