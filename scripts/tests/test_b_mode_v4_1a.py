# -*- coding: utf-8 -*-
"""test_b_mode_v4_1a.py · V4-1a B asset 表格 + 段落样式保真测试

覆盖:
- _build_style_id_map(src_doc, dst_doc): name-based styleId 映射
- _rebind_or_strip_refs(elem, style_map, part_name): tblStyle/pStyle rebind +
  numId strip 真行为
- e2e 注入 (绕过 provider, 直接走 deepcopy 循环): 用 anchor 1.docx 模拟
  _handle_asset_lookup 的 element-level 搬运 + rebind/strip, 写盘 + zipfile
  解 document.xml 做 XML 断言

anchor (用户提供, fixture-first):
  assets/公司资质/own_demo/_raw/1.docx (WPS 业绩证明表, 9x7 表格 / dxa 列宽 /
  tblStyle=12 (name="Table Grid") / pStyle=3 (name="heading 2") / numId=2)

XML 断言用确切值, 禁泛断言。列宽断言 7 个 dxa 值逐值对照。

人工目检项: 测试末尾打印明确提示 (不假装自动验视觉), 真生成注入后 docx
供 Hugin 打开核对表格边框 / 列宽视觉 / Heading 样式 / 跨 WPS 渲染。

运行:
    .venv/Scripts/python.exe scripts/tests/test_b_mode_v4_1a.py
"""
from __future__ import annotations

import copy
import re
import sys
import tempfile
import zipfile
from pathlib import Path

REPO = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(REPO / "scripts"))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

from b_mode_fill import (  # noqa: E402
    _apply_explicit_fonts_to_runs,
    _apply_table_cell_size_to_runs,
    _build_style_id_map,
    _normalize_run_size,
    _rebind_or_strip_refs,
)
from docx_builder import apply_default_styles  # noqa: E402

ANCHOR = REPO / "assets" / "公司资质" / "own_demo" / "_raw" / "1.docx"

# 源 1.docx 实测列宽 (Phase 0 dxa 值, 由 V4-1a anchor 实测固定)
EXPECTED_COL_WIDTHS = ["551", "965", "1314", "2268", "993", "992", "1213"]
EXPECTED_ROWS = 9
EXPECTED_COLS = 7


def _inject_anchor_to_tmp_doc() -> Path:
    """模拟 _handle_asset_lookup 的 OXML 注入路径, 落到 tmp docx。

    V4-1a.8: 补 apply_default_styles 跟生产路径 (b_mode_fill._main_body
    L404-405) 对齐。原夹具未调 apply_default_styles, 致 master Heading2
    保留 python-docx 默认 accent1 蓝色 (4F81BD), Hugin 目检产物时误判
    "标题蓝色是源/rebind bug"; 实测生产 demo assembled.docx Heading2
    已是黑色 (apply_default_styles 把 color 覆盖为 000000)。
    """
    src_doc = Document(str(ANCHOR))
    dst_doc = Document()
    apply_default_styles(dst_doc)  # V4-1a.8: 对齐生产路径
    style_map = _build_style_id_map(src_doc, dst_doc)
    src_body = src_doc.element.body
    dst_body = dst_doc.element.body
    dst_sectPr = dst_body.find(qn("w:sectPr"))
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")
    for child in src_body.iterchildren():
        if child.tag == p_tag:
            texts = "".join(t.text or "" for t in child.iter(qn("w:t")))
            if not texts.strip():
                continue
        elif child.tag == tbl_tag:
            pass
        else:
            continue
        new_elem = copy.deepcopy(child)
        _rebind_or_strip_refs(new_elem, style_map, part_name="V4-1a anchor 测试")
        # V4-1a.6 fix: 补显式 rFonts (跟 _handle_asset_lookup 同序: rebind 后字体)
        _apply_explicit_fonts_to_runs(new_elem)
        # V4-1a.9 字号归化: 删 sz/szCs 让字号回落 master Normal
        _normalize_run_size(new_elem)
        # V4-1a.12 表格 cell 字号 per-run 真修 (跟 _handle_asset_lookup 同序)
        _apply_table_cell_size_to_runs(new_elem)
        if dst_sectPr is not None:
            dst_sectPr.addprevious(new_elem)
        else:
            dst_body.append(new_elem)
    tmp = Path(tempfile.mkstemp(suffix=".docx", prefix="v4_1a_test_")[1])
    dst_doc.save(str(tmp))
    return tmp


def _read_document_xml(docx_path: Path) -> str:
    with zipfile.ZipFile(docx_path) as z:
        return z.read("word/document.xml").decode("utf-8")


def _read_styles_xml(docx_path: Path) -> str:
    with zipfile.ZipFile(docx_path) as z:
        return z.read("word/styles.xml").decode("utf-8")


def main() -> int:
    fails = 0
    cases = 0
    print("test_b_mode_v4_1a (V4-1a B asset 表格 + 段落样式保真)")
    print()

    if not ANCHOR.exists():
        print(f"[FAIL] anchor 缺失: {ANCHOR}")
        print("      V4-1a fixture-first 依赖该 anchor (用户提供的真实业绩表)")
        return 1
    print(f"  anchor: {ANCHOR.relative_to(REPO)} ({ANCHOR.stat().st_size:,} bytes)")
    print()

    # ─── case 1: _build_style_id_map name-based 映射 ───
    cases += 1
    src = Document(str(ANCHOR))
    dst = Document()
    style_map = _build_style_id_map(src, dst)
    ok = style_map.get("12") == "TableGrid" and style_map.get("3") == "Heading2"
    print(f"  [{'PASS' if ok else 'FAIL'}] case_1_style_map: "
          f"'12'→{style_map.get('12')!r}, '3'→{style_map.get('3')!r} "
          f"(expect TableGrid / Heading2)")
    if not ok:
        print(f"           style_map size={len(style_map)}, full={style_map}")
        fails += 1

    # ─── e2e 注入 (后续 case 共用产物) ───
    out_path = _inject_anchor_to_tmp_doc()
    doc_xml = _read_document_xml(out_path)

    # ─── case 2: 列宽逐值对照 (V4-1a 招牌断言) ───
    cases += 1
    grid_cols = re.findall(r'<w:gridCol w:w="(\d+)"', doc_xml)
    ok = grid_cols == EXPECTED_COL_WIDTHS
    print(f"  [{'PASS' if ok else 'FAIL'}] case_2_col_widths (招牌目标): "
          f"got {grid_cols}, expect {EXPECTED_COL_WIDTHS}")
    if not ok:
        fails += 1

    # ─── case 3: 表格行列数 ───
    cases += 1
    tbl_count = doc_xml.count("<w:tbl>")
    tr_count = doc_xml.count("<w:tr ") + doc_xml.count("<w:tr>")
    # 每行 <w:tc> 数 = 总 tc / tr
    tc_count = doc_xml.count("<w:tc>") + doc_xml.count("<w:tc ")
    cols_per_row = tc_count // tr_count if tr_count else 0
    ok = (tbl_count >= 1 and tr_count == EXPECTED_ROWS
          and cols_per_row == EXPECTED_COLS)
    print(f"  [{'PASS' if ok else 'FAIL'}] case_3_table_shape: "
          f"tbl={tbl_count}, tr={tr_count} (expect {EXPECTED_ROWS}), "
          f"tc/row={cols_per_row} (expect {EXPECTED_COLS})")
    if not ok:
        fails += 1

    # ─── case 4: tblStyle rebind 到 master TableGrid ───
    cases += 1
    tbl_styles = re.findall(r'<w:tblStyle w:val="([^"]+)"', doc_xml)
    ok = tbl_styles == ["TableGrid"]
    print(f"  [{'PASS' if ok else 'FAIL'}] case_4_tblStyle_rebind: "
          f"got {tbl_styles}, expect ['TableGrid'] (源端 styleId='12' 已 rebind)")
    if not ok:
        fails += 1

    # ─── case 5: pStyle rebind 到 master Heading2 ───
    cases += 1
    p_styles = re.findall(r'<w:pStyle w:val="([^"]+)"', doc_xml)
    ok = "Heading2" in p_styles and "3" not in p_styles
    print(f"  [{'PASS' if ok else 'FAIL'}] case_5_pStyle_rebind: "
          f"got {p_styles}, expect contains 'Heading2' & no '3' "
          f"(源端 pStyle='3' 已 rebind)")
    if not ok:
        fails += 1

    # ─── case 6: numId strip ───
    cases += 1
    num_refs = re.findall(r'<w:numId w:val="(\d+)"', doc_xml)
    ok = num_refs == []
    print(f"  [{'PASS' if ok else 'FAIL'}] case_6_numId_strip: "
          f"got {num_refs}, expect [] (V4-1a 主动降级, master 无 multilevel "
          f"decimal 可复用; 段落语义由 pStyle Heading2 rebind 保留)")
    if not ok:
        fails += 1

    # ─── case 7: pStyle Heading2 段落即承载源端"业绩证明文件"的样式 ───
    # 注: 源 anchor 的 run 级无 <w:b/> 直接标签 (Phase 0 实测 rFonts 只 1 处
    # 且仅 hint="eastAsia"), 加粗由 pStyle Heading 2 样式定义提供。
    # rebind 后 master Heading2 (docx_builder DEFAULT_STYLES) 同样含
    # b=True 15pt 黑体, 故视觉加粗保留 = pStyle rebind 成功的等价证据。
    cases += 1
    ok = "Heading2" in p_styles  # 实质=case 5 重述, 但提示语义不同
    print(f"  [{'PASS' if ok else 'FAIL'}] case_7_bold_via_pStyle: "
          f"Heading2 ∈ pStyle (源端 anchor 加粗依赖 pStyle, 非 run 级 <w:b/>; "
          f"rebind 后 master Heading2 同样 bold=True 15pt)")
    if not ok:
        fails += 1

    # ─── case 8: V4-1a.6 字体 fix · 注入 run 补显式 rFonts 防 fallback ───
    # 验证 anchor 注入后表格内 run 的 rFonts 都补了显式字体名,
    # 不再"只 w:hint 无字体名" (该状态会触发 WPS/Word fallback 到 MS 明朝)。
    cases += 1
    # 表格内所有 rFonts 字面
    tbl_xml_m = re.search(r'<w:tbl>.*?</w:tbl>', doc_xml, re.DOTALL)
    tbl_xml = tbl_xml_m.group(0) if tbl_xml_m else ""
    rfonts_in_tbl = re.findall(r'<w:rFonts[^/]*/>', tbl_xml)
    # 每个 rFonts 实测 attribute (检查 4 字体名是否全在)
    has_eastasia_song = sum(1 for r in rfonts_in_tbl if 'w:eastAsia="宋体"' in r)
    has_ascii_tnr = sum(1 for r in rfonts_in_tbl if 'w:ascii="Times New Roman"' in r)
    has_hansi_tnr = sum(1 for r in rfonts_in_tbl if 'w:hAnsi="Times New Roman"' in r)
    has_cs_tnr = sum(1 for r in rfonts_in_tbl if 'w:cs="Times New Roman"' in r)
    # 仅 hint 无字体名 (fallback 隐患) 的 count = 0
    only_hint_pat = re.compile(
        r'<w:rFonts(?:\s+w:hint="[^"]*")?\s*/>'
    )
    only_hint_no_font = sum(
        1 for r in rfonts_in_tbl
        if only_hint_pat.fullmatch(r)
    )
    n = len(rfonts_in_tbl)
    ok = (
        n > 0
        and has_eastasia_song == n
        and has_ascii_tnr == n
        and has_hansi_tnr == n
        and has_cs_tnr == n
        and only_hint_no_font == 0
    )
    print(f"  [{'PASS' if ok else 'FAIL'}] case_8_explicit_fonts (V4-1a.6 字体 fix): "
          f"表格 rFonts {n} 个 / eastAsia=宋体 {has_eastasia_song}/{n}, "
          f"ascii=Times New Roman {has_ascii_tnr}/{n}, "
          f"hAnsi=Times New Roman {has_hansi_tnr}/{n}, "
          f"cs=Times New Roman {has_cs_tnr}/{n} / "
          f"仅 hint 无字体名 (fallback 隐患) {only_hint_no_font} (expect 0)")
    if not ok:
        print(f"           sample rFonts: {sorted(set(rfonts_in_tbl))[:3]}")
        fails += 1

    # ─── case 9: V4-1a.9 字号归化 + 边界守正 (字体/段距/缩进不误剥) ───
    # 验证 anchor 注入后 sz/szCs 已剥 (表外段落回落 master Normal) + 保留项
    # 均未被误剥。V4-1a.12 起表内 cell run 由 _apply_table_cell_size_to_runs
    # set sz/szCs=24 (12pt), 全文 sz count 不再为 0; case_9 改验"表外 sz=0"
    # 才是 V4-1a.9 真语义 (归化只对表外段落生效, 表内 cell 走 V4-1a.12 路径)。
    cases += 1
    # 把表格内容从 doc_xml 抠出, 剩余 = 表外内容
    outside_xml = re.sub(r'<w:tbl>.*?</w:tbl>', '', doc_xml, flags=re.DOTALL)
    sz_outside = len(re.findall(r'<w:sz w:val=', outside_xml))
    szCs_outside = len(re.findall(r'<w:szCs w:val=', outside_xml))
    rfonts_eastasia = len(re.findall(r'w:eastAsia="宋体"', doc_xml))
    rfonts_only_hint = len(re.compile(
        r'<w:rFonts(?:\s+w:hint="[^"]*")?\s*/>'
    ).findall(doc_xml))
    spacing_count = len(re.findall(r'<w:spacing\s+[^/]+/>', doc_xml))
    ind_count = len(re.findall(r'<w:ind\s+[^/]+/>', doc_xml))
    # V4-1a.9 前 anchor 实测: sz=103 + szCs=103; spacing=64 + ind=69; rFonts 全补
    ok = (
        sz_outside == 0           # 归化: 表外 sz 已剥 (回落 Normal 14pt)
        and szCs_outside == 0     # 归化: 表外 szCs 已剥
        and rfonts_eastasia > 0   # 守正: 字体 (V4-1a.6) 未被误剥
        and rfonts_only_hint == 0  # 守正: 0 个仅 hint (MS 明朝隐患仍清零)
        and spacing_count > 0     # 守正: 段距保留 (中文排版裁定窄)
        and ind_count > 0         # 守正: 缩进保留
    )
    print(f"  [{'PASS' if ok else 'FAIL'}] case_9_size_normalize_and_boundary "
          f"(V4-1a.9 表外字号归化 + 守正)")
    print(f"           归化: 表外 sz={sz_outside} szCs={szCs_outside} (expect 0/0, "
          f"V4-1a.12 表内 cell sz=24 由 case_10 验)")
    print(f"           守正: rFonts eastAsia=宋体 {rfonts_eastasia} (expect >0), "
          f"仅 hint {rfonts_only_hint} (expect 0)")
    print(f"           守正: spacing={spacing_count} ind={ind_count} (expect >0/>0, "
          f"段距缩进裁定窄保留)")
    if not ok:
        fails += 1

    # ─── case 10: V4-1a.12 表格 cell 字号 per-run 真修 (验生效层非存在层) ───
    # 历史: V4-1a.10/.11 原 case_10 验 "TableNormal style 内 sz=24 写进去"
    # (验存在), 经 Hugin 目检 WPS 显示 14pt 暴露走错 OXML 字号继承层级 (段落
    # style Normal sz=28 压过表格 style TableNormal sz=24, TableNormal sz 对
    # cell 文字无效)。V4-1a.12 改 per-run set 给 cell run, run-level rPr.sz
    # 优先级最高, XML 有则 WPS 必取 -- 本 case_10 改写为验生效层:
    # (a) 注入表格每个 cell run rPr 含 <w:sz w:val="24"/> + <w:szCs w:val="24"/>
    # (b) master TableNormal style 无 sz (V4-1a.10 死代码已回滚)
    # run-level rPr.sz 是 OXML 字号继承优先级最高一层, XML 断言验"存在"即等价
    # 验"WPS 渲染必为 12pt", 不再有"写存在不验生效" R10 风险。
    cases += 1
    from docx_builder import DEFAULT_TABLE_SIZE_PT  # noqa: E402
    styles_xml = _read_styles_xml(out_path)
    expected_sz = str(DEFAULT_TABLE_SIZE_PT * 2)  # 半 point

    # (a) 注入表格每个 cell run rPr 含 sz=24 + szCs=24 (per-run set 生效)
    tbl_xml_m = re.search(r'<w:tbl>.*?</w:tbl>', doc_xml, re.DOTALL)
    tbl_xml = tbl_xml_m.group(0) if tbl_xml_m else ''
    table_runs = re.findall(r'<w:r\b[^>]*>.*?</w:r>', tbl_xml, re.DOTALL)
    sz_pattern = '<w:sz w:val="' + expected_sz + '"/>'
    szCs_pattern = '<w:szCs w:val="' + expected_sz + '"/>'
    runs_with_sz = sum(1 for r in table_runs if sz_pattern in r)
    runs_with_szCs = sum(1 for r in table_runs if szCs_pattern in r)

    # (b) V4-1a.10 死代码回滚验证 (TableNormal style 无 sz)
    tn_match = re.search(
        r'<w:style[^>]*w:styleId="TableNormal"[^>]*>.*?</w:style>',
        styles_xml, re.DOTALL,
    )
    tn_sz_in_style = re.search(r'<w:sz w:val=', tn_match.group(0)) if tn_match else None

    ok = (
        len(table_runs) > 0
        and runs_with_sz == len(table_runs)
        and runs_with_szCs == len(table_runs)
        and tn_sz_in_style is None  # TableNormal sz 已回滚 (无 sz)
    )
    print(f"  [{'PASS' if ok else 'FAIL'}] case_10_cell_size_per_run_set "
          f"(V4-1a.12 cell run sz/szCs 直接 set {DEFAULT_TABLE_SIZE_PT}pt 生效)")
    print(f"           生效层: cell run 含 sz={expected_sz}: "
          f"{runs_with_sz}/{len(table_runs)} (expect 全数)")
    print(f"           生效层: cell run 含 szCs={expected_sz}: "
          f"{runs_with_szCs}/{len(table_runs)} (expect 全数)")
    print(f"           回滚: TableNormal style 内 <w:sz>: "
          f"{'absent' if tn_sz_in_style is None else 'PRESENT (V4-1a.10 死代码未删)'}")
    if not ok:
        fails += 1

    # ─── 汇总 + 人工目检 surface ───
    print()
    print("=" * 70)
    print(f"汇总: {cases - fails} 通过 / {fails} 失败")
    print()
    print("【人工目检项 surface (不在自动断言范围)】")
    print(f"  注入产物 (供打开核对): {out_path}")
    print("  请用 WPS / Word 打开, 视觉核对以下项:")
    print("    1. 9x7 表格边框是否完整 (single 单线 0.5pt 对齐源 1.docx)")
    print("    2. 7 列列宽视觉是否跟源 1.docx 一致 (vs 自动验 dxa 值已 PASS)")
    print("    3. 标题段 '业绩证明文件' 是否走 Heading 2 样式 (加粗 15pt 黑体)")
    print("    4. 跨 WPS / Word 不同版本渲染无字体 fallback (CJK 主题字体灾难类型)")
    print("    5. 单元格内容无 OXML 错位 (multi-p / 中文换行正确)")
    print("=" * 70)
    return 1 if fails else 0


if __name__ == "__main__":
    sys.exit(main())
