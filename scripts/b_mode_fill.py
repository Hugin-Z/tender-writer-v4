# -*- coding: utf-8 -*-
"""
b_mode_fill.py · B 模式 Part 组装(V61)

职责:读 manifest.yaml 和 assets_provider 字段,按 assembly_order 逐段处理,
产出 assembled.docx + .pending_marker 空文件。

source_type 分流:
- inline_template: 占位文字 + 引用 c_mode 填充能力(本轮不实际调用,产占位段落说明)
- asset_lookup: 调 AssetsProvider.lookup + resolve → 从返回的占位 docx 拷贝段落
- self_drafted: 产占位段落 "[本节需供应商自撰: <section_title>]"
- 未知 source_type: raise ValueError

用法:
    python scripts/b_mode_fill.py --project <项目> --part <N>
"""

from __future__ import annotations

import argparse
import copy
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError:
    print("[错误] 缺少 python-docx 依赖,请先双击 install.bat 安装依赖。", file=sys.stderr)
    sys.exit(1)

try:
    import yaml
except ImportError:
    print("[错误] 缺少 PyYAML 依赖。", file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent


def _safe_name(name: str) -> str:
    s = re.sub(r'^[一二三四五六七八九十\d]+[、.]\s*', '', name)
    s = re.sub(r'[（(].+?[)）]', '', s)
    s = s.strip()
    return s or 'part'


def _append_heading(doc, text: str, level: int = 2):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def _append_placeholder_body(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True


def _handle_inline_template(doc, spec: dict, part_name: str):
    """inline_template: 招标文件给模板(基本情况表/信誉声明等带变量占位的格式)。

    行为: 产占位段, 由使用者投标前手工按 items 字段清单填写。同 self_drafted
    定位 — 工具不自动化, 投标方人填。

    历史: V3-1 期 docstring 提"未来由 b_mode_fill 自动调 c_mode_fill 完成
    填充 (留 V4)" — 但 V4 全期无真实需求驱动 (低频, 投标方人填 2 分钟内
    搞定, 跨模块自动化 ROI 低), 决策不做。V4 收口后 (V4-7a / v4_backlog
    ee596be 期间) 顺手诚实化删 "留 V4" 承诺字面。
    """
    title = spec.get('section_title', '<untitled>')
    _append_heading(doc, f"{spec.get('section_id', '')} {title}")
    _append_placeholder_body(
        doc,
        f"[inline_template 占位] 本段为招标文件给定模板的变量填充。"
        f"产出 assembled.docx 后请使用者手工按 items 字段清单填写。"
    )
    items = spec.get('items') or []
    if items:
        _append_placeholder_body(doc, f"字段清单: {', '.join(str(i) for i in items)}")


def _build_style_id_map(src_doc, dst_doc) -> dict[str, str]:
    """V4-1a S2: 建 {源 styleId → master styleId} 映射, 按 w:name 对齐。

    源 docx (e.g. WPS 导出) 用数字 styleId (如 "12"), master (python-docx
    Document) 用 Word 标准 styleId (如 "TableGrid"), 但 w:name 一致
    ("Table Grid")。本函数按 w:name 做 name-based rebind 映射。

    返 dict {src_styleId: master_styleId}; 源端 styleId 在 master 无同 w:name
    则该 styleId 不入 dict (caller _rebind_or_strip_refs 走 strip + warning
    路径)。
    """
    src_styles = src_doc.styles.element
    dst_styles = dst_doc.styles.element
    src_id_to_name: dict[str, str] = {}
    for s in src_styles.findall(qn('w:style')):
        sid = s.get(qn('w:styleId'))
        nm = s.find(qn('w:name'))
        if sid and nm is not None:
            name_val = nm.get(qn('w:val'))
            if name_val:
                src_id_to_name[sid] = name_val
    name_to_dst_id: dict[str, str] = {}
    for s in dst_styles.findall(qn('w:style')):
        sid = s.get(qn('w:styleId'))
        nm = s.find(qn('w:name'))
        if sid and nm is not None:
            name_val = nm.get(qn('w:val'))
            if name_val:
                name_to_dst_id[name_val] = sid
    return {sid: name_to_dst_id[name]
            for sid, name in src_id_to_name.items()
            if name in name_to_dst_id}


def _rebind_or_strip_refs(elem, style_map: dict[str, str],
                          part_name: str = '') -> tuple[int, int, int]:
    """V4-1a S2: 对 deepcopy 后的 element 做跨 part ref 处理。

    - <w:tblStyle w:val=X> / <w:pStyle w:val=X>:
      * X 在 style_map 内 → 改 X 为 map[X] (name-based rebind, 零视觉损失)
      * 否则 → strip 该 ref + emit warning (回退 master 默认样式)
    - <w:numPr> (含 numId / ilvl): strip 整 numPr + emit warning
      (V4-1a 主动降级, master 端无 multilevel decimal 可复用; 段落字面
      无编号前缀, 标题语义由 pStyle rebind 保留; 详见 plans/v4-1a.md)

    rebind/strip 对 elem (deepcopy 副本) 操作, 不动源 docx。

    返 (rebind_count, strip_style_count, strip_num_count) 供统计。
    """
    rebind_n = 0
    strip_style_n = 0
    strip_num_n = 0
    for tag_qn in (qn('w:tblStyle'), qn('w:pStyle')):
        for ref in list(elem.iter(tag_qn)):
            val = ref.get(qn('w:val'))
            if val is None:
                continue
            if val in style_map:
                ref.set(qn('w:val'), style_map[val])
                rebind_n += 1
            else:
                ref.getparent().remove(ref)
                strip_style_n += 1
                print(f"[警告] V4-1a B asset 注入 [{part_name}]: styleId={val!r} "
                      f"在 master 端无同名样式, strip 后回退 master 默认",
                      file=sys.stderr)
    for numPr in list(elem.iter(qn('w:numPr'))):
        numPr.getparent().remove(numPr)
        strip_num_n += 1
        print(f"[警告] V4-1a B asset 注入 [{part_name}]: numPr (自动编号前缀) "
              f"strip, 如需可手补编号字面", file=sys.stderr)
    return rebind_n, strip_style_n, strip_num_n


def _apply_explicit_fonts_to_runs(elem) -> int:
    """V4-1a.6 fix: 给搬入 element 的每个 <w:r> 补显式 <w:rFonts> 字体名,
    防 deepcopy 注入后中文 fallback 到 MS 明朝 (L4 灾难复现) + 西文 fallback
    到 theme minorLatin。

    源 docx (e.g. WPS 导出 1.docx) run 的 rFonts 通常仅 w:hint="eastAsia"
    无显式字体名, 靠 theme minorEastAsia / minorLatin 解字体; 注入 master 后
    master theme ea="" → WPS/Word fallback 到系统 CJK 默认 (MS 明朝, 日文)。

    本函数补显式 rFonts (跟 master apply_default_styles 给 Normal 设的字体一致):
    - w:eastAsia → 宋体 (master 中文)
    - w:ascii    → Times New Roman (master 西文 ASCII)
    - w:hAnsi    → Times New Roman (master 西文高位 ANSI)
    - w:cs       → Times New Roman (复杂脚本, 跟 set_run_font 同型)

    设计要点: **仅在对应属性不存在时 set** (不覆盖源 docx 已有显式字体)。
    源若某 run 显式指定了字体名 (非靠主题), 保留源的; 只补"靠主题"的 run。
    本 bug 修的是"靠主题导致 fallback", 不是"强制统一所有字体"。

    返回补了多少个 attribute (供日志/测试统计)。
    """
    fonts = (
        (qn("w:eastAsia"), "宋体"),
        (qn("w:ascii"),    "Times New Roman"),
        (qn("w:hAnsi"),    "Times New Roman"),
        (qn("w:cs"),       "Times New Roman"),
    )
    n = 0
    for r in elem.iter(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            r.insert(0, rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr_qn, font_val in fonts:
            if rfonts.get(attr_qn) is None:
                rfonts.set(attr_qn, font_val)
                n += 1
    return n


def _normalize_run_size(elem) -> int:
    """V4-1a.9 字号归化: 删搬入 element 内所有 <w:rPr> 下的 <w:sz> 和 <w:szCs>
    子元素, 让字号回落 master Normal style (14pt) 实现字号统一。

    源 docx run 通常带显式字号 (e.g. 1.docx 全 sz=24=12pt, szCs=21=10.5pt),
    跟 master 不一致导致注入后表格字号飘。删 sz/szCs 让字号回落 style 系统,
    显示 master Normal 字号。

    设计要点 (跟 V4-1a.6 字体 fix "显式 set" 互补):
    - 字体怕 fallback (主题解不到 → MS 明朝灾难), 所以 V4-1a.6 显式 set 字体名
    - 字号有 master style 兜底 (style 系统 cascade 到 Normal 14pt),
      所以这里"删标签回落" 更干净, 不假装"已对齐"

    边界 (绝不误伤):
    - 不碰 <w:rFonts> (V4-1a.6 补的显式字体, 删了 MS 明朝 fallback 回来)
    - 不碰 <w:b> / <w:i> / <w:u> (强调, 尺度乙保留)
    - 不碰 pPr 的 <w:spacing> / <w:ind> (段距/缩进, 中文排版裁定窄保留)
    - 不碰表格 tblGrid 列宽

    遍历位置: elem.iter(<w:rPr>) → 找 rPr 直接子的 <w:sz> 和 <w:szCs> 删之。
    会扫到 run 级 rPr (在 <w:r> 内) + 段落标记字符的 rPr (在 <w:pPr> 内) +
    style def 内 rPr (本函数只处理搬入 element 副本, style def 不在范围)。

    返回剥了多少 sz/szCs 标签 (供日志/测试统计)。
    """
    n = 0
    sz_qn = qn("w:sz")
    szCs_qn = qn("w:szCs")
    for rpr in elem.iter(qn("w:rPr")):
        for tag in (sz_qn, szCs_qn):
            for el in rpr.findall(tag):
                rpr.remove(el)
                n += 1
    return n


def _paragraph_has_drawing(p_elem) -> bool:
    """V4-skel.4 V4-1b 占位: 只读检测段落是否含 <w:drawing> (不修改 element)。

    检测用 element 树 find, 仅扫 descendants 不操作 element。跟 V4-1b 实现层
    _extract_inline_images (将做 OXML 拷贝 + media part 复制 + rId 重分配) 完全
    不重叠 — 本 helper 是结构层只读 R10 监测, 让 missing_elements.yaml 知道
    本 Part 有几个含图段落; 实线图片保真留 V4-1b 实现层。
    """
    return p_elem.find(f'.//{qn("w:drawing")}') is not None


def _count_drawings_in(p_elem) -> int:
    """V4-skel.4 V4-1b 占位: 只读统计段落内 <w:drawing> 数 (不修改)。

    跟 _paragraph_has_drawing 同型纯只读。len(list(iter)) 是常见 lxml 模式,
    无 element 重组 / 无 media 操作。
    """
    return len(list(p_elem.iter(qn('w:drawing'))))


def _insert_image_placeholder_paragraph(doc, n_drawings: int, source_path):
    """V4-skel.5 V4-1b 占位 (R10 最关键): 含图段落整段跳过, 在原位置插入显式
    占位段, 替代被跳过的段落。

    业务语义 (Hugin 业务现实): 资质/证书章节内, 含图段是证书扫描件 — 投标前
    由用户人工放入原件 (V4-4 公章红线同源: 签章/原件是用户法律动作, 工具不
    代劳)。占位段聚焦 "缺证书图, 投标前手放原件", 不强调 "丢了文字"
    (含图段文字基本是图注, 真要文字 AI 自己组织, 强调字数会误导)。

    不操作 <w:drawing> 剥离 / 不操作段落重组 / 不操作 media part —
    踏一步就是 V4-1a 第二季入口。整段跳过 + 占位段是结构层 R10 红线。
    """
    text = (
        f'[V4-1b 占位:此处缺 {n_drawings} 张证书图(源 {source_path}),'
        f'V4-1b 未自动搬运,投标前请人工放入原件]'
    )
    p = doc.add_paragraph()  # python-docx add_paragraph 自动 insert sectPr 前
    run = p.add_run(text)
    run.italic = True


def _scan_missing_elements(src_doc, source_asset_path: Path) -> dict:
    """V4-skel.4 V4-1b 占位: 纯只读扫源 docx, 返回 missing_elements 字典。

    扫描点 (全只读, 不剥 element / 不操作 media):
    - 直接子 <w:p> 内含 <w:drawing> 计数 (drawings + 含图段数)
    - 直接子 <w:tbl> 内 cell <w:p> 含 <w:drawing> 粗粒度计数 (cell-level 不展开)
    - len(src_doc.sections) 代理 header/footer 存在性 (V4-1b 实现层精化)

    返回 dict 结构对齐 missing_elements.yaml schema:
        source_asset / images / tables_with_potential_images /
        headers / footers / placeholders_inserted /
        v4_1b_implementation_pending
    """
    src_body = src_doc.element.body
    p_tag = qn('w:p')
    tbl_tag = qn('w:tbl')

    total_drawings = 0
    image_paragraphs = 0
    tables_with_images = 0
    for child in src_body.iterchildren():
        if child.tag == p_tag:
            if _paragraph_has_drawing(child):
                image_paragraphs += 1
                total_drawings += _count_drawings_in(child)
        elif child.tag == tbl_tag:
            if any(_paragraph_has_drawing(p) for p in child.iter(p_tag)):
                tables_with_images += 1

    n_sections = len(src_doc.sections)
    return {
        'source_asset': str(source_asset_path),
        'images': {
            'total_drawings_skipped': total_drawings,
            'image_paragraphs_skipped': image_paragraphs,
        },
        'tables_with_potential_images': tables_with_images,
        'headers': n_sections,  # 粗粒度: section 数代理 header/footer 存在性
        'footers': n_sections,
        'placeholders_inserted': image_paragraphs,  # V4-skel.5 每个含图段都插一个占位段
        'v4_1b_implementation_pending': True,
    }


def _apply_table_cell_size_to_runs(elem) -> int:
    """V4-1a.12 表格 cell 字号 per-run 真修: 给搬入 element 内 <w:tbl> 范围下
    所有 <w:r> 的 rPr 显式 set <w:sz> + <w:szCs> = DEFAULT_TABLE_SIZE_PT (12pt),
    确保注入表格 cell 文字在 WPS/Word 渲染为约定的 12pt。

    根因背景 (Phase 0 实测):
    - V4-1a.10 给 master TableNormal style 加 sz=24 期望"兜底"全表格字号,
      实测 OXML 字号继承优先级: 段落 style Normal sz=28 (14pt) 压过表格 style
      TableNormal sz=24, TableNormal sz 永远被 Normal 覆盖, 对 cell 文字无效。
    - add_table 生产表格 12pt 实际靠 per-run set_run_font(size_pt=12) 生效
      (run-level rPr.sz 优先级最高, 直接生效, 不依赖继承链)。
    - 注入表格 cell 段落无 pStyle → 隐式 Normal → run 无 sz (V4-1a.9 已剥)
      → 取 Normal sz=28 = 14pt (错)。
    修法: 给注入 cell run 跟 add_table 同款 per-run set 12pt, run-level rPr.sz
    优先级最高, XML 有则 WPS 必取, 不赌继承链。

    设计要点:
    - **只对 <w:tbl> 范围内 run set**, 不动表外段落 run (表外字号继续走
      V4-1a.9 剥后回落 Normal 14pt 不动)
    - **set 等价 add** (前空): V4-1a.9 已剥 cell run 的 sz + szCs 两者
      (实测 _normalize_run_size L228-234 同剥), 所以这里 set 实际是从空到 12pt,
      非"覆盖式"语义 — 但即使将来 V4-1a.9 改成只剥 sz 不剥 szCs, 这里覆盖式
      set 也正确 (字号约定就是要统一 12pt, 不容源端 szCs 残留)
    - 跟 V4-1a.6 字体 fix 同型 (run-level 显式 set 不赌继承, 字号约定确定性
      生效)

    返回 set 的 attribute count (供日志/测试统计)。
    """
    # 延迟 import 跟 _main_body 内 apply_default_styles 调用同型,
    # 避免 b_mode_fill 顶层依赖 docx_builder 形成耦合。
    from docx_builder import DEFAULT_TABLE_SIZE_PT
    val_str = str(DEFAULT_TABLE_SIZE_PT * 2)  # 半 point
    n = 0
    for tbl in elem.iter(qn("w:tbl")):
        for r in tbl.iter(qn("w:r")):
            rpr = r.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                r.insert(0, rpr)
            for tag_prefix in ("w:sz", "w:szCs"):
                existing = rpr.find(qn(tag_prefix))
                if existing is not None:
                    rpr.remove(existing)
                sz_el = OxmlElement(tag_prefix)  # OxmlElement 需 prefix string
                sz_el.set(qn("w:val"), val_str)
                rpr.append(sz_el)
                n += 1
    return n


def _handle_asset_lookup(doc, spec: dict, provider, part_name: str,
                         missing_acc: list | None = None):
    """asset_lookup: 走 AssetsProvider 接口。

    V3-1:把 lookup_priority / year_filter / asset_query.type 从 spec 传给
    provider.lookup;真实命中时追加来源溯源行(便于审核位回查),不再追加占位文字。
    """
    from docx import Document as _DocxDocument

    title = spec.get('section_title', '<untitled>')
    asset_type = spec.get('asset_type', 'unknown')
    _append_heading(doc, f"{spec.get('section_id', '')} {title}")

    asset_query = spec.get('asset_query', {}) or {}
    ref = provider.lookup(
        asset_type,
        asset_query_type=asset_query.get('type', ''),
        lookup_priority=spec.get('lookup_priority', 'latest_year_first'),
        year_filter=spec.get('year_filter'),
        # 旧字段保留(PlaceholderAssetsProvider 用于构造 lookup_key)
        part=part_name,
        section_id=spec.get('section_id', ''),
        source=spec.get('source', ''),
    )
    resolved_path = provider.resolve(ref)

    # V4-1a: 从源 docx OXML element 级 deepcopy 整段/整表搬入 (替代 V3-1 仅
    # add_run(para.text) 文本搬运)。表格 (含 tblGrid 列宽) / 段落级样式 /
    # run 级字体加粗 等 element 自带,deepcopy 即得保真。
    # 跨 part ref (tblStyle / pStyle / numId) 由 S2 的 _rebind_refs 处理。
    src_doc = _DocxDocument(str(resolved_path))
    src_body = src_doc.element.body
    # V4-skel.4 V4-1b: 只读扫源 docx 含图段 / 表内含图 / sections 数, 累加给
    # _main_body 末尾统一写 missing_elements.yaml。R10: stderr emit 占位喊话。
    # 注: C4 仅监测 + 喊话, iter 循环不动 (含图段落仍走 V4-1a deepcopy, 带
    # drawing 破引用进 master); C5 改 iter 含图 continue + 插占位段后才修
    # 破引用。
    if missing_acc is not None:
        scan_result = _scan_missing_elements(src_doc, resolved_path)
        missing_acc.append(scan_result)
        print(
            f"[V4-1b 占位] 注入 {resolved_path.name}: 跳过 "
            f"{scan_result['images']['total_drawings_skipped']} 张图段落 / "
            f"{scan_result['tables_with_potential_images']} 个表格可能含图 / "
            f"{scan_result['headers']} 个 section header/footer, "
            f"见 missing_elements.yaml",
            file=sys.stderr,
        )
    # V4-1a S2: 建源 → master 的 name-based styleId 映射 (per-asset 一次)
    style_map = _build_style_id_map(src_doc, doc)
    # 目标 body 末尾通常是 <w:sectPr> (section properties), 新 element 要
    # insert 在它之前; 若无 sectPr 则直接 append。同 python-docx
    # add_paragraph() 内部 _insert_p 策略。
    dst_body = doc.element.body
    dst_sectPr = dst_body.find(qn('w:sectPr'))
    p_tag = qn('w:p')
    tbl_tag = qn('w:tbl')
    rebind_total = 0
    strip_style_total = 0
    strip_num_total = 0
    font_apply_total = 0
    size_strip_total = 0
    cell_size_set_total = 0
    for child in src_body.iterchildren():
        if child.tag == p_tag:
            # 跳过纯空段落 (沿 V3-1 trim 行为, 避免 DEMO PLACEHOLDER 类
            # 尾部空段污染 assembled.docx)
            texts = ''.join(t.text or '' for t in child.iter(qn('w:t')))
            if not texts.strip():
                continue
            # V4-skel.5 V4-1b: 含图段落整段跳过 + 插显式占位段
            # (R10 最关键: 不 deepcopy 避免带 drawing 破引用进 master,
            # 用占位段告诉用户"此处缺证书图, 投标前人工放原件")
            if _paragraph_has_drawing(child):
                n_drawings = _count_drawings_in(child)
                _insert_image_placeholder_paragraph(doc, n_drawings, resolved_path)
                continue
        elif child.tag == tbl_tag:
            # 表格无条件保留 (V4-1a 的核心保真目标)
            pass
        else:
            # <w:sectPr> 等其他 element (section properties) 不搬,
            # 目标 doc 用自己的 sectPr
            continue
        new_elem = copy.deepcopy(child)
        # V4-1a S2: deepcopy 副本上做跨 part ref 处理 (rebind/strip)
        r, ss, sn = _rebind_or_strip_refs(new_elem, style_map, part_name)
        rebind_total += r
        strip_style_total += ss
        strip_num_total += sn
        # V4-1a.6 fix: 补显式 rFonts 防 WPS/Word fallback (中文 MS 明朝 / 西文 Cambria)
        font_apply_total += _apply_explicit_fonts_to_runs(new_elem)
        # V4-1a.9 字号归化: 删 sz/szCs 让字号回落 master Normal (统一注入内容字号)
        size_strip_total += _normalize_run_size(new_elem)
        # V4-1a.12 表格 cell 字号真修: 给 <w:tbl> 内 run 显式 set 12pt
        # (run-level 优先级最高, 不赌继承链; V4-1a.10 的 TableNormal style sz
        # 被 Normal style 覆盖无效)
        cell_size_set_total += _apply_table_cell_size_to_runs(new_elem)
        if dst_sectPr is not None:
            dst_sectPr.addprevious(new_elem)
        else:
            dst_body.append(new_elem)
    if (rebind_total or strip_style_total or strip_num_total
            or font_apply_total or size_strip_total or cell_size_set_total):
        print(f"[信息] V4-1a [{part_name}] ref 处理: rebind={rebind_total}, "
              f"strip_style={strip_style_total}, strip_num={strip_num_total}, "
              f"font_apply={font_apply_total}, size_strip={size_strip_total}, "
              f"cell_size_set={cell_size_set_total}",
              file=sys.stderr)

    if ref.is_placeholder:
        _append_placeholder_body(
            doc,
            f"(AssetRef 占位: is_placeholder={ref.is_placeholder}, "
            f"lookup_key={ref.lookup_key})"
        )
    else:
        # V3-1:真实命中时追加溯源行(便于审核位回查实际选用的 asset)
        src_filename = ref.metadata.get("filename", "")
        src_kind = ref.metadata.get("kind", "")
        if src_filename:
            _append_placeholder_body(
                doc, f"(asset 来源: {src_filename}, kind={src_kind})"
            )
        # V4-skel.2 V4-2b 占位: 检查 asset_query.rationale, 未填则溯源行后
        # 追加占位喊话 (R10 红线: 字段缺失/__PENDING_AI__ 不静默放过)
        rationale = asset_query.get('rationale')
        if rationale is None or rationale == '__PENDING_AI__':
            _append_placeholder_body(
                doc,
                "[V4-2b 占位] rationale 未填 "
                "(asset_query.rationale 缺失或 __PENDING_AI__), "
                "实线 V4-2b 实现层补"
            )


def _handle_self_drafted(doc, spec: dict, part_name: str):
    """self_drafted: 供应商自撰,产占位段落。"""
    title = spec.get('section_title', part_name)
    _append_heading(doc, f"{spec.get('section_id', '')} {title}".strip())
    _append_placeholder_body(
        doc,
        f"[本节需供应商自撰: {title}]"
    )
    if spec.get('source'):
        _append_placeholder_body(doc, f"招标文件原文提示: {spec.get('source')}")


def _dispatch(doc, spec: dict, provider, part_name: str,
              missing_acc: list | None = None):
    st = spec.get('source_type')
    if st == 'inline_template':
        _handle_inline_template(doc, spec, part_name)
    elif st == 'asset_lookup':
        _handle_asset_lookup(doc, spec, provider, part_name,
                             missing_acc=missing_acc)
    elif st == 'self_drafted':
        _handle_self_drafted(doc, spec, part_name)
    else:
        raise ValueError(
            f"未知 source_type='{st}' 在 assembly_order 项 section_id="
            f"{spec.get('section_id', '?')}。"
            f"当前合法值: inline_template / asset_lookup / self_drafted。"
            f"AI 判出新形态由用户 review 决定是否扩展 b_mode_fill。"
        )


def main():
    parser = argparse.ArgumentParser(
        description='B 模式 Part 组装(V61)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument('--project', required=True, help='项目名')
    parser.add_argument('--part', type=int, required=True, help='response_file_parts 索引')
    parser.add_argument(
        '--non-interactive',
        action='store_true',
        help='V3-1:CuratedLocalAssetsProvider 多命中时不弹 stdin 提示,'
             '按 lookup_priority 自动选(适用于 CI / 自动化场景)',
    )
    args = parser.parse_args()

    project_dir = ROOT / 'projects' / args.project

    # V3-3: timing hook
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from _timing_hook import stage_timer

    with stage_timer("b_mode_fill", project_dir):
        _main_body(args, project_dir)


def _main_body(args, project_dir):
    # V53: 未 review 则硬失败
    from brief_schema import ensure_reviewed
    ensure_reviewed(project_dir / 'output')

    import json as _json
    with open(project_dir / 'output' / 'tender_brief.json', 'r', encoding='utf-8') as f:
        brief = _json.load(f)
    parts = brief.get('response_file_parts', [])
    if args.part < 0 or args.part >= len(parts):
        print(f"[错误] --part 超出范围", file=sys.stderr)
        sys.exit(1)
    part = parts[args.part]
    if part.get('production_mode') != 'B':
        print(f"[错误] Part[{args.part}] production_mode="
              f"{part.get('production_mode')} 不是 B", file=sys.stderr)
        sys.exit(1)

    part_dir_name = _safe_name(part['name'])
    b_dir = project_dir / 'output' / 'b_mode' / part_dir_name

    manifest_path = b_dir / 'manifest.yaml'
    if not manifest_path.exists():
        print(f"[错误] 找不到 manifest.yaml: {manifest_path}", file=sys.stderr)
        print(f"  请先跑 b_mode_extract.py 产出 manifest", file=sys.stderr)
        sys.exit(1)

    with open(manifest_path, 'r', encoding='utf-8') as f:
        manifest = yaml.safe_load(f)

    provider_name = manifest.get('assets_provider', 'placeholder')
    from assets_provider import get_provider

    # V3-1:把 brief.extracted.bidding_entity 透传给 provider 当 company_id,
    # 让 CuratedLocalAssetsProvider 扫到正确公司目录。bidding_entity 缺省时
    # 走 provider 自身的默认值(CuratedLocalAssetsProvider 默认 own_default)。
    provider_kwargs: dict = {}
    bidding_entity = (brief.get('extracted', {}) or {}).get('bidding_entity', '')
    if provider_name == 'curated_local':
        if bidding_entity:
            provider_kwargs['company_id'] = bidding_entity
        if args.non_interactive:
            provider_kwargs['non_interactive'] = True

    provider = get_provider(provider_name, **provider_kwargs)

    assembly_order = manifest.get('assembly_order', [])
    print(f"[信息] 组装 Part[{args.part}] '{part['name']}' (B 模式)"
          f",共 {len(assembly_order)} 项")
    print(f"  assets_provider: {provider_name}")

    # v2 补丁 2(字体修复):为 assembled.docx 的 Normal/Heading 样式设字体,
    # 避免 run fallback 到 docDefaults 主题字体(WPS 会渲染为日文 MS 明朝)
    from docx_builder import apply_default_styles

    doc = Document()
    apply_default_styles(doc)
    # 第一段:Part 标题
    p = doc.add_paragraph()
    run = p.add_run(part['name'])
    run.bold = True
    run.font.size = Pt(16)

    stats = {'inline_template': 0, 'asset_lookup': 0, 'self_drafted': 0}
    # V4-skel.4 V4-1b 占位: 累加每个 asset 的 missing_elements 扫描结果, 末尾统一写盘
    missing_acc: list = []
    for i, spec in enumerate(assembly_order):
        _dispatch(doc, spec, provider, part['name'], missing_acc=missing_acc)
        st = spec.get('source_type')
        if st in stats:
            stats[st] += 1

    assembled_path = b_dir / 'assembled.docx'
    doc.save(str(assembled_path))

    # V4-skel.4 V4-1b 占位: 写 missing_elements.yaml (跟 V4-4 attachments.yaml
    # 同型 sidecar)。仅当有 asset_lookup 真扫到 asset 时才写 (missing_acc 非空)。
    if missing_acc:
        missing_elements_path = b_dir / 'missing_elements.yaml'
        with open(missing_elements_path, 'w', encoding='utf-8') as f:
            yaml.safe_dump({'assets': missing_acc}, f,
                           allow_unicode=True, sort_keys=False)

    # .pending_marker 空文件
    marker_path = b_dir / '.pending_marker'
    marker_path.write_text('', encoding='utf-8')

    print()
    print("=" * 60)
    print(f"[完成] assembled.docx: {assembled_path}")
    print(f"[完成] .pending_marker: {marker_path} (占位状态标记)")
    print(f"  inline_template 段: {stats['inline_template']}")
    print(f"  asset_lookup 段:    {stats['asset_lookup']}")
    print(f"  self_drafted 段:    {stats['self_drafted']}")
    if missing_acc:
        print(f"[V4-1b 占位] missing_elements.yaml 已写 {len(missing_acc)} 个 asset 扫描结果")
    print()
    print("用户填充完成真实内容后,请手工删除 .pending_marker。")
    print("V45 整标合并器会检测 marker,列入 pending_manual_work.md。")
    print("=" * 60)


if __name__ == '__main__':
    main()
