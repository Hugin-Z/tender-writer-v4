# -*- coding: utf-8 -*-
"""
compliance_check.py · 合规终审脚本（阶段 5）

输入:
    output/tender_response.docx
    output/scoring_matrix.csv

输出:
    output/compliance_report.md

说明:
    这是自动化初筛工具，目标是尽量把明显的漏答、模板残留和格式风险提出来。
    它不能替代标书经理或法务的人工终审。
"""

from __future__ import annotations

import argparse
import csv
import json
import re
import sys
from datetime import datetime, timezone, timedelta
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from brief_schema import build_part_maps

TEMPLATE_RESIDUES = [
    "xxx 公司",
    "XXX公司",
    "XX公司",
    "xxx 项目",
    "XXX项目",
    "XX项目",
    # 占位符模式必须带尖括号或方括号,避免误伤正文里合法出现的"甲方单位"等词
    "<甲方单位>",
    "<甲方名称>",
    "<乙方名称>",
    "<项目名称>",
    "<投标人名称>",
    "[甲方单位]",
    "[甲方名称]",
    "[乙方名称]",
    "[项目名称]",
    "[投标人名称]",
    "【甲方单位】",
    "【项目名称】",
    "【投标人名称】",
    "【公司名称】",
    "TODO",
    "todo",
    "【待补充】",
    "【待填写】",
    "【待人工确认】",
    "示例文本",
]

SUBSTANTIAL_KEYWORDS = ["完全响应", "完全满足", "实质性响应", "无偏离", "满足要求"]
FORMAT_CHECKS = {
    "封面": ["投标文件", "投标人"],
    "目录": ["目录"],
}
NORMALIZE_PATTERN = re.compile(r"[\s\u3000，,。；;：:（）()\[\]【】\-—_/]")

# ── V3-8 fontTable 字体常量 ────────────────────────────────────────────
# 中文字体规范名 → 别名清单（归一化目标）
_CHINESE_CANONICAL: dict[str, list[str]] = {
    "宋体":     ["宋体", "SimSun", "宋体-简", "NSimSun", "STSong", "宋体-繁"],
    "黑体":     ["黑体", "SimHei", "STHeiti"],
    "仿宋":     ["仿宋", "FangSong", "仿宋_GB2312", "STFangsong"],
    "微软雅黑": ["微软雅黑", "Microsoft YaHei", "Microsoft YaHei UI"],
}
# 归一化后的合法字体集合（快速查找用）
_FONT_ALLOWED_NORMALIZED: frozenset[str] = frozenset(
    list(_CHINESE_CANONICAL.keys()) + ["Times New Roman", "Arial", "Calibri", "Cambria"]
)
# python-docx 默认 boilerplate 自带的字体（知道存在，容忍不报 warn）
_FONT_BOILERPLATE_TOLERATED: frozenset[str] = frozenset({
    "Symbol",
    "Courier",
    "ＭＳ 明朝",    # ＭＳ 明朝（全角，MS Mincho，python-docx 默认）
    "ＭＳ ゴシック",  # ＭＳ ゴシック（全角，MS Gothic，python-docx 默认）
})


def normalize_for_match(text: str) -> str:
    return NORMALIZE_PATTERN.sub("", text)


def read_docx_text(docx_path: Path) -> tuple[str, dict]:
    try:
        from docx import Document
    except ImportError:
        print("[错误] 缺少 python-docx 依赖，请先双击 install.bat 安装依赖。", file=sys.stderr)
        sys.exit(1)

    document = Document(str(docx_path))
    paragraph_texts: list[str] = []
    paragraph_styles: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            paragraph_texts.append(text)
            paragraph_styles.append(getattr(para.style, "name", ""))

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraph_texts.append(" | ".join(cells))

    full_text = "\n".join(paragraph_texts)
    format_info = {
        "paragraph_count": len(document.paragraphs),
        "non_empty_paragraphs": len(paragraph_texts),
        "heading_count": sum(1 for name in paragraph_styles if name.startswith("Heading")),
        "cover_preview": paragraph_texts[:12],
        "section_margins": [],
    }

    for section in document.sections:
        format_info["section_margins"].append({
            "top_cm": round(section.top_margin.cm, 2) if section.top_margin else None,
            "bottom_cm": round(section.bottom_margin.cm, 2) if section.bottom_margin else None,
            "left_cm": round(section.left_margin.cm, 2) if section.left_margin else None,
            "right_cm": round(section.right_margin.cm, 2) if section.right_margin else None,
        })

    return full_text, format_info


def load_matrix(matrix_csv: Path) -> list[dict]:
    rows: list[dict] = []
    with open(matrix_csv, "r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            rows.append(row)
    return rows


def extract_candidates(row: dict) -> list[str]:
    item = (row.get("评分项") or "").strip()
    keywords_field = (row.get("关键词") or "").strip()

    candidates: list[str] = []
    if item:
        candidates.append(item)
    if keywords_field:
        candidates.extend([part.strip() for part in re.split(r"[;；、/]+", keywords_field) if part.strip()])

    for token in re.split(r"[（）()\[\]【】、，,；;：:\s]+", item):
        token = token.strip()
        if 2 <= len(token) <= 10:
            candidates.append(token)

    unique: list[str] = []
    seen: set[str] = set()
    for candidate in candidates:
        normalized = normalize_for_match(candidate)
        if len(normalized) < 2 or normalized in seen:
            continue
        seen.add(normalized)
        unique.append(candidate)
    return unique


def check_coverage(
    docx_text: str,
    matrix: list[dict],
    mode_map: dict[str, str] | None = None,
) -> list[dict]:
    """评分项覆盖度校验。mode_map: {part_name: production_mode}，
    production_mode != 'A' 的评分项标 out_of_scope 跳过。"""
    normalized_doc = normalize_for_match(docx_text)
    results: list[dict] = []

    for row in matrix:
        item = (row.get("评分项") or "").strip()
        score = (row.get("分值") or "").strip()
        chapter = (row.get("应答章节") or "").strip()
        attribution = (row.get("评分项归属") or "").strip()
        candidates = extract_candidates(row)

        # 非 v1.1 范围（production_mode != A）跳过
        if mode_map and attribution:
            pm = mode_map.get(attribution, "")
            if pm and pm != "A":
                results.append({
                    "item": item,
                    "score": score,
                    "chapter": chapter,
                    "status": "out_of_scope",
                    "candidates": candidates,
                    "matched_candidates": [],
                })
                continue

        exact_hit = item and normalize_for_match(item) in normalized_doc
        matched_candidates = [
            candidate for candidate in candidates
            if normalize_for_match(candidate) in normalized_doc
        ]

        if exact_hit or len(matched_candidates) >= 2:
            status = "covered"
        elif len(matched_candidates) == 1:
            status = "partial"
        else:
            status = "missing"

        results.append({
            "item": item,
            "score": score,
            "chapter": chapter,
            "status": status,
            "candidates": candidates,
            "matched_candidates": matched_candidates,
        })

    return results


def check_substantial_response(docx_text: str, matrix: list[dict]) -> list[dict]:
    results: list[dict] = []
    lowered_doc = docx_text

    for row in matrix:
        item = (row.get("评分项") or "").strip()
        risk = (row.get("风险提示") or "").strip()
        if not any(marker in item or marker in risk for marker in ("★", "▲", "实质性响应")):
            continue

        candidates = extract_candidates(row)
        responded = False
        evidence = ""

        for candidate in candidates:
            start = 0
            while True:
                idx = lowered_doc.find(candidate, start)
                if idx == -1:
                    break
                window = lowered_doc[max(0, idx - 80): idx + len(candidate) + 120]
                matched_phrase = next((phrase for phrase in SUBSTANTIAL_KEYWORDS if phrase in window), "")
                if matched_phrase:
                    responded = True
                    evidence = f"{candidate} 附近出现“{matched_phrase}”"
                    break
                start = idx + len(candidate)
            if responded:
                break

        if not responded:
            global_phrase = next((phrase for phrase in SUBSTANTIAL_KEYWORDS if phrase in lowered_doc), "")
            if global_phrase:
                evidence = f"全文出现过“{global_phrase}”，但未能与该条款建立近邻关联"
            else:
                evidence = "未找到明确响应表述"

        results.append({
            "item": item,
            "responded": responded,
            "evidence": evidence,
        })

    return results


def check_template_residues(docx_text: str) -> list[str]:
    found: list[str] = []
    for residue in TEMPLATE_RESIDUES:
        if residue in docx_text:
            idx = docx_text.find(residue)
            context = docx_text[max(0, idx - 20): idx + len(residue) + 20].replace("\n", " ")
            found.append(f"残留“{residue}” 上下文: ...{context}...")
    return found


def should_section_only(docx_filename: str, explicit_flag: bool) -> bool:
    """V3-5:决定本次 compliance_check 是否走 section-only 模式。

    优先级:
    1. 用户显式 --section-only → True
    2. 文件名含 'tender_response' 且不含 'final_response' → True
    3. 否则 False(全项检查;final_response.docx 等终端产物走此路径)

    section-only 模式跳过 FORMAT_CHECKS 字典(封面/目录关键词),
    保留 section margins + heading count 检查。
    """
    if explicit_flag:
        return True
    name = docx_filename.lower()
    if "tender_response" in name and "final_response" not in name:
        return True
    return False


def check_format(docx_text: str, format_info: dict, section_only: bool = False) -> list[str]:
    issues: list[str] = []
    expected_margins = {"top_cm": 2.54, "bottom_cm": 2.54, "left_cm": 3.17, "right_cm": 3.17}

    margins = format_info.get("section_margins", [])
    if not margins:
        issues.append("未能读取到 section 页边距信息")
    else:
        for index, section_margin in enumerate(margins):
            for key, expected in expected_margins.items():
                actual = section_margin.get(key)
                if actual is None:
                    issues.append(f"section[{index}] 缺少 {key}")
                elif abs(actual - expected) > 0.05:
                    issues.append(f"section[{index}] {key} 实际 {actual} cm，期望 {expected} cm")

    # V3-5:section-only 模式跳过封面/目录关键词检查
    # (适用于 tender_response.docx 等分章节产物,完整封面/目录在终端 final_response.docx)
    if not section_only:
        normalized_doc = normalize_for_match(docx_text)
        for check_name, keywords in FORMAT_CHECKS.items():
            if not all(normalize_for_match(keyword) in normalized_doc for keyword in keywords):
                kw_list = ', '.join(keywords)
                issues.append(f'missing {check_name} keywords: {kw_list}')

    if format_info.get("heading_count", 0) == 0:
        issues.append("未检测到 Heading 样式标题，目录和层级结构可能未按 docx_builder 规范生成")

    return issues



def _normalize_font_alias(font_name: str) -> str:
    """把字体别名归一化到 canonical name；未识别则原样返回。"""
    for canonical, aliases in _CHINESE_CANONICAL.items():
        if font_name in aliases:
            return canonical
    return font_name


def _check_fonttable(docx_path) -> list[str]:
    """V3-8：解析 word/fontTable.xml，返回非白名单且非 boilerplate 字体的 warn 列表。"""
    import zipfile as _zipfile
    import xml.etree.ElementTree as _ET

    NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    try:
        with _zipfile.ZipFile(docx_path) as z:
            if "word/fontTable.xml" not in z.namelist():
                return []
            xml_bytes = z.read("word/fontTable.xml")
    except (_zipfile.BadZipFile, KeyError):
        return ["字体检查(fontTable)：docx 解压失败"]

    try:
        root = _ET.fromstring(xml_bytes.decode("utf-8"))
    except _ET.ParseError:
        return ["字体检查(fontTable)：xml 解析失败"]

    fonts = [f.get(f"{NS}name") for f in root.findall(f"{NS}font")]
    fonts = [f for f in fonts if f]

    suspicious = []
    for font in fonts:
        normalized = _normalize_font_alias(font)
        if normalized in _FONT_ALLOWED_NORMALIZED:
            continue
        if font in _FONT_BOILERPLATE_TOLERATED or normalized in _FONT_BOILERPLATE_TOLERATED:
            continue
        suspicious.append(font)

    if not suspicious:
        return []
    return [
        f"字体检查(fontTable warn)：docx 声明非白名单字体 {suspicious}，"
        f"非 python-docx 默认 boilerplate；若被实际引用可能导致 WPS fallback 渲染问题。"
        f"白名单 + boilerplate 容忍见 compliance_check.py _FONT_ALLOWED_NORMALIZED / _FONT_BOILERPLATE_TOLERATED。"
    ]


def check_font_safety(docx_path: Path) -> list[str]:
    """
    v2 补丁 2:字体安全检查(回归检查项)。

    扫 docx 的 word/styles.xml + word/document.xml,确保:
    - Normal 样式必须有 <w:rFonts> 且 w:eastAsia 在中文白名单
      (宋体 / 仿宋 / 仿宋_GB2312 / 黑体 / 微软雅黑)
    - 所有 run 级 <w:rFonts> 的 w:eastAsia 在上述白名单
    - 所有 run 级 <w:rFonts> 的 w:ascii 在 ("Times New Roman",)

    返回 issues 列表,空表示通过。

    根因背景:若 Normal 样式未设字体,Word/WPS 会 fallback 到 docDefaults 的
    minorEastAsia 主题字体,在主题文件 themeFontLang 为日语或 CJK 解析失败时
    渲染为 MS 明朝(日文字体),导致 WPS 提示"字体缺失"。
    """
    import re as _re
    import zipfile as _zipfile

    issues: list[str] = []
    ALLOWED_CJK = {a for aliases in _CHINESE_CANONICAL.values() for a in aliases}
    ALLOWED_ASCII = {"Times New Roman"}

    try:
        with _zipfile.ZipFile(docx_path) as z:
            styles_xml = z.read("word/styles.xml").decode("utf-8")
            document_xml = z.read("word/document.xml").decode("utf-8")
    except (KeyError, _zipfile.BadZipFile) as exc:
        issues.append(f"字体检查:docx 解压失败 {exc}")
        return issues

    # 检查 Normal 样式
    m = _re.search(
        r'<w:style[^>]*w:styleId="Normal"[^>]*>(.*?)</w:style>',
        styles_xml, _re.DOTALL)
    if not m:
        issues.append("字体检查:styles.xml 中未找到 Normal 样式")
    else:
        normal_body = m.group(1)
        r = _re.search(r'<w:rFonts\s[^/]*/?>', normal_body)
        if not r:
            issues.append(
                "字体检查:Normal 样式无 <w:rFonts> 标签,run 会 fallback 到 "
                "docDefaults minorEastAsia 主题字体,WPS 可能渲染为日文 MS 明朝"
            )
        else:
            ea_m = _re.search(r'w:eastAsia="([^"]+)"', r.group(0))
            if not ea_m:
                issues.append("字体检查:Normal 样式 <w:rFonts> 缺 w:eastAsia 属性")
            elif ea_m.group(1) not in ALLOWED_CJK:
                issues.append(
                    f"字体检查:Normal 样式 eastAsia={ea_m.group(1)!r} "
                    f"不在白名单 {sorted(ALLOWED_CJK)}"
                )

    # 抽样检查 run 级 rFonts 异常字体
    run_rfonts = _re.findall(r'<w:rFonts\s[^/]*/?>', document_xml)
    bad_ea: dict = {}
    bad_ascii: dict = {}
    for tag in run_rfonts:
        ea_m = _re.search(r'w:eastAsia="([^"]+)"', tag)
        if ea_m and ea_m.group(1) not in ALLOWED_CJK:
            bad_ea[ea_m.group(1)] = bad_ea.get(ea_m.group(1), 0) + 1
        asc_m = _re.search(r'w:ascii="([^"]+)"', tag)
        if asc_m and asc_m.group(1) not in ALLOWED_ASCII:
            bad_ascii[asc_m.group(1)] = bad_ascii.get(asc_m.group(1), 0) + 1

    for font, cnt in bad_ea.items():
        issues.append(
            f"字体检查:检测到非白名单 eastAsia 字体 {font!r} × {cnt} 处 "
            f"(白名单 {sorted(ALLOWED_CJK)})"
        )
    for font, cnt in bad_ascii.items():
        issues.append(
            f"字体检查:检测到非白名单 ascii 字体 {font!r} × {cnt} 处 "
            f"(白名单 {sorted(ALLOWED_ASCII)})"
        )

    # ── V3-8 新增：fontTable.xml 级检查（warn 级）
    issues.extend(_check_fonttable(docx_path))

    return issues


def check_title_keyword_coverage(
    matrix: list[dict],
    chapters_dir: Path,
    brief_path: Path,
) -> list[dict]:
    """校验评分项关键词是否出现在章节 markdown 标题中。

    仅校验 production_mode==A 的 Part。
    返回: [{"item", "keywords", "matched_title", "status", "note"}, ...]
    status: "pass" / "fail" / "skip" / "no_keyword"
    """
    if not brief_path.exists():
        print(f"[警告] 标题覆盖校验跳过：未找到 {brief_path}",
              file=sys.stderr)
        return []

    brief_data = json.loads(brief_path.read_text(encoding="utf-8"))
    part_name_map, _ = build_part_maps(brief_data.get("response_file_parts", []))
    mode_a_parts: dict[str, str] = {}
    for name, part in part_name_map.items():
        if part.get("production_mode") == "A":
            mode_a_parts[name] = part["id"]

    if not mode_a_parts:
        return []

    results: list[dict] = []

    for part_name, part_id in mode_a_parts.items():
        part_dir = chapters_dir / part_id
        if not part_dir.exists():
            for row in matrix:
                if (row.get("评分项归属") or "").strip() == part_name:
                    results.append({
                        "item": (row.get("评分项") or "").strip(),
                        "keywords": (row.get("关键词") or "").strip(),
                        "matched_title": "",
                        "status": "skip",
                        "note": f"章节目录 {part_id}/ 不存在",
                    })
            continue

        # 收集所有章节文件的标题行
        headings: list[str] = []
        for md_file in sorted(part_dir.glob("*.md")):
            for line in md_file.read_text(encoding="utf-8").splitlines():
                if line.startswith("#"):
                    headings.append(line.lstrip("#").strip())

        # 逐条评分项校验
        for row in matrix:
            attr = (row.get("评分项归属") or "").strip()
            if attr != part_name:
                continue

            item_name = (row.get("评分项") or "").strip()
            keywords_str = (row.get("关键词") or "").strip()

            if not keywords_str:
                results.append({
                    "item": item_name,
                    "keywords": "",
                    "matched_title": "",
                    "status": "no_keyword",
                    "note": "关键词缺失，无法校验",
                })
                continue

            kws = [k.strip() for k in re.split(r"[;；、/]+", keywords_str)
                   if k.strip()]

            matched_title = ""
            for kw in kws:
                for heading in headings:
                    if kw in heading:
                        matched_title = heading
                        break
                if matched_title:
                    break

            results.append({
                "item": item_name,
                "keywords": keywords_str,
                "matched_title": matched_title if matched_title else "",
                "status": "pass" if matched_title else "fail",
                "note": "",
            })

    return results


def check_mandatory_element_coverage(
    matrix: list[dict],
    chapters_dir: Path,
    brief_path: Path,
) -> list[dict]:
    """校验每个 A 模式评分项的必备要素是否出现在对应章节中。

    返回: [{"item", "elements", "hit", "miss", "rate", "status"}, ...]
    status: "pass" / "fail" / "skip" / "no_elements"
    """
    if not brief_path.exists():
        return []

    brief_data = json.loads(brief_path.read_text(encoding="utf-8"))
    part_name_map, _ = build_part_maps(brief_data.get("response_file_parts", []))
    mode_a_parts: dict[str, str] = {}
    for name, part in part_name_map.items():
        if part.get("production_mode") == "A":
            mode_a_parts[name] = part["id"]

    if not mode_a_parts:
        return []

    results: list[dict] = []

    for part_name, part_id in mode_a_parts.items():
        part_dir = chapters_dir / part_id
        # Collect all chapter text in this Part
        all_text = ""
        if part_dir.exists():
            for md_file in sorted(part_dir.glob("*.md")):
                all_text += md_file.read_text(encoding="utf-8") + "\n"

        for row in matrix:
            attr = (row.get("评分项归属") or "").strip()
            if attr != part_name:
                continue

            item_name = (row.get("评分项") or "").strip()
            elements_str = (row.get("必备要素") or "").strip()

            if not elements_str or elements_str == "__PENDING_AI__":
                results.append({
                    "item": item_name,
                    "elements": "",
                    "hit": [],
                    "miss": [],
                    "rate": 0.0,
                    "status": "no_elements",
                })
                continue

            if not part_dir.exists():
                results.append({
                    "item": item_name,
                    "elements": elements_str,
                    "hit": [],
                    "miss": [],
                    "rate": 0.0,
                    "status": "skip",
                })
                continue

            elements = [e.strip() for e in re.split(r"[;；、/]+", elements_str)
                        if e.strip()]
            hit = [e for e in elements if e in all_text]
            miss = [e for e in elements if e not in all_text]
            rate = len(hit) / len(elements) * 100 if elements else 0.0

            results.append({
                "item": item_name,
                "elements": elements_str,
                "hit": hit,
                "miss": miss,
                "rate": rate,
                "status": "pass" if not miss else "fail",
            })

    return results


def render_report(
    coverage: list[dict],
    substantial: list[dict],
    residues: list[str],
    format_issues: list[str],
    title_keyword: list[dict],
    mandatory_elements: list[dict],
    docx_path: Path,
    matrix_csv: Path,
) -> str:
    out_of_scope = [row for row in coverage if row["status"] == "out_of_scope"]
    in_scope = [row for row in coverage if row["status"] != "out_of_scope"]
    total = len(in_scope)
    covered = [row for row in in_scope if row["status"] == "covered"]
    partial = [row for row in in_scope if row["status"] == "partial"]
    missing = [row for row in in_scope if row["status"] == "missing"]

    sub_total = len(substantial)
    sub_responded = [row for row in substantial if row["responded"]]
    sub_missing = [row for row in substantial if not row["responded"]]

    lines: list[str] = []
    lines.append("# 合规终审报告（compliance_report）")
    lines.append("")
    lines.append(f"> **被检 docx**: {docx_path}")
    lines.append(f"> **评分矩阵**: {matrix_csv}")
    lines.append("> **说明**: 本报告是自动化初筛结果，不能替代标书经理和法务的人工终审。")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## 一、评分项覆盖度检查")
    lines.append("")
    lines.append(f"- 评分项总数: {total}")
    lines.append(f"- 明确覆盖: {len(covered)}")
    lines.append(f"- 弱覆盖: {len(partial)}")
    lines.append(f"- 未覆盖: {len(missing)}")
    lines.append(f"- 覆盖率: {(len(covered) / total * 100) if total else 0:.1f}%")
    lines.append("")

    lines.append("### 未覆盖清单（必须返回阶段 4 补写）")
    lines.append("")
    if missing:
        for row in missing:
            preview = "，".join(row["candidates"][:5])
            lines.append(f"- ❌ **{row['item']}**（分值 {row['score']}）")
            lines.append(f"  候选关键词: {preview}")
            lines.append(f"  目标章节: {row['chapter']}")
    else:
        lines.append("- ✅ 未发现明确漏答项")
    lines.append("")

    lines.append("### 弱覆盖清单（建议人工复核）")
    lines.append("")
    if partial:
        for row in partial:
            matched = "，".join(row["matched_candidates"])
            lines.append(f"- ⚠️ **{row['item']}** 仅命中关键词: {matched}")
    else:
        lines.append("- ✅ 未发现弱覆盖项")
    lines.append("")

    if out_of_scope:
        # v2 P13:对外不用"v1.1 范围"这类内部版本术语
        lines.append("### 不适用的评分项(归属非技术部分,跳过自动覆盖度检查)")
        lines.append("")
        for row in out_of_scope:
            lines.append(f"- ⏭️ **{row['item']}**(分值 {row['score']})"
                         f" [归属非技术部分,由素材组装/模板填充类 Part 承载,"
                         f"跳过自动覆盖度检查]")
        lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## 二、★/▲ 条款响应检查")
    lines.append("")
    lines.append(f"- ★/▲ 条款总数: {sub_total}")
    lines.append(f"- 明确响应: {len(sub_responded)}")
    lines.append(f"- 未明确响应: {len(sub_missing)}")
    lines.append("")

    lines.append("### 未明确响应清单")
    lines.append("")
    if sub_missing:
        for row in sub_missing:
            lines.append(f"- ❌ **{row['item']}**：{row['evidence']}")
    else:
        if sub_total == 0:
            lines.append("- ⚠️ 评分矩阵中未标出任何 ★/▲ 条款，请人工复核 tender_brief.md")
        else:
            lines.append("- ✅ 所有 ★/▲ 条款均检测到近邻响应表述")
    lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## 三、模板残留检查")
    lines.append("")
    if residues:
        for residue in residues:
            lines.append(f"- ❌ {residue}")
    else:
        lines.append("- ✅ 未检测到常见模板残留（如 XXX 公司、TODO、【待补充】）")
    lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## 四、格式检查")
    lines.append("")
    if format_issues:
        for issue in format_issues:
            lines.append(f"- ⚠️ {issue}")
    else:
        lines.append("- ✅ 页边距、封面/目录特征、标题样式未发现明显问题")
    lines.append("")

    # 五、标题关键词覆盖校验
    lines.append("---")
    lines.append("")
    lines.append("## 五、标题关键词覆盖校验")
    lines.append("")

    tk_pass = [r for r in title_keyword if r["status"] == "pass"]
    tk_fail = [r for r in title_keyword if r["status"] == "fail"]
    tk_skip = [r for r in title_keyword if r["status"] == "skip"]
    tk_no_kw = [r for r in title_keyword if r["status"] == "no_keyword"]
    tk_checkable = len(tk_pass) + len(tk_fail)

    if not title_keyword:
        lines.append("- ⚠️ 标题覆盖校验未执行(缺少 tender_brief.json 或无主撰类 Part)")
    else:
        lines.append("| 评分项 | 关键词 | 命中标题 | 状态 |")
        lines.append("|---|---|---|---|")
        for r in title_keyword:
            item = r["item"] or "[未填]"
            kws = r["keywords"] or "[缺失]"
            matched = r["matched_title"] if r["matched_title"] else "(未命中)"
            if r["status"] == "pass":
                status_icon = "✓"
            elif r["status"] == "fail":
                status_icon = "✗"
            elif r["status"] == "skip":
                status_icon = f"跳过({r['note']})"
            else:
                status_icon = r["note"]
            lines.append(f"| {item} | {kws} | {matched} | {status_icon} |")
        lines.append("")
        if tk_checkable > 0:
            rate = len(tk_pass) / tk_checkable * 100
            lines.append(f"总计：{len(tk_pass)}/{tk_checkable} 评分项标题覆盖通过"
                         f"（覆盖率 {rate:.0f}%）")
        if tk_skip:
            lines.append(f"跳过：{len(tk_skip)} 项（章节目录不存在）")
        if tk_no_kw:
            lines.append(f"无法校验：{len(tk_no_kw)} 项（关键词缺失）")
    lines.append("")

    # 六、必备要素覆盖校验
    me_pass = [r for r in mandatory_elements if r["status"] == "pass"]
    me_fail = [r for r in mandatory_elements if r["status"] == "fail"]
    me_skip = [r for r in mandatory_elements if r["status"] == "skip"]
    me_no = [r for r in mandatory_elements if r["status"] == "no_elements"]
    me_checkable = len(me_pass) + len(me_fail)

    lines.append("---")
    lines.append("")
    lines.append("## 六、必备要素覆盖校验")
    lines.append("")
    if not mandatory_elements:
        lines.append("- ⚠️ 必备要素校验未执行(缺少 tender_brief.json 或无主撰类 Part)")
    elif me_checkable == 0:
        lines.append("- ⚠️ 所有评分项的必备要素列为空或待填充，无法校验")
    else:
        lines.append("| 评分项 | 必备要素 | 命中 | 未命中 | 覆盖率 | 状态 |")
        lines.append("|---|---|---|---|---|---|")
        for r in mandatory_elements:
            item = r["item"] or "[未填]"
            elements = r["elements"] or "[空]"
            if r["status"] == "pass":
                lines.append(f"| {item} | {elements} | {len(r['hit'])} | 0 | {r['rate']:.0f}% | ✓ |")
            elif r["status"] == "fail":
                miss_str = "; ".join(r["miss"])
                lines.append(f"| {item} | {elements} | {len(r['hit'])} | {len(r['miss'])} | {r['rate']:.0f}% | ✗ 未命中: {miss_str} |")
            elif r["status"] == "no_elements":
                lines.append(f"| {item} | [待填充] | - | - | - | 跳过 |")
            else:
                lines.append(f"| {item} | {elements} | - | - | - | 跳过({r.get('note', '章节不存在')}) |")
        lines.append("")
        if me_checkable > 0:
            avg_rate = sum(r["rate"] for r in mandatory_elements if r["status"] in ("pass", "fail")) / me_checkable
            lines.append(f"总计：{len(me_pass)}/{me_checkable} 评分项必备要素覆盖通过"
                         f"（平均覆盖率 {avg_rate:.0f}%）")
        if me_no:
            lines.append(f"待填充：{len(me_no)} 项（必备要素列为空）")
    lines.append("")

    # 七、结论
    title_kw_failures = len(tk_fail)
    me_failures = len(me_fail)
    serious_issues = len(missing) + len(sub_missing) + len(residues) + title_kw_failures + me_failures
    lines.append("---")
    lines.append("")
    lines.append("## 七、结论")
    lines.append("")
    if serious_issues == 0:
        lines.append("- ✅ 自动检查未发现严重问题，但仍需完成人工终审后再提交。")
    else:
        parts = []
        if missing:
            parts.append(f"漏答 {len(missing)} 项")
        if sub_missing:
            parts.append(f"★/▲ 未明确响应 {len(sub_missing)} 项")
        if residues:
            parts.append(f"模板残留 {len(residues)} 项")
        if title_kw_failures:
            parts.append(f"标题关键词未覆盖 {title_kw_failures} 项")
        if me_failures:
            parts.append(f"必备要素未覆盖 {me_failures} 项")
        lines.append(
            f"- ❌ 自动检查发现 {serious_issues} 项严重问题"
            f"（{'、'.join(parts)}），"
            "必须返回阶段 4 修正后重新检查。"
        )

    return "\n".join(lines) + "\n"


def build_compliance_metrics(
    coverage: list[dict],
    substantial: list[dict],
    residues: list[str],
    format_issues: list[str],
    title_keyword: list[dict],
    mandatory_elements: list[dict],
    project: str,
) -> dict:
    """从内部统计数据直接构建 compliance_metrics dict。"""
    in_scope = [r for r in coverage if r["status"] != "out_of_scope"]
    covered = [r for r in in_scope if r["status"] == "covered"]
    partial = [r for r in in_scope if r["status"] == "partial"]
    missing = [r for r in in_scope if r["status"] == "missing"]
    total = len(in_scope)

    # 分离 ★ 和 ▲ 条款
    star_items = [r for r in substantial if "★" in r["item"]]
    triangle_items = [r for r in substantial if "▲" in r["item"] and "★" not in r["item"]]

    # 模板残留: 提取残留内容(不含上下文)
    residue_items = []
    for r in residues:
        m = re.search(r'残留"([^"]+)"', r)
        if m:
            residue_items.append(m.group(1))

    # 标题关键词覆盖
    tk_pass = [r for r in title_keyword if r["status"] == "pass"]
    tk_fail = [r for r in title_keyword if r["status"] == "fail"]
    tk_checkable = len(tk_pass) + len(tk_fail)

    tz = timezone(timedelta(hours=8))
    now = datetime.now(tz)

    return {
        "generated_at": now.strftime("%Y-%m-%dT%H:%M:%S+08:00"),
        "project": project,
        "coverage_pct": round(len(covered) / total * 100, 1) if total else 0.0,
        "coverage_total": total,
        "coverage_explicit": len(covered),
        "coverage_weak": len(partial),
        "coverage_missing": len(missing),
        "star_total": len(star_items),
        "star_explicit": sum(1 for r in star_items if r["responded"]),
        "star_unclear": sum(1 for r in star_items if not r["responded"]),
        "triangle_total": len(triangle_items),
        "triangle_explicit": sum(1 for r in triangle_items if r["responded"]),
        "triangle_unclear": sum(1 for r in triangle_items if not r["responded"]),
        "template_residue_count": len(residues),
        "template_residue_items": residue_items,
        "format_issues_count": len(format_issues),
        "format_issues_items": list(format_issues),
        "title_keyword_coverage_pct": round(len(tk_pass) / tk_checkable * 100, 1) if tk_checkable else 0.0,
        "title_keyword_hit": len(tk_pass),
        "title_keyword_total": tk_checkable,
        "mandatory_element_coverage_pct": round(
            sum(r["rate"] for r in mandatory_elements if r["status"] in ("pass", "fail"))
            / max(1, sum(1 for r in mandatory_elements if r["status"] in ("pass", "fail")))
            , 1),
        "mandatory_element_total": sum(
            len(r["hit"]) + len(r["miss"])
            for r in mandatory_elements if r["status"] in ("pass", "fail")),
        "mandatory_element_hit": sum(
            len(r["hit"])
            for r in mandatory_elements if r["status"] in ("pass", "fail")),
        "mandatory_element_missed_items": [
            {"item": r["item"], "missed": r["miss"]}
            for r in mandatory_elements if r["status"] == "fail"
        ],
    }


def main() -> None:
    parser = argparse.ArgumentParser(
        description="合规终审（阶段 5）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("docx_path", help="最终标书 docx 路径")
    parser.add_argument("matrix_csv", help="评分矩阵 CSV 路径")
    parser.add_argument("--out", default="output", help="报告输出目录，默认 output/")
    parser.add_argument(
        "--section-only", action="store_true",
        help="V3-5:跳过封面/目录关键词检查(适用于 tender_response.docx 等分章节"
             "产物)。文件名含 tender_response 时自动启用,本 flag 显式覆盖。",
    )
    args = parser.parse_args()

    # V3-3: timing hook
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from _timing_hook import stage_timer
    # matrix 在 project/output/scoring_matrix.csv,推 project_dir
    _project_dir = Path(args.matrix_csv).resolve().parent.parent

    with stage_timer("compliance_check", _project_dir):
        _main_body(args)


def _main_body(args) -> None:
    docx_path = Path(args.docx_path)
    matrix_csv = Path(args.matrix_csv)

    # V53: 未 review 则硬失败
    from brief_schema import ensure_reviewed
    ensure_reviewed(Path(args.out))

    if not docx_path.exists():
        print(f"[错误] 找不到 docx 文件: {docx_path}", file=sys.stderr)
        sys.exit(1)
    if not matrix_csv.exists():
        print(f"[错误] 找不到评分矩阵: {matrix_csv}", file=sys.stderr)
        sys.exit(1)

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    print(f"[信息] 正在读取 docx: {docx_path}")
    docx_text, format_info = read_docx_text(docx_path)
    print(f"[信息] 正在读取评分矩阵: {matrix_csv}")
    matrix = load_matrix(matrix_csv)

    # 构建 part_name → production_mode 映射(用于区分 v1.1 范围)
    brief_path = matrix_csv.parent / "tender_brief.json"
    mode_map: dict[str, str] = {}
    brief_parts: list[dict] = []
    if brief_path.exists():
        brief_data = json.loads(brief_path.read_text(encoding="utf-8"))
        brief_parts = brief_data.get("response_file_parts", [])
        part_name_map, _ = build_part_maps(brief_parts)
        for name, part in part_name_map.items():
            pm = part.get("production_mode", "")
            if name and pm:
                mode_map[name] = pm

    # V3-5:决定是否走 section-only 模式
    section_only = should_section_only(docx_path.name, args.section_only)
    if section_only:
        print(
            f"[信息] section-only 模式启用(跳过封面/目录关键词检查;触发原因:"
            f"{'显式 --section-only' if args.section_only else '文件名含 tender_response'})",
            file=sys.stderr,
        )

    coverage = check_coverage(docx_text, matrix, mode_map)
    substantial = check_substantial_response(docx_text, matrix)
    residues = check_template_residues(docx_text)
    format_issues = check_format(docx_text, format_info, section_only=section_only)
    # v2 补丁 2:字体安全检查(Normal 样式字体 + run 级 rFonts 白名单)
    font_issues = check_font_safety(docx_path)
    if font_issues:
        format_issues.extend(font_issues)

    # 标题关键词覆盖校验（自动从 matrix_csv 同目录检测 brief + chapters）
    brief_path = matrix_csv.parent / "tender_brief.json"
    chapters_dir = matrix_csv.parent / "chapters"
    title_keyword = check_title_keyword_coverage(matrix, chapters_dir, brief_path)
    tk_fail = sum(1 for r in title_keyword if r["status"] == "fail")
    print(f"[信息] 标题关键词覆盖校验: {len(title_keyword)} 条评分项, "
          f"{tk_fail} 条未覆盖", file=sys.stderr)

    # 必备要素覆盖校验
    mandatory_elements = check_mandatory_element_coverage(matrix, chapters_dir, brief_path)
    me_fail = sum(1 for r in mandatory_elements if r["status"] == "fail")
    print(f"[信息] 必备要素覆盖校验: {len(mandatory_elements)} 条评分项, "
          f"{me_fail} 条未覆盖", file=sys.stderr)

    report_path = out_dir / "compliance_report.md"
    report_path.write_text(
        render_report(coverage, substantial, residues, format_issues,
                      title_keyword, mandatory_elements, docx_path, matrix_csv),
        encoding="utf-8",
    )

    # 推断项目名(从 matrix_csv 路径: projects/{project}/output/...)
    project_name = ""
    for parent in matrix_csv.resolve().parents:
        if parent.parent.name == "projects":
            project_name = parent.name
            break
    if not project_name:
        project_name = matrix_csv.parent.name

    metrics_dict = build_compliance_metrics(
        coverage, substantial, residues, format_issues,
        title_keyword, mandatory_elements, project_name,
    )
    metrics_path = out_dir / "compliance_metrics.json"
    with open(metrics_path, "w", encoding="utf-8") as mf:
        json.dump(metrics_dict, mf, ensure_ascii=False, indent=2)
        mf.write("\n")

    print(f"[完成] 合规报告已写入: {report_path}")
    print(f"[完成] 合规指标已写入: {metrics_path}")
    print()
    print("=" * 60)
    in_scope = [row for row in coverage if row["status"] != "out_of_scope"]
    oos_count = sum(1 for row in coverage if row["status"] == "out_of_scope")
    print(f"自动检查统计: 明确漏答 {sum(1 for row in in_scope if row['status'] == 'missing')} 项，"
          f"弱覆盖 {sum(1 for row in in_scope if row['status'] == 'partial')} 项，"
          f"范围外跳过 {oos_count} 项，"
          f"★/▲ 未明确响应 {sum(1 for row in substantial if not row['responded'])} 项，"
          f"模板残留 {len(residues)} 项，格式问题 {len(format_issues)} 项，"
          f"标题关键词未覆盖 {tk_fail} 项")
    print("=" * 60)


if __name__ == "__main__":
    main()
