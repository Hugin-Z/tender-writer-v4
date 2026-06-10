# -*- coding: utf-8 -*-
"""make_1_docx.py · 生成 V4-1a 测试 anchor 1.docx (合成测试 fixture)

合成 V4-1a 测试 anchor, 供 scripts/tests/test_b_mode_v4_1a.py 验证 B asset
表格 + 段落样式 + 字体 + 字号保真链路 (10 case)。所有业务字段为合成占位
("示例项目 A-H" / "示例采购单位 N" 等)。

结构锚 (test_b_mode_v4_1a.py 实测固定):
- 9 行 x 7 列表格, tblStyle styleId="12" name="Table Grid"
- 列宽 dxa: 551 / 965 / 1314 / 2268 / 993 / 992 / 1213
- 表外 1 段 Heading 段, pStyle styleId="3" name="heading 2",
  含 numPr (numId=2 multilevel decimal) + spacing + ind 字面
- 所有 run rFonts 只 hint="eastAsia" 无字体名 (触发 V4-1a.6 字体 fix 路径)
- 真写 numbering.xml: abstractNumId=0 multilevel decimal 7 ilvl + num numId=2
  (case_6 numId strip 路径必须有真 numId 才有 strip 对象)

运行: .venv/Scripts/python.exe assets/公司资质/own_demo/_raw/_generators/make_1_docx.py
"""
from __future__ import annotations

import argparse
import re
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap as _nsmap  # noqa: F401
from lxml import etree

# REPO = __file__.parents[5]: _generators -> _raw -> own_demo -> 公司资质 -> assets -> REPO
REPO = Path(__file__).resolve().parents[5]
DEFAULT_OUT = REPO / "assets" / "公司资质" / "own_demo" / "_raw" / "1.docx"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSMAP = {"w": W_NS}

COL_WIDTHS_DXA = ["551", "965", "1314", "2268", "993", "992", "1213"]
HEADER_ROW = ["序号", "完成时间", "项目名称", "合同金额", "甲方", "联系人", "联系电话"]
DATA_ROWS = [
    ["1", "2024-01-15", "示例项目 A", "100 万元", "示例采购单位 1", "示例联系人 1", "__PENDING_USER__"],
    ["2", "2024-02-15", "示例项目 B", "200 万元", "示例采购单位 2", "示例联系人 2", "__PENDING_USER__"],
    ["3", "2024-03-15", "示例项目 C", "300 万元", "示例采购单位 3", "示例联系人 3", "__PENDING_USER__"],
    ["4", "2024-04-15", "示例项目 D", "400 万元", "示例采购单位 4", "示例联系人 4", "__PENDING_USER__"],
    ["5", "2024-05-15", "示例项目 E", "500 万元", "示例采购单位 5", "示例联系人 5", "__PENDING_USER__"],
    ["6", "2024-06-15", "示例项目 F", "600 万元", "示例采购单位 6", "示例联系人 6", "__PENDING_USER__"],
    ["7", "2024-07-15", "示例项目 G", "700 万元", "示例采购单位 7", "示例联系人 7", "__PENDING_USER__"],
    ["8", "2024-08-15", "示例项目 H", "800 万元", "示例采购单位 8", "示例联系人 8", "__PENDING_USER__"],
]


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def _make_elem(tag: str, **attrs) -> etree._Element:
    e = etree.SubElement.__self__  # dummy ref to silence linter
    el = etree.Element(_w(tag), nsmap=NSMAP)
    for k, v in attrs.items():
        el.set(_w(k), v)
    return el


def _ensure_styles(doc: Document) -> None:
    """加自定义 styleId='12' (name='Table Grid') + styleId='3' (name='heading 2')。

    docx_builder.apply_default_styles 后 master 端 TableGrid + Heading2 styleId
    已存在; 但源 anchor 字面用数字 styleId, 让 _build_style_id_map 按 name 做 rebind。
    所以 fixture 必须有 styleId='12' name='Table Grid' 和 styleId='3' name='heading 2'。
    """
    styles_el = doc.styles.element

    # 删除可能冲突的 default Heading2 / TableGrid 让源端字面 12/3 唯一指向
    # (实际 default Document() 不带 Heading2 / TableGrid, 不冲突)

    # styleId="12" name="Table Grid" (table style, 仿 master TableGrid)
    s12 = etree.SubElement(styles_el, _w("style"))
    s12.set(_w("type"), "table")
    s12.set(_w("styleId"), "12")
    n12 = etree.SubElement(s12, _w("name"))
    n12.set(_w("val"), "Table Grid")
    bp12 = etree.SubElement(s12, _w("basedOn"))
    bp12.set(_w("val"), "TableNormal")
    # 表格边框
    tblPr = etree.SubElement(s12, _w("tblPr"))
    tblBorders = etree.SubElement(tblPr, _w("tblBorders"))
    for b in ("top", "left", "bottom", "right", "insideH", "insideV"):
        be = etree.SubElement(tblBorders, _w(b))
        be.set(_w("val"), "single")
        be.set(_w("sz"), "4")
        be.set(_w("space"), "0")
        be.set(_w("color"), "auto")

    # styleId="3" name="heading 2" (paragraph style)
    s3 = etree.SubElement(styles_el, _w("style"))
    s3.set(_w("type"), "paragraph")
    s3.set(_w("styleId"), "3")
    n3 = etree.SubElement(s3, _w("name"))
    n3.set(_w("val"), "heading 2")
    bp3 = etree.SubElement(s3, _w("basedOn"))
    bp3.set(_w("val"), "Normal")
    nx3 = etree.SubElement(s3, _w("next"))
    nx3.set(_w("val"), "Normal")


def _ensure_numbering(doc: Document) -> None:
    """真写 numbering.xml: abstractNum 0 multilevel decimal 7 ilvl + num 2 → 0。

    case_6 验证注入后 doc.xml 内 numId 全 strip; 源 anchor 必须有 numId
    引用才有 strip 对象, 所以 fixture 必须真写 numbering part 且含 abstractNum
    + num 真定义 (V4-1a.9 主动降级路径才能被触发执行)。
    """
    # python-docx 不直接暴露 numbering_part, 通过 part 注册
    from docx.opc.constants import CONTENT_TYPE, RELATIONSHIP_TYPE as RT
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    numbering_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:abstractNum w:abstractNumId="0">
    <w:multiLevelType w:val="multilevel"/>
""" + "\n".join(
        f"""    <w:lvl w:ilvl="{i}">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="{'.'.join(['%' + str(k + 1) for k in range(i + 1)])}"/>
      <w:lvlJc w:val="left"/>
      <w:pPr>
        <w:ind w:left="{(i + 1) * 425}" w:hanging="425"/>
      </w:pPr>
    </w:lvl>"""
        for i in range(7)
    ) + """
  </w:abstractNum>
  <w:num w:numId="2">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>
"""
    # 优先复用已存在 numbering part (避免 docx zip dup); Document() 默认无该 part
    pkg = doc.part.package
    existing = None
    for p in pkg.iter_parts():
        if p.partname == "/word/numbering.xml":
            existing = p
            break
    if existing is not None:
        existing._blob = numbering_xml.encode("utf-8")
    else:
        part = Part(
            partname=PackURI("/word/numbering.xml"),
            content_type=CONTENT_TYPE.WML_NUMBERING,
            blob=numbering_xml.encode("utf-8"),
            package=pkg,
        )
        doc.part.relate_to(part, RT.NUMBERING)


def _make_heading_para() -> etree._Element:
    """造 Heading2 段 '业绩证明文件': pStyle='3' + numPr (numId=2, ilvl=1)
    + spacing + ind 字面 + 1 个 run rFonts 仅 hint=eastAsia 无字体名。"""
    p = etree.Element(_w("p"), nsmap=NSMAP)
    pPr = etree.SubElement(p, _w("pPr"))

    pStyle = etree.SubElement(pPr, _w("pStyle"))
    pStyle.set(_w("val"), "3")

    numPr = etree.SubElement(pPr, _w("numPr"))
    ilvl = etree.SubElement(numPr, _w("ilvl"))
    ilvl.set(_w("val"), "1")
    numId = etree.SubElement(numPr, _w("numId"))
    numId.set(_w("val"), "2")

    spacing = etree.SubElement(pPr, _w("spacing"))
    spacing.set(_w("before"), "240")
    spacing.set(_w("after"), "120")

    ind = etree.SubElement(pPr, _w("ind"))
    ind.set(_w("firstLine"), "420")

    r = etree.SubElement(p, _w("r"))
    rPr = etree.SubElement(r, _w("rPr"))
    rFonts = etree.SubElement(rPr, _w("rFonts"))
    rFonts.set(_w("hint"), "eastAsia")
    t = etree.SubElement(r, _w("t"))
    t.text = "业绩证明文件"
    return p


def _make_run_no_font(text: str) -> etree._Element:
    """run + rFonts hint=eastAsia 无字体名 (触发 V4-1a.6 字体 fix 补显式字体)。"""
    r = etree.Element(_w("r"), nsmap=NSMAP)
    rPr = etree.SubElement(r, _w("rPr"))
    rFonts = etree.SubElement(rPr, _w("rFonts"))
    rFonts.set(_w("hint"), "eastAsia")
    t = etree.SubElement(r, _w("t"))
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def _make_table() -> etree._Element:
    """9 行 x 7 列, tblStyle='12', 列宽 dxa 7 值, 每 cell 1 段 + 1 run。"""
    tbl = etree.Element(_w("tbl"), nsmap=NSMAP)

    tblPr = etree.SubElement(tbl, _w("tblPr"))
    tblStyle = etree.SubElement(tblPr, _w("tblStyle"))
    tblStyle.set(_w("val"), "12")
    tblW = etree.SubElement(tblPr, _w("tblW"))
    tblW.set(_w("w"), "0")
    tblW.set(_w("type"), "auto")
    tblLook = etree.SubElement(tblPr, _w("tblLook"))
    tblLook.set(_w("val"), "04A0")

    tblGrid = etree.SubElement(tbl, _w("tblGrid"))
    for w in COL_WIDTHS_DXA:
        gc = etree.SubElement(tblGrid, _w("gridCol"))
        gc.set(_w("w"), w)

    all_rows = [HEADER_ROW] + DATA_ROWS
    assert len(all_rows) == 9 and all(len(r) == 7 for r in all_rows), \
        f"行列数错: {len(all_rows)} x {[len(r) for r in all_rows]}"

    for row_data in all_rows:
        tr = etree.SubElement(tbl, _w("tr"))
        for i, cell_text in enumerate(row_data):
            tc = etree.SubElement(tr, _w("tc"))
            tcPr = etree.SubElement(tc, _w("tcPr"))
            tcW = etree.SubElement(tcPr, _w("tcW"))
            tcW.set(_w("w"), COL_WIDTHS_DXA[i])
            tcW.set(_w("type"), "dxa")
            p = etree.SubElement(tc, _w("p"))
            pPr = etree.SubElement(p, _w("pPr"))
            sp = etree.SubElement(pPr, _w("spacing"))
            sp.set(_w("line"), "240")
            sp.set(_w("lineRule"), "auto")
            p.append(_make_run_no_font(cell_text))
    return tbl


def _prune_docx(docx_path: Path) -> None:
    """Post-process: 删 stylesWithEffects.xml + prune styles.xml 仅留用到 + basedOn 链。

    python-docx Document() 起手带 Word 全套 built-in styles (~350KB styles.xml +
    ~438KB stylesWithEffects.xml), fixture 实际只用到 12 / 3 / TableNormal /
    Normal 几个。prune 让 fixture 更接近原 anchor 精简形态 (从 ~840KB 解压
    打到 ~12KB 量级), 公开仓 fixture 是刻意构造的最小样本。
    """
    with zipfile.ZipFile(docx_path) as z:
        parts = {n.filename: z.read(n.filename) for n in z.infolist()
                 if not n.is_dir()}

    document_xml = parts["word/document.xml"]
    numbering_xml = parts.get("word/numbering.xml", b"")

    # 收集 document.xml + numbering.xml 内显式引用的 styleId
    referenced: set[str] = set()
    for blob in (document_xml, numbering_xml):
        text = blob.decode("utf-8")
        for tag in ("pStyle", "tblStyle", "rStyle", "numStyleLink", "styleLink"):
            referenced.update(
                re.findall(r'<w:' + tag + r'\s+w:val="([^"]+)"', text)
            )

    # 解析 styles.xml 建 styleId → element 字典 + 链上引用 (basedOn/linkStyle/next)
    styles_root = etree.fromstring(parts["word/styles.xml"])
    style_by_id: dict[str, etree._Element] = {}
    deps: dict[str, set[str]] = {}
    for s in styles_root.findall(qn("w:style")):
        sid = s.get(qn("w:styleId"))
        if not sid:
            continue
        style_by_id[sid] = s
        d: set[str] = set()
        for tag in ("basedOn", "link", "next"):
            for ref in s.findall(qn(f"w:{tag}")):
                v = ref.get(qn("w:val"))
                if v:
                    d.add(v)
        deps[sid] = d

    # 递归闭包: referenced ∪ deps[referenced] ∪ ...
    # 加默认必保: Normal / DefaultParagraphFont / TableNormal / NoList
    keep = set(referenced) | {"Normal", "DefaultParagraphFont", "TableNormal", "NoList"}
    frontier = set(keep)
    while frontier:
        nxt: set[str] = set()
        for sid in frontier:
            for dep in deps.get(sid, set()):
                if dep not in keep:
                    nxt.add(dep)
        keep |= nxt
        frontier = nxt

    # 删 styles.xml 内不在 keep 的 w:style
    pruned_n = kept_n = 0
    for s in list(styles_root.findall(qn("w:style"))):
        sid = s.get(qn("w:styleId"))
        if sid in keep:
            kept_n += 1
        else:
            styles_root.remove(s)
            pruned_n += 1

    parts["word/styles.xml"] = etree.tostring(
        styles_root, xml_declaration=True, encoding="UTF-8", standalone=True,
    )

    # 删 stylesWithEffects.xml (Word 14+ 副本, 主流 reader 不依赖)
    parts.pop("word/stylesWithEffects.xml", None)

    # 同步删 [Content_Types].xml 内 stylesWithEffects override
    ct_xml = parts["[Content_Types].xml"].decode("utf-8")
    ct_xml = re.sub(
        r'<Override\s+PartName="/word/stylesWithEffects\.xml"[^/]*/>',
        "", ct_xml,
    )
    parts["[Content_Types].xml"] = ct_xml.encode("utf-8")

    # 同步删 word/_rels/document.xml.rels 内 stylesWithEffects rel
    rels_path = "word/_rels/document.xml.rels"
    if rels_path in parts:
        rels_xml = parts[rels_path].decode("utf-8")
        rels_xml = re.sub(
            r'<Relationship[^>]*Target="stylesWithEffects\.xml"[^/]*/>',
            "", rels_xml,
        )
        parts[rels_path] = rels_xml.encode("utf-8")

    # 重写 zip (tmp 放在输出同 dir, 避免 Windows 跨 drive shutil.move 失败)
    fd, tmp_name = tempfile.mkstemp(suffix=".docx", dir=str(docx_path.parent))
    import os
    os.close(fd)
    tmp = Path(tmp_name)
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as z:
        for name, blob in parts.items():
            z.writestr(name, blob)
    docx_path.unlink()
    tmp.rename(docx_path)

    print(f"[prune] styles.xml: kept {kept_n} / dropped {pruned_n} built-in styles; "
          f"stylesWithEffects.xml removed")


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", type=Path, default=DEFAULT_OUT,
                    help=f"输出路径 (默认 {DEFAULT_OUT})")
    args = ap.parse_args()
    out: Path = args.out

    doc = Document()
    _ensure_styles(doc)
    _ensure_numbering(doc)

    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))

    # 删默认 Document() 起手的空段
    for p in list(body.findall(qn("w:p"))):
        body.remove(p)

    heading = _make_heading_para()
    table = _make_table()

    if sectPr is not None:
        sectPr.addprevious(heading)
        sectPr.addprevious(table)
    else:
        body.append(heading)
        body.append(table)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    _prune_docx(out)
    print(f"[OK] 合成 fixture 已写: {out} ({out.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
