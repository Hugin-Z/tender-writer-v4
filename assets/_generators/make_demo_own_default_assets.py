# -*- coding: utf-8 -*-
"""
make_demo_own_default_assets.py · 生成 demo 用占位 asset(V3-1)

为让 demo_cadre_training 项目跑 b_mode_fill 时能演示 CuratedLocalAssetsProvider
真实命中,在 demo 项目的 bidding_entity(own_demo)公司下摄入若干
"DEMO PLACEHOLDER"标记的 docx。

产出:
- 公司资质/own_demo/_raw/20250101_demo_business_license.docx
- 公司资质/own_demo/_raw/20250215_demo_iso9001.docx
- 公司资质/own_demo/_raw/20250420_demo_audit_report.docx
- 公司资质/own_demo/_raw/20250520_demo_tax_social_security.docx
- 团队简历/own_demo/_raw/20250301_demo_zhang_san.docx
- 团队简历/own_demo/_raw/20250315_demo_li_si.docx

每份 docx 内容明确标注"DEMO PLACEHOLDER, NOT REAL ASSET",防止误导用户认为是真实素材。
真实使用时,删除这些 demo docx,通过正常摄入流程入库真实资料。

注意:文件名保留脚本名为 make_demo_own_default_assets.py(commit 5 已入库),
内部 COMPANY_ID 从 own_default 改为 own_demo(对齐 demo 项目实际 bidding_entity)。

运行:
    .venv/Scripts/python.exe assets/_generators/make_demo_own_default_assets.py
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


ASSETS_ROOT = Path(__file__).resolve().parent.parent
COMPANY_ID = "own_demo"  # V3-1: demo 项目的 bidding_entity 是 own_demo,不是 own_default
DEMO_BANNER = "==== DEMO PLACEHOLDER, NOT REAL ASSET ===="


def make_docx(path: Path, title: str, paragraphs: list[str]):
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph(DEMO_BANNER)
    doc.add_paragraph(title)
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.add_paragraph(DEMO_BANNER)
    doc.save(str(path))


def main() -> int:
    qual_raw = ASSETS_ROOT / "公司资质" / COMPANY_ID / "_raw"
    cv_raw = ASSETS_ROOT / "团队简历" / COMPANY_ID / "_raw"

    # 公司资质 ===========================
    make_docx(
        qual_raw / "20250101_demo_business_license.docx",
        "营业执照(DEMO)",
        [
            "统一社会信用代码:DEMO-91XXXXXXXXXXXXXXXX",
            "公司名称:示例投标方 ALPHA(DEMO)",
            "经营范围:信息技术服务、软件开发(DEMO 占位文本,非真实资质)",
            "注册资本:1000 万元",
            "颁证日期:2025-01-01",
            "本文件由 V3-1 demo 生成器产出,真实使用时请走正常摄入流程替换为真实营业执照扫描件。",
        ],
    )
    make_docx(
        qual_raw / "20250215_demo_iso9001.docx",
        "ISO 9001 质量管理体系认证证书(DEMO)",
        [
            "证书号:DEMO-ISO9001-2025",
            "颁证机构:示范认证中心(DEMO)",
            "颁证日期:2025-02-15",
            "有效期至:2028-02-14",
            "适用范围:软件开发与系统集成服务(DEMO 占位)",
            "本文件由 V3-1 demo 生成器产出,真实使用时请替换为真实 ISO 9001 证书扫描件。",
        ],
    )
    make_docx(
        qual_raw / "20250420_demo_audit_report.docx",
        "2024 年度财务审计报告(DEMO)",
        [
            "审计单位:示范会计师事务所(DEMO)",
            "审计基准日:2024-12-31",
            "营业收入:5000 万元(DEMO 数据)",
            "净利润:300 万元(DEMO 数据)",
            "审计意见:无保留意见(DEMO)",
            "本文件由 V3-1 demo 生成器产出,真实使用时请替换为真实审计报告。",
        ],
    )
    make_docx(
        qual_raw / "20250520_demo_tax_social_security.docx",
        "完税与社保缴纳证明(DEMO)",
        [
            "纳税人识别号:DEMO-91XXXXXXXXXXXXXXXX",
            "完税证明月份:2025-04",
            "社保缴纳人数:50 人(DEMO)",
            "社保缴纳基数:DEMO 占位数据",
            "本文件由 V3-1 demo 生成器产出,真实使用时请替换为真实税务/社保证明。",
        ],
    )

    # 团队简历 ===========================
    make_docx(
        cv_raw / "20250301_demo_zhang_san.docx",
        "张三 简历(DEMO)",
        [
            "姓名:张三(DEMO)",
            "职务:项目经理",
            "工作年限:10 年",
            "项目经验:负责 X / Y / Z 三个示例项目(DEMO 占位)",
            "本文件由 V3-1 demo 生成器产出,真实使用时请替换为真实简历。",
        ],
    )
    make_docx(
        cv_raw / "20250315_demo_li_si.docx",
        "李四 简历(DEMO)",
        [
            "姓名:李四(DEMO)",
            "职务:技术骨干",
            "工作年限:8 年",
            "技术专长:后端架构 / 数据库设计(DEMO 占位)",
            "本文件由 V3-1 demo 生成器产出,真实使用时请替换为真实简历。",
        ],
    )

    print(f"[OK] demo own_default fixture 生成完毕")
    print(f"  公司资质:{qual_raw} (4 份)")
    print(f"  团队简历:{cv_raw} (2 份)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
